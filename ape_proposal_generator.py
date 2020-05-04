# -*- coding: utf-8 -*- 
from __future__ import print_function
import os
import re
import random
import numpy
import pandas as pd
import sys
import array
from datetime import date
from mailmerge import MailMerge
import docx
import locale
import pyexcel as pe
from openpyxl import workbook
from openpyxl import load_workbook
from DatabaseTesting import *
locale.setlocale(locale.LC_ALL, 'en_US.utf8')
from docx import Document
from PyQt5 import QtWidgets, uic,QtGui, QtCore
from PyQt5.QtWidgets import QWidget, QListWidget, QListWidgetItem, QLabel, QPushButton, QApplication, QTableWidgetItem, QDesktopWidget, QMainWindow, QFileDialog
from PyQt5.QtGui import  QIcon, QPixmap  #(This would not allow me to import QTableWidgetItem)
from Assay_Class import *
from Products_Class import *
from Unit_Operations_Classes import *
from Useful_Functions import *
from Utilities_Class import *
from Refinery_Class import *
from addingScansunits import *
from Sub_Classes import *
#This is solely for graphing
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure
import mammoth
import shutil
from docxcompose.composer import Composer


# PyQt5 Cheat Sheet TESTTEST
#
# Receiving Input:
#       Text boxes: dlg.__TEXT BOX NAME__.text()
#       Check box/ Radio button: dlg.___BUTTON NAME__.isChecked()
#       Combo box: dlg.__BOX NAME__.currentText()
#       Spin box: dlg.__BOX NAME___.value()
# Changing Value:
#       Text boxes: dlg.__TEXT BOX NAME__.setText(...)
#       Check box/Radio button: dlg.___BUTTON NAME___.setChecked( True/False )
#       Combo box: dlg.___BOX NAME___.addItem(...)
#       List box: 
#           dlg.___LIST NAME__.addItem(...)
#           dlg.__LIST NAME__.clear()
# Connecting functions:
#       Click of Button: dlg.__BUTTON NAME__.clicked.connect(___FUNCTION NAME___)
#       Combo Box Selection: dlg.__BOX NAME__.currentIndexChanged.connect(__FUNCTION NAME___)
#       Radio Button Selected: dlg.___BUTTON NAME___.toggled.connect(___FUNCTION NAME___)
# Variables Cheat Sheet:
#       text boxes: txt_...
#       push button: cmd_...
#       check boxes: chk_...
#       spin boxes: spin_...
#       combo boxes: cmb_...
#       radio button: btn_...


#class wdg_TBPCurve(QWidget):

#    def __init__(self, parent = None):

#        QWidget.__init__(self,parent)

#        self.canvas = FigureCanvas(Figure())

#        vertical_layout = QVBoxLayout()
#        vertical_layout.addWidget(self.Canvas)

#        self.canvas.axes = self.canvas.figure.add_subplot(111)
#        self.setLayout(vertical_layout)


class Canvas(FigureCanvas):
    def __init__(self, parent = None, width = 5, height = 5, dpi = 100):
        fig = Figure(figsize = (width, height), dpi = dpi)
        self.axes = fig.add_subplot(111)

        FigureCanvas.__init__(self, fig)
        self.setParent(parent)
    def plot(self):
        x = numpy.array([50,30,40])
        labels = ["Apples", "Bananas","Melons"]
        ax = self.figure.add_subplot(111)
        ax.pie(x, labels = labels)

class Dashboard:

    def __init__(self, dlg): #The usage of dlg here is what separates the user interface (UI) from the calculation engine. If you want to use this in a different interface. Replace dlg with the variable assignment you please
        self.dlg = dlg
        self.MyDatabase = AssayDatabase()
        self.CustomRefinery = Refinery(True)
        self.OverallRefinery = Refinery(True)
        self.Refinery_List = dict({})
        self.Proposal_Dictionary_List = dict({})
        self.SelectedRefinery = ''
        self.Comments = []
        self.Dictionary = dict({})
        self.scanner = None #refers to the OCR scanner. Look at addingScansunits.py to look at the file associated with the OCR scanner
        self.Temperature_Cut_Point_List = [self.dlg.txt_FuelGasUpperTemperature, self.dlg.txt_LightNaphthaUpperTemperature, self.dlg.txt_HeavyNaphthaUpperTemperature, self.dlg.txt_KeroseneUpperTemperature,
                                           self.dlg.txt_DieselUpperTemperature, self.dlg.txt_AGOUpperTemperature, self.dlg.txt_VGOUpperTemperature] #The user has the power to change these
        #these are the units you choose when filling out the assay
        self.Temperature_Units_List = ['Fahrenheit', 'Celsius', 'Kelvin', 'Rankine']
        self.TAN_Units_List = ['mg KOH/g']
        self.Weight_Percent_Units_List = ['ppm', 'wt%', 'mg/g']
        self.Percentage_Units_List = ['vol%', 'wt%']
        self.Density_Units_List = ['SG', 'API', 'kg/L','kg/m^3', 'lb/gal', 'lb/bbl', 'lb/ft3']
        self.Distance_Units_List = ['mm', 'ft', 'cm', 'm']
        self.Viscosity_Units_List = ['cSt']
        self.Molecular_Weight_Units_List = ['g/mol', 'lb/mol']
        self.Pressure_Units_List = ['kPa', 'psi']
        self.dlg.btn_ocr.setChecked(True)
        self.dlg.cmb_SourceName.addItem('Select Source Name')
        self.dlg.btn_domestic.setChecked(True)
        self.dlg.btn_fixed.setChecked(True)
        self.dlg.spin_commission.setValue(15)
        ###Joe: Begin Settings and Price Index Database Implementation
        self.MyPriceIndexDatabase = PriceIndexDatabase()
        self.IsPriceIndexDateSearch = True
        ###Joe: Added compliment definition study cost and process design study cost with original feasibility study cost created by Adam, also see AssignDashboardVariables, ProposalDictionary, DomesticSelected, InternationalSelected, GenerateClicked for similar additions
        self.dlg.txt_DefinitionStudyCost.setText('75000')
        self.dlg.txt_FeasibilityStudyCost.setText('212750')
        self.dlg.txt_ProcessDesignStudyCost.setText('125000')
        ###Joe: End Settings and Price Index Database Implementation
        self.FixedSelected()
        self.IsRegionSearch = True
        self.IsAssaySearch = False
        self.IsInternational = False
        self.IsDomestic = True
        self.SetTable()
        #hide unecessary parts of the dashboard
        #Hides the Maximized products
        self.dlg.list_MaximizedProducts.hide()
        self.dlg.cmd_MoveUpMaximizedProduct.hide()
        self.dlg.cmd_MoveDownMaximizedProduct.hide()
        self.dlg.cmd_ClearMaximizedProducts.hide()
        self.dlg.lbl_MaximizationHierarchy.hide()
        #Hides the Blinded Side Draws
        self.dlg.list_BlindedSideDraws.hide()
        self.dlg.lbl_BlindedSideDraws.hide()
        self.dlg.cmd_ClearBlindedSideDraws.hide()
        self.dlg.chk_blend.hide()
        self.dlg.txt_blendPercent1.hide()
        self.dlg.lbl_blendPercent1.hide()
        self.dlg.txt_blendPercent2.hide()
        self.dlg.txt_blendPercent3.hide()
        self.dlg.txt_blendPercent4.hide()
        self.dlg.txt_blendPercent5.hide()
        self.dlg.txt_blendPercent6.hide()
        self.dlg.txt_blendPercent7.hide()
        self.dlg.lbl_blendPercent2.hide()
        self.dlg.lbl_blendPercent3.hide()
        self.dlg.lbl_blendPercent4.hide()
        self.dlg.lbl_blendPercent5.hide()
        self.dlg.lbl_blendPercent6.hide()
        self.dlg.lbl_blendPercent7.hide()
        self.dlg.cmb_contact.addItems(['Johnny Hallford', 'Matt Rodgers', 'Joseph DeSpain', 'Other'])
        self.dlg.cmb_companys.addItems(['Chemex Modular, LLC', 'Pan Africa Chemex Limited', 'Chemex Modular India Private Limited'])
        self.dlg.cmb_CustomCurrentlySelectedRefinery.addItems(['Select Refinery'])
        self.dlg.cmb_CurrentlySelectedRefinery.addItems(['Select Refinery'])
        self.dlg.cmb_SettingsSelectedRefinery.addItems(['Select Refinery'])
        self.dlg.lbl_name.hide()
        self.dlg.txt_name.hide()
        self.dlg.lbl_title.hide()
        self.dlg.txt_title.hide()
        self.dlg.lbl_email.hide()
        self.dlg.txt_email.hide()
        self.dlg.lbl_phone.hide()
        self.dlg.txt_phone.hide()
        #Add units to the Assay Page
        self.dlg.cmb_MethaneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_EthaneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_PropaneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_iButaneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_nButaneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_nPentaneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_iPentaneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_CyclopentaneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_CyclohexaneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_BenzeneUnit.addItems(self.Percentage_Units_List)
        self.dlg.cmb_T1Unit.addItems(self.Temperature_Units_List)
        self.dlg.cmb_T2Unit.addItems(self.Temperature_Units_List)
        self.dlg.cmb_T3Unit.addItems(self.Temperature_Units_List)
        self.dlg.cmb_T4Unit.addItems(self.Temperature_Units_List)
        self.dlg.cmb_T5Unit.addItems(self.Temperature_Units_List)
        self.dlg.cmb_IBPUnit.addItems(self.Temperature_Units_List)
        self.dlg.cmb_FBPUnit.addItems(self.Temperature_Units_List)
        #Add Units to the Assay Table
        self.Assay_Table_Column_Labels = ['Initial Temp\n', 'Final Temp\n', 'Weight Yield\n', 'Volume Yield\n', 'Density\n', 'Sulfur\n', 'Mercaptan Sulfur\n', 'Smoke Point\n', 'Viscosity @ T1\n', 'Viscosity @ T2\n',
                                          'Viscosity @ T2\n', 'TAN\n', 'Molecular Weight\n', 'Salt Content\n', 'Flash Point\n', 'Freeze Point\n', 'Cloud Point\n', 'Pour Point\n', 'RON\n', 'Cetane\n', 'RVP\n',
                                          'Aromatics\n', 'Olefins\n', 'Paraffins\n', 'Vanadium\n', 'Nickel\n', 'Nitrogen\n'] #Each of these have different units so we do not iterate through the list here. However, it will be utilized later
        #Add units to our assay table
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Initial Temp\n')).addItems(self.Temperature_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Final Temp\n')).addItems(self.Temperature_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Weight Yield\n')).addItem('wt%')
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Volume Yield\n')).addItem('vol%')
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Density\n')).addItems(self.Density_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Sulfur\n')).addItems(self.Weight_Percent_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Mercaptan Sulfur\n')).addItems(self.Weight_Percent_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Smoke Point\n')).addItems(self.Distance_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Viscosity @ T1\n')).addItems(self.Viscosity_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Viscosity @ T2\n')).addItems(self.Viscosity_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('TAN\n')).addItems(self.TAN_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Molecular Weight\n')).addItems(self.Molecular_Weight_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Salt Content\n')).addItems(['ptb'])
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Flash Point\n')).addItems(self.Temperature_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('RVP\n')).addItems(self.Pressure_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('TVP\n')).addItems(self.Pressure_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Freeze Point\n')).addItems(self.Temperature_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Cloud Point\n')).addItems(self.Temperature_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Aromatics\n')).addItems(self.Percentage_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Olefins\n')).addItems(self.Percentage_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Paraffins\n')).addItems(self.Percentage_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Naphthenes\n')).addItems(self.Percentage_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Vanadium\n')).addItems(self.Weight_Percent_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Nickel\n')).addItems(self.Weight_Percent_Units_List)   
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Pour Point\n')).addItems(self.Temperature_Units_List)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Nitrogen\n')).addItems(self.Weight_Percent_Units_List)

        #The following is in regards to the Input Data To Database Section on the Assay Viewer
        self.SourceNumber = 1
        self.PopupCheck = False

        self.DirectoryFolderName = os.getcwd()
        self.ConnectObjects()# initializes user interface
        self.dlg.cmb_AssayName.addItems(['Select Assay Name'])
        self.ShowSourceOptions()
        #self.UploadInfoFromOCRScanner()
        #self.dlg.setFixedSize(self.dlg.sizeHint())
        return

    def SetTable(self):
        #NOT USED ANYWHERE IN THE CODE. THIS IS SOLELY FOR PRACTICE/DEMONSTRATION PURPOSES ONLY
        self.dlg.tbl_assay.insertRow(0)
        for i in range(0,self.dlg.tbl_assay.columnCount()):
            combobox = QtWidgets.QComboBox()
            self.dlg.tbl_assay.setCellWidget(0,i,combobox)
            self.dlg.tbl_assay.setItem(0,i,QtWidgets.QTableWidgetItem('uwot'))
        self.dlg.tbl_assay.setVerticalHeaderLabels(['Units','Whole Crude'])
        return
    # Products Breakdown on the Dashboard
    def AssignDashboardVariables(self): #This function is not connected to a button
        #Receives information on dashboard (Mostly product info) and assigns it to the currently selected Refinery
        # receives basic information and adds merge values to dictionary
        self.clientCompany = self.dlg.txt_clientCompany.text()
        self.clientName = self.dlg.txt_clientName.text()
        self.clientTitle = self.dlg.txt_clientTitle.text()
        self.companyAbbrev = self.dlg.txt_clientAbbrev.text()
        self.installLocation = self.dlg.txt_clientInstallLocation.text()
        self.DefinitionStudyCost = self.dlg.txt_DefinitionStudyCost.text()
        self.FeasibilityStudyCost = self.dlg.txt_FeasibilityStudyCost.text()
        self.ProcessDesignStudyCost = self.dlg.txt_ProcessDesignStudyCost.text()
        self.commissionValue = self.dlg.spin_commission.value()
        self.Refinery_List[self.SelectedRefinery].CommissionRate = self.dlg.spin_commission.value()
        self.Dictionary.update({'AssayName': self.dlg.cmb_AssayName.currentText(), 'Construction_Time': str(self.dlg.spin_constructionTime.value()), "Delivery_Time": str(self.dlg.spin_startupTime.value()), 'CompanyAbbrev': str(self.dlg.txt_clientAbbrev.text()), 'ClientTitle': self.dlg.txt_clientTitle.text(), 
                          'ClientName': self.clientName, 'ClientCompany': self.clientCompany, "Install_Location": self.installLocation })

        if len(self.Refinery_List) == 0: #Error Checking
            return

        if self.dlg.chk_LPD.isChecked():
            self.VolumeUnit = "Liters per day"
            self.VolumeUnitLabel = "LPD"
        elif self.dlg.chk_GPD.isChecked():
            self.VolumeUnit = "Gallons per day"
            self.VolumeUnitLabel = "gpd"
        elif self.dlg.chk_MTPA.isChecked():
            self.VolumeUnit = "Metric tons per annum"
            self.VolumeUnitLabel = "MTPA"
        else:
            self.VolumeUnit = None
            self.VolumeUnitLabel = "bpd"
        self.Dictionary.update({'VolumeUnitLabel': self.VolumeUnitLabel })
        try: 
            if self.dlg.txt_totalCapacity.text() != '':
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.CapacityValue = int(self.dlg.txt_totalCapacity.text())
            else:
                # QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please Input Capacity', QtWidgets.QMessageBox.Ok)
                return True
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please enter a correct value for capacity', QtWidgets.QMessageBox.Ok)

        # products to maximize
        self.maximizedProducts = []
        for i in range(0, self.dlg.list_MaximizedProducts.count()):
            self.maximizedProducts.append(self.dlg.list_MaximizedProducts.item(i).text()) 

        #Blinded Side Draws
        if len(self.dlg.list_BlindedSideDraws.findItems('Kerosene',QtCore.Qt.MatchExactly)) > 0: self.Refinery_List[self.SelectedRefinery].ProductsInstance.KeroseneBlindedSideDrawCheck = True
        if len(self.dlg.list_BlindedSideDraws.findItems('Jet A',QtCore.Qt.MatchExactly)) > 0: self.Refinery_List[self.SelectedRefinery].ProductsInstance.JetABlindedSideDrawCheck = True
        if len(self.dlg.list_BlindedSideDraws.findItems('Diesel',QtCore.Qt.MatchExactly)) > 0: self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselBlindedSideDrawCheck = True
        if len(self.dlg.list_BlindedSideDraws.findItems('ULS Diesel',QtCore.Qt.MatchExactly)) > 0: self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselBlindedSideDrawCheck = True
        if len(self.dlg.list_BlindedSideDraws.findItems('HS Diesel',QtCore.Qt.MatchExactly)) > 0: self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselBlindedSideDrawCheck = True
        if len(self.dlg.list_BlindedSideDraws.findItems('LSMDO',QtCore.Qt.MatchExactly)) > 0: self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOBlindedSideDrawCheck = True
        if len(self.dlg.list_BlindedSideDraws.findItems('MDO',QtCore.Qt.MatchExactly)) > 0: self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOBlindedSideDrawCheck = True
        if len(self.dlg.list_BlindedSideDraws.findItems('AGO',QtCore.Qt.MatchExactly)) > 0: self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOBlindedSideDrawCheck = True
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.NumBlindedSideDraws = self.dlg.list_BlindedSideDraws.count()

        # index of assays to test (then use database to find assay information) (Not sure what this is)
        self.assaysToTest = []
        for j in range(0, self.dlg.list_assays.count()):
            self.assaysToTest.append(self.dlg.list_assays.item(j).text()[(self.dlg.list_assays.item(j).text()).find(',') + 1 : len(self.dlg.list_assays.item(j).text())]) 

        # desired products group box
        self.IDNum = self.dlg.txt_IDNum.text()
        if self.dlg.chk_FuelGas.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasFuelGas = True
        if self.dlg.chk_LPG.isChecked():
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLPG = True
        if self.dlg.chk_AGO.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasAGO = True
        if self.dlg.chk_ATB.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasATB = True
        if self.dlg.chk_LightNaphtha.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLightNaphtha = True
        if self.dlg.chk_HeavyNaphtha.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasHeavyNaphtha = True
        if self.dlg.chk_SRNaphtha.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasSRNaphtha = True
        if self.dlg.chk_Gasoline.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasGasoline = True
        if self.dlg.chk_Kerosene.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasKerosene = True
        if self.dlg.chk_JetA.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasJetA = True
        if self.dlg.chk_Diesel.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasDiesel = True
        if self.dlg.chk_HSDiesel.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasHSDiesel = True
        if self.dlg.chk_ULSDiesel.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasULSDiesel = True
        if self.dlg.chk_MDO.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasMDO = True
        if self.dlg.chk_LSMDO.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLSMDO = True
        if self.dlg.chk_LVGO.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLVGO = True
        if self.dlg.chk_HVGO.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasHVGO = True
        if self.dlg.chk_VTB.isChecked(): 
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasVTB = True
        
        if to_str(self.dlg.txt_crudeUnits.text()) != '':
            try:
                self.Refinery_List[self.SelectedRefinery].NumberOfSpecifiedCDU = int(self.dlg.txt_crudeUnits.text())
            except:
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please enter number for Number of Crude Units', QtWidgets.QMessageBox.Ok)
                raise 
                return
        for i in range(0,len(self.Temperature_Cut_Point_List)):
            try:    
                if self.Temperature_Cut_Point_List[i].text() != '':
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.Specified_Cut_Temps[i][1] = float(self.Temperature_Cut_Point_List[i].text())
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.Specified_Cut_Temps[i+1][0] = float(self.Temperature_Cut_Point_List[i].text())
                    #print("&*(^*(&^&*^&*^&*^&*^&*^&*(^^&^&^&^&^&^&^&^&^&^&^&^&^&" + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.Specified_Cut_Temps))
            except:
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please enter numbers for cut point values', QtWidgets.QMessageBox.Ok)
                raise
                return
        return
    def ClearDashboardVariables(self):
        #Clears everything off the dashboard
        self.dlg.txt_FuelGasBPD.clear()
        self.dlg.txt_FuelGasLPD.clear()
        self.dlg.txt_FuelGasGPD.clear()
        self.dlg.txt_FuelGasMTPA.clear()

        self.dlg.txt_LPGBPD.clear()
        self.dlg.txt_LPGLPD.clear()
        self.dlg.txt_LPGGPD.clear()
        self.dlg.txt_LPGMTPA.clear()

        self.dlg.txt_LightNaphthaBPD.clear()
        self.dlg.txt_LightNaphthaLPD.clear()
        self.dlg.txt_LightNaphthaGPD.clear()
        self.dlg.txt_LightNaphthaMTPA.clear()

        self.dlg.txt_HeavyNaphthaBPD.clear()
        self.dlg.txt_HeavyNaphthaLPD.clear()
        self.dlg.txt_HeavyNaphthaGPD.clear()
        self.dlg.txt_HeavyNaphthaMTPA.clear()

        self.dlg.txt_SRNaphthaBPD.clear()
        self.dlg.txt_SRNaphthaLPD.clear()
        self.dlg.txt_SRNaphthaGPD.clear()
        self.dlg.txt_SRNaphthaMTPA.clear()

        self.dlg.txt_GasolineBPD.clear()
        self.dlg.txt_GasolineLPD.clear()
        self.dlg.txt_GasolineGPD.clear()
        self.dlg.txt_GasolineMTPA.clear()

        self.dlg.txt_KeroseneBPD.clear()
        self.dlg.txt_KeroseneLPD.clear()
        self.dlg.txt_KeroseneGPD.clear()
        self.dlg.txt_KeroseneMTPA.clear()

        self.dlg.txt_JetABPD.clear()
        self.dlg.txt_JetALPD.clear()
        self.dlg.txt_JetAGPD.clear()
        self.dlg.txt_JetAMTPA.clear()

        self.dlg.txt_HSDieselBPD.clear()
        self.dlg.txt_HSDieselLPD.clear()
        self.dlg.txt_HSDieselGPD.clear()
        self.dlg.txt_HSDieselMTPA.clear()

        self.dlg.txt_DieselBPD.clear()
        self.dlg.txt_DieselLPD.clear()
        self.dlg.txt_DieselGPD.clear()
        self.dlg.txt_DieselMTPA.clear()

        self.dlg.txt_ULSDieselBPD.clear()
        self.dlg.txt_ULSDieselLPD.clear()
        self.dlg.txt_ULSDieselGPD.clear()
        self.dlg.txt_ULSDieselMTPA.clear()

        self.dlg.txt_AGOBPD.clear()
        self.dlg.txt_AGOLPD.clear()
        self.dlg.txt_AGOGPD.clear()
        self.dlg.txt_AGOMTPA.clear()

        self.dlg.txt_ATBBPD.clear()
        self.dlg.txt_ATBLPD.clear()
        self.dlg.txt_ATBGPD.clear()
        self.dlg.txt_ATBMTPA.clear()

        self.dlg.txt_MDOBPD.clear()
        self.dlg.txt_MDOLPD.clear()
        self.dlg.txt_MDOGPD.clear()
        self.dlg.txt_MDOMTPA.clear()

        self.dlg.txt_LSMDOBPD.clear()
        self.dlg.txt_LSMDOLPD.clear()
        self.dlg.txt_LSMDOGPD.clear()
        self.dlg.txt_LSMDOMTPA.clear()

        self.dlg.txt_LVGOBPD.clear()
        self.dlg.txt_LVGOLPD.clear()
        self.dlg.txt_LVGOGPD.clear()
        self.dlg.txt_LVGOMTPA.clear()

        self.dlg.txt_HVGOBPD.clear()
        self.dlg.txt_HVGOLPD.clear()
        self.dlg.txt_HVGOGPD.clear()
        self.dlg.txt_HVGOMTPA.clear()

        self.dlg.txt_VTBBPD.clear()
        self.dlg.txt_VTBLPD.clear()
        self.dlg.txt_VTBGPD.clear()
        self.dlg.txt_VTBMTPA.clear()

        self.dlg.chk_FuelGas.setChecked(False)
        self.dlg.chk_LPG.setChecked(False)
        self.dlg.chk_LightNaphtha.setChecked(False)
        self.dlg.chk_HeavyNaphtha.setChecked(False)
        self.dlg.chk_SRNaphtha.setChecked(False)
        self.dlg.chk_Gasoline.setChecked(False)
        self.dlg.chk_Kerosene.setChecked(False)
        self.dlg.chk_JetA.setChecked(False)
        self.dlg.chk_Diesel.setChecked(False)
        self.dlg.chk_ULSDiesel.setChecked(False)
        self.dlg.chk_HSDiesel.setChecked(False)
        self.dlg.chk_LSMDO.setChecked(False)
        self.dlg.chk_MDO.setChecked(False)
        self.dlg.chk_AGO.setChecked(False)
        self.dlg.chk_ATB.setChecked(False)
        self.dlg.chk_LVGO.setChecked(False)
        self.dlg.chk_HVGO.setChecked(False)
        self.dlg.chk_VTB.setChecked(False)

        self.dlg.txt_ProductsBreakdownTotal.clear()
        self.dlg.txt_ProductsBreakdownCost.clear()
        self.dlg.txt_totalCapacity.clear()
        self.dlg.list_specialComments.clear()
        self.dlg.txt_RONMinimum.setText('')

        #Clears the Boiling Point Ranges off the Sandbox/Blending Manager Tabs
        self.dlg.txt_SettingsIBP.clear()
        for i in range(0,len(self.Temperature_Cut_Point_List)):
            self.Temperature_Cut_Point_List[i].clear()
        self.dlg.txt_SettingsFBP.clear()
        return
    def FillOutDashboardVariables(self):
        #Takes the currently selected refinery and fills out the dashboard so that the user can visually see what data that refinery is holding

        if len(self.Refinery_List) == 0: #Error Checking
            return False

        self.dlg.txt_FuelGasBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.FuelGasBPD))
        self.dlg.txt_FuelGasLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.FuelGasLPD))
        self.dlg.txt_FuelGasGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.FuelGasGPD))
        self.dlg.txt_FuelGasMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.FuelGasMTPA))

        self.dlg.txt_LPGBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LPGBPD))
        self.dlg.txt_LPGLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LPGLPD))
        self.dlg.txt_LPGGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LPGGPD))
        self.dlg.txt_LPGMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LPGMTPA))

        self.dlg.txt_LightNaphthaBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LightNaphthaBPD))
        self.dlg.txt_LightNaphthaLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LightNaphthaLPD))
        self.dlg.txt_LightNaphthaGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LightNaphthaGPD))
        self.dlg.txt_LightNaphthaMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LightNaphthaMTPA))

        self.dlg.txt_HeavyNaphthaBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HeavyNaphthaBPD))
        self.dlg.txt_HeavyNaphthaLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HeavyNaphthaLPD))
        self.dlg.txt_HeavyNaphthaGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HeavyNaphthaGPD))
        self.dlg.txt_HeavyNaphthaMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HeavyNaphthaMTPA))

        self.dlg.txt_SRNaphthaBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.SRNaphthaBPD))
        self.dlg.txt_SRNaphthaLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.SRNaphthaLPD))
        self.dlg.txt_SRNaphthaGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.SRNaphthaGPD))
        self.dlg.txt_SRNaphthaMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.SRNaphthaMTPA))

        self.dlg.txt_GasolineBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.GasolineBPD))
        self.dlg.txt_GasolineLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.GasolineLPD))
        self.dlg.txt_GasolineGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.GasolineGPD))
        self.dlg.txt_GasolineMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.GasolineMTPA))

        self.dlg.txt_KeroseneBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.KeroseneBPD))
        self.dlg.txt_KeroseneLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.KeroseneLPD))
        self.dlg.txt_KeroseneGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.KeroseneGPD))
        self.dlg.txt_KeroseneMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.KeroseneMTPA))

        self.dlg.txt_JetABPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.JetABPD))
        self.dlg.txt_JetALPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.JetALPD))
        self.dlg.txt_JetAGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.JetAGPD))
        self.dlg.txt_JetAMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.JetAMTPA))

        self.dlg.txt_HSDieselBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselBPD))
        self.dlg.txt_HSDieselLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselLPD))
        self.dlg.txt_HSDieselGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselGPD))
        self.dlg.txt_HSDieselMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselMTPA))

        self.dlg.txt_DieselBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselBPD))
        self.dlg.txt_DieselLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselLPD))
        self.dlg.txt_DieselGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselGPD))
        self.dlg.txt_DieselMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselMTPA))

        self.dlg.txt_ULSDieselBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselBPD))
        self.dlg.txt_ULSDieselLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselLPD))
        self.dlg.txt_ULSDieselGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselGPD))
        self.dlg.txt_ULSDieselMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselMTPA))

        self.dlg.txt_AGOBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOBPD))
        self.dlg.txt_AGOLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOLPD))
        self.dlg.txt_AGOGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOGPD))
        self.dlg.txt_AGOMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOMTPA))

        self.dlg.txt_ATBBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.ATBBPD))
        self.dlg.txt_ATBLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.ATBLPD))
        self.dlg.txt_ATBGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.ATBGPD))
        self.dlg.txt_ATBMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.ATBMTPA))

        self.dlg.txt_MDOBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOBPD))
        self.dlg.txt_MDOLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOLPD))
        self.dlg.txt_MDOGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOGPD))
        self.dlg.txt_MDOMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOMTPA))

        self.dlg.txt_LSMDOBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOBPD))
        self.dlg.txt_LSMDOLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOLPD))
        self.dlg.txt_LSMDOGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOGPD))
        self.dlg.txt_LSMDOMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOMTPA))

        self.dlg.txt_LVGOBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LVGOBPD))
        self.dlg.txt_LVGOLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LVGOLPD))
        self.dlg.txt_LVGOGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LVGOGPD))
        self.dlg.txt_LVGOMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.LVGOMTPA))

        self.dlg.txt_HVGOBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HVGOBPD))
        self.dlg.txt_HVGOLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HVGOLPD))
        self.dlg.txt_HVGOGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HVGOGPD))
        self.dlg.txt_HVGOMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HVGOMTPA))

        self.dlg.txt_VTBBPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.VTBBPD))
        self.dlg.txt_VTBLPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.VTBLPD))
        self.dlg.txt_VTBGPD.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.VTBGPD))
        self.dlg.txt_VTBMTPA.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.VTBMTPA))

        self.dlg.chk_FuelGas.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasFuelGas)
        self.dlg.chk_LPG.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLPG)
        self.dlg.chk_LightNaphtha.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLightNaphtha)
        self.dlg.chk_HeavyNaphtha.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasHeavyNaphtha)
        self.dlg.chk_SRNaphtha.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasSRNaphtha)
        self.dlg.chk_Gasoline.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasGasoline)
        self.dlg.chk_Kerosene.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasKerosene)
        self.dlg.chk_JetA.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasJetA)
        self.dlg.chk_Diesel.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasDiesel)
        self.dlg.chk_ULSDiesel.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasULSDiesel)
        self.dlg.chk_HSDiesel.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasHSDiesel)
        self.dlg.chk_LSMDO.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLSMDO)
        self.dlg.chk_MDO.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasMDO)
        self.dlg.chk_AGO.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasAGO)
        self.dlg.chk_ATB.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasATB)
        self.dlg.chk_LVGO.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLVGO)
        self.dlg.chk_HVGO.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasHVGO)
        self.dlg.chk_VTB.setChecked(self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasVTB)

        self.dlg.txt_RONMinimum.setText(special_format(self.Refinery_List[self.SelectedRefinery].ProductsInstance.GasolineRON))
        self.dlg.txt_ProductsBreakdownTotal.setText(special_format(self.Refinery_List[self.SelectedRefinery].TotalBPD))
        if self.IsDomestic:
            self.dlg.txt_ProductsBreakdownCost.setText("$ " + special_format(self.Refinery_List[self.SelectedRefinery].DomesticTotalPrice))
        elif self.IsInternational:
            self.dlg.txt_ProductsBreakdownCost.setText("$ " + special_format(self.Refinery_List[self.SelectedRefinery].OverseasTotalPrice))
        self.dlg.txt_totalCapacity.setText(to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.CapacityValue))
        self.dlg.spin_commission.setValue(self.Refinery_List[self.SelectedRefinery].CommissionRate)
        self.dlg.list_specialComments.clear()
        for comment in self.Refinery_List[self.SelectedRefinery].SpecialConsiderations:
            self.Comments.append(comment)
            self.dlg.list_specialComments.addItem('* ' + comment)

        self.ClearBlindedSideDraws()
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.KeroseneBlindedSideDrawCheck == True: self.KeroseneBlindedSideDraw()
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.JetABlindedSideDrawCheck == True: self.JetABlindedSideDraw()
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselBlindedSideDrawCheck == True: self.DieselBlindedSideDraw()
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselBlindedSideDrawCheck == True: self.ULSDieselBlindedSideDraw()
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselBlindedSideDrawCheck == True: self.HSDieselBlindedSideDraw()
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOBlindedSideDrawCheck == True: self.LSMDOBlindedSideDraw()
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOBlindedSideDrawCheck == True: self.MDOBlindedSideDraw()
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOBlindedSideDrawCheck == True: self.AGOBlindedSideDraw()

        # Sets the Boiling Point Cut Temperatures on the Blending Manager/Sandbox tab
        self.dlg.txt_SettingsIBP.setText(to_str(round(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBP,2)))
        self.dlg.txt_SettingsFBP.setText(to_str(round(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FBP,2)))
        for i in range(0,len(self.Temperature_Cut_Point_List)):
            self.Temperature_Cut_Point_List[i].setText(str(round(self.Refinery_List[self.SelectedRefinery].ProductsInstance.Specified_Cut_Temps[i][1],2)))
        return
    def ProductsCheckBoxCheck(self):
        #Checks to make sure all of the selections are legal and will not mess up the products calculations. An output of an error message will appear if this check fails
        if self.dlg.chk_FuelGas.isChecked() and self.dlg.chk_LPG.isChecked():
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please only select one off-gas product', QtWidgets.QMessageBox.Ok)
            return False
        if self.dlg.chk_LightNaphtha.isChecked() or self.dlg.chk_HeavyNaphtha.isChecked():
            if not (self.dlg.chk_LightNaphtha.isChecked() and self.dlg.chk_HeavyNaphtha.isChecked()):
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'If choosing to split Naphtha, you must select both Light and Heavy Naphtha', QtWidgets.QMessageBox.Ok)
                return False
        if not(self.dlg.chk_LightNaphtha.isChecked() or self.dlg.chk_HeavyNaphtha.isChecked() or self.dlg.chk_Gasoline.isChecked() or self.dlg.chk_SRNaphtha.isChecked()):
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please Select a Naphtha Cut', QtWidgets.QMessageBox.Ok)
            return False
        if self.dlg.chk_SRNaphtha.isChecked():
            if self.dlg.chk_LightNaphtha.isChecked() or self.dlg.chk_HeavyNaphtha.isChecked() or self.dlg.chk_Gasoline.isChecked():
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'If choosing SR Naphtha, you are not allowed to select another Naphtha Cut', QtWidgets.QMessageBox.Ok)
                return False
        if self.dlg.chk_Gasoline.isChecked():
            if self.dlg.chk_LightNaphtha.isChecked() or self.dlg.chk_HeavyNaphtha.isChecked() or self.dlg.chk_SRNaphtha.isChecked():
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'If choosing Gasoline, you are not allowed to select another Naphtha Cut', QtWidgets.QMessageBox.Ok)
                return False
        if self.dlg.chk_Kerosene.isChecked() or self.dlg.chk_JetA.isChecked():
            if self.dlg.chk_Kerosene.isChecked() and self.dlg.chk_JetA.isChecked():
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please Deselect one of your Kerosene range cuts', QtWidgets.QMessageBox.Ok)
                return False
        if self.dlg.chk_Diesel.isChecked() or self.dlg.chk_HSDiesel.isChecked() or self.dlg.chk_ULSDiesel.isChecked():
            CheckNumber = 0
            if self.dlg.chk_Diesel.isChecked():  CheckNumber += 1
            if self.dlg.chk_ULSDiesel.isChecked():  CheckNumber += 1
            if self.dlg.chk_HSDiesel.isChecked():  CheckNumber += 1
            if CheckNumber > 1:
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'If choosing a Diesel, you are not allowed to select more than 1 Diesel Cut', QtWidgets.QMessageBox.Ok)
                return False
        if self.dlg.chk_LSMDO.isChecked() or self.dlg.chk_MDO.isChecked() or self.dlg.chk_AGO.isChecked() or self.dlg.chk_LVGO.isChecked():
            CheckNumber = 0
            if self.dlg.chk_LSMDO.isChecked():  CheckNumber += 1
            if self.dlg.chk_MDO.isChecked():  CheckNumber += 1
            if self.dlg.chk_AGO.isChecked():  CheckNumber += 1
            if self.dlg.chk_LVGO.isChecked():  CheckNumber += 1
            if CheckNumber > 1:
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'If choosing a Light/Medium Gas Oil, you are not allowed to select more than 1 Gas Oil Cut', QtWidgets.QMessageBox.Ok)
                return False
        if self.dlg.chk_ATB.isChecked():
            if self.dlg.chk_LVGO.isChecked() or self.dlg.chk_HVGO.isChecked() or self.dlg.chk_VTB.isChecked():
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'If choosing Atmospheric Tower Bottoms, you are not allowed to select a Vacuum Cut', QtWidgets.QMessageBox.Ok)
                return False
        if self.dlg.chk_VTB.isChecked() or self.dlg.chk_HVGO.isChecked():
            if not (self.dlg.chk_VTB.isChecked() and self.dlg.chk_HVGO.isChecked()):
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'If choosing a Vacuum Cut, you must select both HVGO and VTB', QtWidgets.QMessageBox.Ok)
                return False
            if not (self.dlg.chk_LSMDO.isChecked() or self.dlg.chk_MDO.isChecked() or self.dlg.chk_LVGO.isChecked()):
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'If choosing a Vacuum Cut, you must select both HVGO and a Lighter Vacuum Cut (LVGO or an MDO)', QtWidgets.QMessageBox.Ok)
                return False
        if not (self.dlg.chk_VTB.isChecked() or self.dlg.chk_ATB.isChecked()):
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'You must choose a Bottoms product', QtWidgets.QMessageBox.Ok)
            return False
        return True
    def CreateDictionary(self):
        
        #Creates the dictionary for merging of the word file. This is the last step before usingthe MailMerge library to automate the generation of a proposal

        if len(self.Refinery_List) == 0: #Error Checking
            return
        for RefineryUnit in self.Refinery_List.values():
            if RefineryUnit.ProductsInstance.IsCalculated == False or RefineryUnit.IsCalculated == False:
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please make sure that all refineries are calculated before attempting to make Dictionary for Merge Fields', QtWidgets.QMessageBox.Ok)
                return
        #We must reset our overall refinery or it will stack the calculations
        self.OverallRefinery.ReInitialize()

        today = date.today()
        Year = to_str(today.strftime('%Y')) #Assigned Correctly
        Month = to_str(today.strftime('%m')) #Assigned Correctly
        Day = to_str(today.strftime('%d')) #Assigned Correctly

        SRUType = "Merichem Lo-Cat" #Not sure on this
        SentBy = self.dlg.cmb_contact.currentText() #Assigned Correctly       
        #InstallLocation = "Puntas de Mata, Monagas States, Venezuela" #Easily Assigned
        ReferenceNum = self.dlg.txt_ReferenceNumber.text() # Correctly assigned (for now)
        RevNumber = self.dlg.txt_RevisionNumber.text() # Correctly Assigned (for now)
        Company = self.dlg.cmb_companys.currentText() #Assigned Correctly
        Octane = to_str(92.9) #Need formula/equation to do this (assigned correctly for now)
        InternationalRate = 0.15

        #This is to find out the largest commission value present so we can accurately determine the feasibility study cost
        for RefineryUnit in self.Refinery_List.values():
            if RefineryUnit.CommissionRate >= self.commissionValue:
                self.commissionValue = RefineryUnit.CommissionRate
         
        KeroMercaptan = None
        KeroSmokePoint = None
        #This next bit of code does not have any current use because kerosene smoke point and mercaptans are not reported in the proposal generator
        #for i in range (0,len(self.MyProducts.Product_Names)):
        #    if self.MyProducts.Product_Names[i] == "Kerosene" or self.MyProducts.Product_Names[i] == "Jet A" or self.MyProducts.Product_Names[i] == "ULS Kerosene":
        #        KeroMercaptan = round_up(self.MyAssay.Mercaptan_Data[i],2)
        #        KeroSmokePoint = round_up(self.MyAssay.Smoke_Point_Data[i],2)
        #        break

        if len(self.CustomRefinery.Specified_Unit_List) >= 1: # We need to add this later for when customers have their own specified refineries costs
            DictionaryUnitList = self.CustomRefinery.Specified_Unit_List

        #Value holders for Maximum Capacities to report on the final proposal
        MaxDesalterCapacity = 0
        MaxDesalterAmount = 0
        MaxDesalterPrice = 0
        MaxCDUCapacity = 0
        MaxCDUAmount = 0
        MaxCDUPrice = 0
        MaxNumSideDraw = 0
        MaxVDUCapacity = 0
        MaxVDUAmount = 0
        MaxVDUPrice = 0
        MaxNHTCapacity = 0
        MaxNHTAmount = 0
        MaxNHTPrice = 0
        MaxNaphthaSplitterCapacity = 0
        MaxNaphthaSplitterAmount = 0
        MaxNaphthaSplitterPrice = 0
        MaxDHTCapacity = 0
        MaxDHTAmount = 0
        MaxDHTPrice = 0
        MaxDistillateSplitterCapacity = 0
        MaxDistillateSplitterAmount = 0
        MaxDistillateSplitterPrice = 0
        MaxCRUCapacity = 0
        MaxCRUAmount = 0
        MaxCRUPrice = 0
        MaxStabilizerCapacity = 0
        MaxStabilizerAmount = 0
        MaxStabilizerPrice = 0
        MaxKHTCapacity = 0
        MaxKHTAmount = 0
        MaxKHTPrice = 0
        MaxCausticTreaterCapacity = 0
        MaxCausticTreaterAmount = 0
        MaxCausticTreaterPrice = 0
        MaxIsomCapacity = 0
        MaxIsomAmount = 0
        MaxIsomPrice = 0
        MaxAdditionalAlloyCapacity = 0
        MaxAdditionalAlloyPrice = 0
        MaxThiolexCapacity = 0
        MaxThiolexAmount = 0
        MaxThiolexPrice = 0
        MaxLoCatCapacity = 0
        MaxLoCatAmount = 0
        MaxLoCatPrice = 0

        #Values for merge field of the name of relevant unit operations
        DesalterUnit = ''
        CrudeDistillationUnit = ''
        VacuumDistillationUnit = ''
        NaphthaStabilizerUnit = ''
        NaphthaSplitterUnit = ''
        DieselHydrotreaterUnit = ''
        NaphthaHydrotreaterUnit = ''
        KeroseneHydrotreaterUnit = ''
        DistillateSplitterUnit = ''
        CatalyticReformingUnit = ''
        Isomerization_Unit = '' #The underscore is to prevent us overwriting the isomerization unit class
        MerichemLoCatUnit = ''
        MerichemThiolexUnit = ''
        ARA_HCU = ''
        ARA_HC = ''
        CausticTreaterUnit = ''



        #Value holder for maximum utilities to be reported on the final proposal
        MaxTotalPrice = 0
        MaxUnits = 0
        MaxSteam = 0
        MaxPower = 0 #kWh/Day
        MaxCoolingWater = 0 #Gallons per day
        MaxCoolingTowerPower = 0 #kWh/day
        MaxDesalterWater = 0 #Gallons per day
        MaxFuel = 0 #MMBtu/day
        MaxSCFH2pBBL = 0 #Standard Cubic feet of hydrogen per barrel of hydrotreated oil
        MaxOperators = 0 #2 per unit op
        #The following list is made so we can calculate the price of the overall refinery. This will be treated as a custom refinery and the units will be calculated using CalculateTotalUtilities()
        Max_Cost_Unit_Op_List = [Desalter(), CDU(), VDU(), NaphthaHydrotreater(), RefluxedSplitter(), DieselHydrotreater(), RefluxedSplitter(), CRU(), RefluxedStabilizer(), KeroseneHydrotreater(), CausticTreater(), IsomerizationUnit(), AdditionalAlloy(), MerichemThiolex(), MerichemLoCat() ] #These will contain the Maxed out unit operations
        #Now we initialize the cost:
        NaphthaSplitterCheck = False #We want one of the splitters tp be a naphtha splitter
        DistillateSplitterCheck = False #We want one of the splitters to be a distillate splitter
        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List):
            MaxedUnitOp.OverseasTotalPrice = 0
            MaxedUnitOp.DomesticTotalPrice = 0
            if MaxedUnitOp.Name == "Refluxed Splitter":
                if MaxedUnitOp.NaphthaCheck == False and MaxedUnitOp.DistillateCheck == False and NaphthaSplitterCheck == False:
                    Max_Cost_Unit_Op_List[i].NaphthaCheck = True
                    NaphthaSplitterCheck = True
                elif MaxedUnitOp.DistillateCheck == False and MaxedUnitOp.NaphthaCheck == False and DistillateSplitterCheck == False:
                    Max_Cost_Unit_Op_List[i].DistillateCheck = True
                    DistillateSplitterCheck = True

        if self.CustomRefinery.IsCalculated == False: #If we have custom units that have been specified, we will not consider the CHEMEX branded refinery/refineries
            RefineryList = self.Refinery_List.values()
        else:
            RefineryList = [self.CustomRefinery]
        #For loop will begin here
        for MyRefinery in RefineryList:
        #for i in range(0, len(self.MyProducts.Product_Names)):
        #    if not(self.MyProducts.Product_Names[i] == "Fuel Gas" or self.MyProducts.Product_Names[i] == "LPG"):
        #        TotalBPD = self.MyProducts.CapacityValue

            TotalBPD = MyRefinery.TotalBPD #This value could be different for each crude potentially
            Capacity = MyRefinery.ProductsInstance.CapacityValue #Loop inside to see maximum capacity value
            #OperatorsAndMaintenanceWorkers = to_str(22) # For now this is 2 operators per process Unit
            for i in range(0,len(MyRefinery.ProductsInstance.Product_Names)):
                if MyRefinery.ProductsInstance.Product_Names[i] == "Gasoline":
                    Unit = "Hydroskimming Refinery"
                    self.files_list=[os.getcwd() + '\\Proposal Templates\\generation\\' + 'Proposal_ExhibitB_Cover.docx', os.getcwd() + '\\Proposal Templates\\generation\\' + 'Generated_Proposal_ExhibitB_Assay.docx', os.getcwd() + '\\Proposal Templates\\generation\\' + 'Proposal_ExhibitC_HYDSKIM.docx']
                    break
                else:
                    Unit = "Crude Distillation Unit" # If there is no gasoline this should default to "Crude Distillation Unit"
                    self.files_list=[os.getcwd() + '\\Proposal Templates\\generation\\' + 'Proposal_ExhibitB_Cover.docx', os.getcwd() + '\\Proposal Templates\\generation\\' + 'Generated_Proposal_ExhibitB_Assay.docx', os.getcwd() + '\\Proposal Templates\\generation\\' + 'Proposal_ExhibitC_CDU.docx']
                    
            AssayName = MyRefinery.ProductsInstance.AssayInstance.AssayName #Easily Assigned

            DesalterAmount = None
            CDUAmount = None
            VDUAmount = None
            NHTAmount = None
            NaphthaSplitterAmount = None
            DHTAmount = None
            DistillateSplitterAmount = None
            CRUAmount = None
            StabilizerAmount = None
            KHTAmount = None
            ARAHCUAmount = None # This is not currently implemented
            ARACHAmount = None # This is not currently implemented
            SMRAmount = None
            SNRAmount = None
            ThiolexAmount = None
            LoCatAmount = None
            GasPlantAmount = None
            CausticTreaterAmount = None
            IsomAmount = None

            DesalterCapacity = None
            CDUCapacity = None
            VDUCapacity = None
            NHTCapacity = None
            NaphthaSplitterCapacity = None
            DHTCapacity = None
            DistillateSplitterCapacity = None
            CRUCapacity = None
            StabilizerCapacity = None
            KHTCapacity = None
            ARAHCUCapacity = None # This is not currently implemented
            ARACHCapacity = None # This is not currently implemented
            AdditionalAlloyCapacity = None
            SMRCapacity = None
            SNRCapacity = None
            ThiolexCapacity = None
            LoCatCapacity = None
            GasPlantCapacity = None
            CausticTreaterCapacity = None
            IsomCapacity = None
            NumSideDraw = None

            DesalterPrice = None
            CDUPrice = None
            VDUPrice = None
            NHTPrice = None
            NaphthaSplitterPrice = None
            DHTPrice = None
            DistillateSplitterPrice = None
            CRUPrice = None
            StabilizerPrice = None
            KHTPrice = None
            ARAHCUPrice = None # This is not currently implemented
            ARACHPrice = None # This is not currently implemented
            AdditionalAlloyPrice = None
            SMRPrice = None
            SNRPrice = None
            ThiolexPrice = None
            LoCatPrice = None
            GasPlantPrice = None
            CausticTreaterPrice = None
            IsomPrice = None

        #This goes in Proposal Generator Class
            if self.IsDomestic:
                if MyRefinery.DomesticTotalPrice > MaxTotalPrice:
                    MaxTotalPrice = MyRefinery.DomesticTotalPrice
                AnnualSpares = 0.02 * MyRefinery.DomesticTotalPrice
            elif self.IsInternational:
                if MyRefinery.OverseasTotalPrice > MaxTotalPrice:
                    MaxTotalPrice = MyRefinery.OverseasTotalPrice
                AnnualSpares = 0.02 * MyRefinery.OverseasTotalPrice
            if MyRefinery.TotalUtilities.Steam > MaxSteam:
                MaxSteam = MyRefinery.TotalUtilities.Steam
            if MyRefinery.TotalUtilities.Power > MaxPower:
                MaxPower = MyRefinery.TotalUtilities.Power
            if MyRefinery.TotalUtilities.CoolingWater > MaxCoolingWater:
                MaxCoolingWater = MyRefinery.TotalUtilities.CoolingWater
            if MyRefinery.TotalUtilities.CoolingTowerPower > MaxCoolingTowerPower:
                MaxCoolingTowerPower = MyRefinery.TotalUtilities.CoolingTowerPower
            if MyRefinery.TotalUtilities.DesalterWater > MaxDesalterWater:
                MaxDesalterWater = MyRefinery.TotalUtilities.DesalterWater
            if MyRefinery.TotalUtilities.Fuel > MaxFuel:
                MaxFuel = MyRefinery.TotalUtilities.Fuel
            if MyRefinery.SCFH2pBBL > MaxSCFH2pBBL:
                MaxSCFH2pBBL = MyRefinery.SCFH2pBBL
            if len(MyRefinery.Unit_List) > MaxUnits:
                MaxUnits = len(MyRefinery.Unit_List)
                MaxOperators = 2*MaxUnits #Just a preliminary estimate
            

            for UnitOp in MyRefinery.Unit_List:
                if UnitOp.Name == "Desalter":
                    DesalterCapacity = round_up(UnitOp.InletFlow,-2)
                    DesalterAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        DesalterPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        DesalterPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    DesalterUnit = 'desalter'
                    if DesalterCapacity*DesalterAmount > MaxDesalterCapacity*MaxDesalterAmount:
                        MaxDesalterCapacity = DesalterCapacity
                        MaxDesalterAmount = DesalterAmount
                        MaxDesalterPrice = DesalterPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #Inneficient. To make this more efficient we can make this into a Dictionary 
                            if MaxedUnitOp.Name == "Desalter":
                                if MaxDesalterCapacity * MaxDesalterAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Crude Distillation Unit":
                    CDUCapacity = round_up(UnitOp.InletFlow,-2)
                    CDUAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        CDUPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        CDUPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    NumSideDraw = UnitOp.NumSideDraws - 3
                    CrudeDistillationUnit = 'crude distillation unit'
                    if CDUCapacity*CDUAmount > MaxCDUCapacity*MaxCDUAmount:
                        MaxCDUCapacity = CDUCapacity
                        MaxCDUAmount = CDUAmount
                        MaxCDUPrice = CDUPrice
                        MaxNumSideDraw = NumSideDraw
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Crude Distillation Unit":
                                if MaxCDUCapacity * MaxCDUAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Vacuum Distillation Unit":
                    VDUCapacity = round_up(UnitOp.InletFlow,-2)
                    VDUAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        VDUPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        VDUPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    VacuumDistillationUnit = 'vacuum distillation unit'
                    if VDUCapacity*VDUAmount > MaxVDUCapacity*MaxVDUAmount:
                        MaxVDUCapacity = VDUCapacity
                        MaxVDUAmount = VDUAmount
                        MaxVDUPrice = VDUPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Vacuum Distillation Unit":
                                if MaxVDUCapacity * MaxVDUAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Naphtha Hydrotreater":
                    NHTCapacity = round_up(UnitOp.InletFlow,-2)
                    NHTAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        NHTPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        NHTPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    NaphthaHydrotreaterUnit = 'naphtha hydrotreater'
                    if NHTCapacity*NHTAmount > MaxNHTCapacity*MaxNHTAmount:
                        MaxNHTCapacity = NHTCapacity
                        MaxNHTAmount = NHTAmount
                        MaxNHTPrice = NHTPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Naphtha Hydrotreater":
                                if MaxNHTCapacity * MaxNHTAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Refluxed Splitter" and UnitOp.NaphthaCheck == True:
                    NaphthaSplitterCapacity = round_up(UnitOp.InletFlow,-2)
                    NaphthaSplitterAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        NaphthaSplitterPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        NaphthaSplitterPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    NaphthaSplitterUnit = 'naphtha splitter'
                    if NaphthaSplitterCapacity*NaphthaSplitterAmount > MaxNaphthaSplitterCapacity*MaxNaphthaSplitterAmount:
                        MaxNaphthaSplitterCapacity = NaphthaSplitterCapacity
                        MaxNaphthaSplitterAmount = NaphthaSplitterAmount
                        MaxNaphthaSplitterPrice = NaphthaSplitterPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Refluxed Splitter" and MaxedUnitOp.NaphthaCheck == True:
                                if MaxNaphthaSplitterCapacity * MaxNaphthaSplitterAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Diesel Hydrotreater":
                    DHTCapacity = round_up(UnitOp.InletFlow,-2)
                    DHTAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        DHTPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        DHTPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    DieselHydrotreaterUnit = 'diesel hydrotreater'
                    if DHTCapacity*DHTAmount > MaxDHTCapacity*MaxDHTAmount:
                        MaxDHTCapacity = DHTCapacity
                        MaxDHTAmount = DHTAmount
                        MaxDHTPrice = DHTPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Diesel Hydrotreater":
                                if MaxDHTCapacity * MaxDHTAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Refluxed Splitter" and UnitOp.DistillateCheck == True:
                    DistillateSplitterCapacity = round_up(UnitOp.InletFlow,-2)
                    DistillateSplitterAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        DistillateSplitterPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        DistillateSplitterPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    DistillateSplitterUnit = 'distillate splitter'
                    if DistillateSplitterCapacity*DistillateSplitterAmount > MaxDistillateSplitterCapacity*MaxDistillateSplitterAmount:
                        MaxDistillateSplitterCapacity = DistillateSplitterCapacity
                        MaxDistillateSplitterAmount = DistillateSplitterAmount
                        MaxDistillateSplitterPrice = DistillateSplitterPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Refluxed Splitter" and MaxedUnitOp.DistillateCheck == True:
                                if MaxDistillateSplitterCapacity * MaxDistillateSplitterAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Catalytic Reforming Unit":
                    CRUCapacity = round_up(UnitOp.InletFlow,-2)
                    CRUAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        CRUPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        CRUPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    CatalyticReformingUnit = 'catalytic reforming unit'
                    if CRUCapacity*CRUAmount > MaxCRUCapacity*MaxCRUAmount:
                        MaxCRUCapacity = CRUCapacity
                        MaxCRUAmount = CRUAmount
                        MaxCRUPrice = CRUPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Catalytic Reforming Unit":
                                if MaxCRUCapacity * MaxCRUAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Refluxed Stabilizer":
                    StabilizerCapacity = round_up(UnitOp.InletFlow,-2)
                    StabilizerAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        StabilizerPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        StabilizerPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    NaphthaStabilizerUnit = 'naphtha stabilizer'
                    if StabilizerCapacity*StabilizerAmount > MaxStabilizerCapacity*MaxStabilizerAmount:
                        MaxStabilizerCapacity = StabilizerCapacity
                        MaxStabilizerAmount = StabilizerAmount
                        MaxStabilizerPrice = StabilizerPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Refluxed Stabilizer":
                                if MaxStabilizerCapacity * MaxStabilizerAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Kerosene Hydrotreater":
                    KHTCapacity = round_up(UnitOp.InletFlow,-2)
                    KHTAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        KHTPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        KHTPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    KeroseneHydrotreaterUnit = "kerosene hydrotreater"
                    if KHTCapacity*KHTAmount > MaxKHTCapacity*MaxKHTAmount:
                        MaxKHTCapacity = KHTCapacity
                        MaxKHTAmount = KHTAmount
                        MaxKHTPrice = KHTPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Kerosene Hydrotreater":
                                if MaxKHTCapacity * MaxKHTAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Caustic Treater":
                    CausticTreaterCapacity = round_up(UnitOp.InletFlow,-2)
                    CausticTreaterAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        CausticTreaterPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        CausticTreaterPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    CausticTreaterUnit = 'caustic treater'
                    if CausticTreaterCapacity*CausticTreaterAmount > MaxCausticTreaterCapacity*MaxCausticTreaterAmount:
                        MaxCausticTreaterCapacity = CausticTreaterCapacity
                        MaxCausticTreaterAmount = CausticTreaterAmount
                        MaxCausticTreaterPrice = CausticTreaterPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Caustic Treater":
                                if MaxCausticTreaterCapacity * MaxCausticTreaterAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Isomerization Unit":
                    IsomCapacity = round_up(UnitOp.InletFlow,-2)
                    IsomAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        IsomPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    if self.IsInternational: 
                        IsomPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    Isomerization_Unit = 'isomerization unit'
                    if IsomCapacity*IsomAmount > MaxIsomCapacity*MaxIsomAmount:
                        MaxIsomCapacity = IsomCapacity
                        MaxIsomAmount = IsomAmount
                        MaxIsomPrice = IsomPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Isomerization Unit":
                                if MaxIsomCapacity * MaxIsomAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Additional Alloy":
                    AdditionalAlloyCapacity = round_up(UnitOp.InletFlow,-2)
                    if self.IsDomestic:
                        AdditionalAlloyPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        AdditionalAlloyPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    if AdditionalAlloyCapacity > MaxAdditionalAlloyCapacity:
                        MaxAdditionalAlloyCapacityCapacity = AdditionalAlloyCapacity
                        MaxAdditionalAlloyPrice = AdditionalAlloyPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Additional Alloy":
                                if MaxAdditionalAlloyCapacity > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Merichem Thiolex Unit":
                    ThiolexCapacity = round_up(UnitOp.InletFlow,2)
                    ThiolexAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        ThiolexPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        ThiolexPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    MerichemThiolexUnit = 'Merichem Thiolex unit'
                    if ThiolexCapacity*ThiolexAmount > MaxThiolexCapacity*MaxThiolexAmount:
                        MaxThiolexCapacity = ThiolexCapacity
                        MaxThiolexAmount = ThiolexAmount
                        MaxThiolexPrice = ThiolexPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Merichem Thiolex Unit":
                                if MaxThiolexCapacity * MaxThiolexAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break
                if UnitOp.Name == "Merichem Lo-Cat Unit":
                    LoCatCapacity = round_up(UnitOp.InletFlow,2)
                    LoCatAmount = UnitOp.NumUnits
                    if self.IsDomestic:
                        LoCatPrice = round_up(UnitOp.DomesticTotalPrice,-1)
                    elif self.IsInternational:
                        LoCatPrice = round_up(UnitOp.OverseasTotalPrice,-1)
                    MerichemLoCatUnit = 'Merichem LoCat unit'
                    if LoCatCapacity*LoCatAmount > MaxLoCatCapacity*MaxLoCatAmount:
                        MaxLoCatCapacity = LoCatCapacity
                        MaxLoCatAmount = LoCatAmount
                        MaxLoCatPrice = LoCatPrice
                        for i, MaxedUnitOp in enumerate(Max_Cost_Unit_Op_List): #We iterate through our maxed unit ops to replace old ones with higher values if they exist
                            if MaxedUnitOp.Name == "Merichem Lo-Cat Unit":
                                if MaxLoCatCapacity * MaxLoCatAmount > MaxedUnitOp.TotalFlow:
                                    Max_Cost_Unit_Op_List[i] = UnitOp
                                break

            #Goes inside the for loop
            OffGasBPD = None
            if MyRefinery.ProductsInstance.LPGBPD != None:
                OffGasBPD = MyRefinery.ProductsInstance.LPGBPD
            if MyRefinery.ProductsInstance.FuelGasBPD != None:
                OffGasBPD = MyRefinery.ProductsInstance.FuelGasBPD

            ProposalDictionary = dict({})

            ProposalDictionary.update({"FuelGas_Yield": to_str(MyRefinery.ProductsInstance.FuelGasNewYield), "LPG_Yield": to_str(MyRefinery.ProductsInstance.LPGNewYield), "LSRNaphtha_Yield": to_str(MyRefinery.ProductsInstance.LightNaphthaNewYield), "SRNaphtha_Yield": to_str(MyRefinery.ProductsInstance.SRNaphthaNewYield), "HSRNaphtha_Yield": to_str(MyRefinery.ProductsInstance.HeavyNaphthaNewYield), "Gasoline_Yield": to_str(MyRefinery.ProductsInstance.GasolineNewYield)})
            ProposalDictionary.update({"SRKerosene_Yield": to_str(MyRefinery.ProductsInstance.KeroseneNewYield), "ULSKerosene_Yield": to_str(MyRefinery.ProductsInstance.ULSKeroseneNewYield), "JetA_Yield": to_str(MyRefinery.ProductsInstance.JetANewYield), "SRDiesel_Yield": to_str(MyRefinery.ProductsInstance.DieselNewYield), "ULSDiesel_Yield": to_str(MyRefinery.ProductsInstance.ULSDieselNewYield), "HSDiesel_Yield": to_str(MyRefinery.ProductsInstance.HSDieselNewYield)})
            ProposalDictionary.update({"AGO_Yield": to_str(MyRefinery.ProductsInstance.AGONewYield), "ATB_Yield": to_str(MyRefinery.ProductsInstance.ATBNewYield), "MDO_Yield": to_str(MyRefinery.ProductsInstance.MDONewYield), "LSMDO_Yield": to_str(MyRefinery.ProductsInstance.LSMDONewYield), "LVGO_Yield": to_str(MyRefinery.ProductsInstance.LVGONewYield), "HVGO_Yield": to_str(MyRefinery.ProductsInstance.HVGONewYield), "VTB_Yield": to_str(MyRefinery.ProductsInstance.VTBNewYield)})
            ProposalDictionary.update({"LVGO_Yield": to_str(MyRefinery.ProductsInstance.LVGONewYield), "FuelGas_SCFPD": special_format(round_up(special_multiply(OffGasBPD,SCFPerBarrel),2)), "LPG_BPD": to_str(MyRefinery.ProductsInstance.LPGBPD), "LSRNaphtha_BPD": to_str(MyRefinery.ProductsInstance.LightNaphthaBPD), "SRNaphtha_BPD": to_str(MyRefinery.ProductsInstance.SRNaphthaBPD), "HSRNaphtha_BPD": to_str(MyRefinery.ProductsInstance.HeavyNaphthaBPD)})
            ProposalDictionary.update({"Gasoline_BPD": to_str(MyRefinery.ProductsInstance.GasolineBPD), "SRKerosene_BPD": to_str(MyRefinery.ProductsInstance.KeroseneBPD), "ULSKerosene_BPD": to_str(MyRefinery.ProductsInstance.ULSKeroseneBPD), "JetA_BPD": to_str(MyRefinery.ProductsInstance.JetABPD), "SRDiesel_BPD": to_str(MyRefinery.ProductsInstance.DieselBPD), "ULSDiesel_BPD": to_str(MyRefinery.ProductsInstance.ULSDieselBPD), "HSDiesel_BPD": to_str(MyRefinery.ProductsInstance.HSDieselBPD)})
            ProposalDictionary.update({"AGO_BPD": to_str(MyRefinery.ProductsInstance.AGOBPD), "ATB_BPD": to_str(MyRefinery.ProductsInstance.ATBBPD), "MDO_BPD": to_str(MyRefinery.ProductsInstance.MDOBPD), "LSMDO_BPD": to_str(MyRefinery.ProductsInstance.LSMDOBPD), "HVGO_BPD": to_str(MyRefinery.ProductsInstance.HVGOBPD), "VTB_BPD": to_str(MyRefinery.ProductsInstance.VTBBPD), "LVGO_BPD": to_str(MyRefinery.ProductsInstance.LVGOBPD)})
            #This data is if a specific unit is picked for the volume flow (LPD or GPD)
            ProposalDictionary.update({"Vol": self.VolumeUnit, "Fuel Gas": special_format(MyRefinery.ProductsInstance.FuelGasFlow.get(self.VolumeUnit)), "LPG": special_format(MyRefinery.ProductsInstance.LPGFlow.get(self.VolumeUnit)), "LSRNaphtha": special_format(MyRefinery.ProductsInstance.LightNaphthaFlow.get(self.VolumeUnit)), "HSRNaphtha": special_format(MyRefinery.ProductsInstance.HeavyNaphthaFlow.get(self.VolumeUnit)), "SRNaphtha": special_format(MyRefinery.ProductsInstance.SRNaphthaFlow.get(self.VolumeUnit))})
            ProposalDictionary.update({"Gasoline": special_format(MyRefinery.ProductsInstance.GasolineFlow.get(self.VolumeUnit)), "SRKerosene": special_format(MyRefinery.ProductsInstance.KeroseneFlow.get(self.VolumeUnit)), "ULSKerosene": special_format(MyRefinery.ProductsInstance.ULSKeroseneFlow.get(self.VolumeUnit)), "JetA": special_format(MyRefinery.ProductsInstance.JetAFlow.get(self.VolumeUnit))})
            ProposalDictionary.update({"SRDiesel": special_format(MyRefinery.ProductsInstance.DieselFlow.get(self.VolumeUnit)), "ULSDiesel": special_format(MyRefinery.ProductsInstance.ULSDieselFlow.get(self.VolumeUnit)), "HSDiesel": special_format(MyRefinery.ProductsInstance.HSDieselFlow.get(self.VolumeUnit)), "AGO": special_format(MyRefinery.ProductsInstance.AGOFlow.get(self.VolumeUnit)), "ATB": special_format(MyRefinery.ProductsInstance.ATBFlow.get(self.VolumeUnit)), "MDO": special_format(MyRefinery.ProductsInstance.MDOFlow.get(self.VolumeUnit))})
            ProposalDictionary.update({"LSMDO": special_format(MyRefinery.ProductsInstance.LSMDOFlow.get(self.VolumeUnit)), "LVGO": special_format(MyRefinery.ProductsInstance.LVGOFlow.get(self.VolumeUnit)), "HVGO": special_format(MyRefinery.ProductsInstance.HVGOFlow.get(self.VolumeUnit)), "VTB": special_format(MyRefinery.ProductsInstance.VTBFlow.get(self.VolumeUnit)), "Vol_Total": special_format(MyRefinery.ProductsInstance.TotalFlow.get(self.VolumeUnit))})
            ProposalDictionary.update({"Yield_Total": to_str('100.0 %'), "BPD_Total": "{:,}".format(TotalBPD),})
            #This is the Pricing of each Unit on each refinery which goes for each individual assay. This will not be reported on the Proposal
            ProposalDictionary.update({"Desalt_Price": special_format(DesalterPrice), "CDU_Price": special_format(CDUPrice), "VDU_Price": special_format(VDUPrice), "NHT_Price": special_format(NHTPrice), "NapSp_Price": special_format(NaphthaSplitterPrice), "DHT_Price": special_format(DHTPrice)})
            ProposalDictionary.update({"DistSp_Price": special_format(DistillateSplitterPrice), "CRU_Price": special_format(CRUPrice), "Stab_Price": special_format(StabilizerPrice), "KHT_Price": special_format(KHTPrice), "ARAHCU_Price": special_format(ARAHCUPrice), "ARACH_Price": special_format(ARACHPrice)})
            ProposalDictionary.update({"Additional_Alloy_Price": special_format(AdditionalAlloyPrice), "SMR_Price": special_format(SMRPrice), "SNR_Price": special_format(SNRPrice), "Thiolex_Price": to_str(ThiolexPrice), "LoCat_Price": to_str(LoCatPrice), "GasPlant_Price": special_format(GasPlantPrice), "Caustic_Price": special_format(CausticTreaterPrice)})
            ProposalDictionary.update({"Isom_Price": special_format(IsomPrice)})
            #This is the Utilities of each Unit for each refinery which goes for each individual assay. This will not be reported on the Proposal
            ProposalDictionary.update({"Unit_Octane": to_str(Octane), "Unit_Steam_lb": special_format(MyRefinery.TotalUtilities.Steam), "Unit_Power_kwh": special_format(MyRefinery.TotalUtilities.Power), "Unit_Cooling_H2O": special_format(MyRefinery.TotalUtilities.CoolingWater)})
            ProposalDictionary.update({"Unit_Cool_Twr": special_format(MyRefinery.TotalUtilities.CoolingTowerPower), "Unit_Desalt_H2O": special_format(MyRefinery.TotalUtilities.DesalterWater), "Unit_Fuel_MMbtu": special_format(MyRefinery.TotalUtilities.Fuel), "Unit_Operators": to_str(len(MyRefinery.Unit_List) * 2), "Unit_Annual_Spares": special_format(AnnualSpares)})
            ProposalDictionary.update({"Unit_TAN#": None, "Unit_Kero_Mercaptan": to_str(KeroMercaptan), "Unit_Kero_Smoke_Point": to_str(KeroSmokePoint), "AssayName": AssayName, "Unit_Hydrogen_scfb": special_format(MyRefinery.SCFH2pBBL)})
            #These are the Utilites Capacity and amounts
            ProposalDictionary.update({"Unit_Desalt_Capacity": special_format(DesalterCapacity), "Unit_CDU_Capacity": special_format(CDUCapacity), "Unit_VDU_Capacity": special_format(VDUCapacity), "Unit_NHT_Capacity": special_format(NHTCapacity), "Unit_NapSp_Capacity": special_format(NaphthaSplitterCapacity), "Unit_DHT_Capacity": special_format(DHTCapacity)})
            ProposalDictionary.update({"Unit_DistSp_Capacity": special_format(DistillateSplitterCapacity), "Unit_CRU_Capacity": special_format(CRUCapacity), "Unit_Stab_Capacity": special_format(StabilizerCapacity), "Unit_KHT_Capacity": special_format(KHTCapacity), "Unit_ARAHCU_Capacity": special_format(ARAHCUCapacity), "Unit_ARACH_Capacity": special_format(ARACHCapacity)})
            ProposalDictionary.update({"Unit_Additional_Alloy_Cap": special_format(AdditionalAlloyCapacity), "Unit_SMR_Capacity": special_format(SMRCapacity), "Unit_SNR_Capacity": special_format(SNRCapacity), "Unit_Thiolex_Capacity": to_str(ThiolexCapacity), "Unit_LoCat_Capacity": to_str(LoCatCapacity), "Unit_GasPlant_Capacity": special_format(GasPlantCapacity), "Unit_Caustic_Capacity": special_format(CausticTreaterCapacity)})
            ProposalDictionary.update({"Unit_Isom_Capacity": special_format(IsomCapacity), "Unit_Num_SideDraw": to_str(NumSideDraw), "Unit_Desalt_Amt": to_str(DesalterAmount), "Unit_CDU_Amt": to_str(CDUAmount), "Unit_VDU_Amt": to_str(VDUAmount), "Unit_NHT_Amt": to_str(NHTAmount), "Unit_NapSp_Amt": to_str(NaphthaSplitterAmount), "Unit_DHT_Amt": to_str(DHTAmount)})
            ProposalDictionary.update({"Unit_DistSp_Amt": to_str(DistillateSplitterAmount), "Unit_CRU_Amt": to_str(CRUAmount), "Unit_Stab_Amt": to_str(StabilizerAmount), "Unit_KHT_Amt": to_str(KHTAmount), "Unit_ARAHCU_Amt": to_str(ARAHCUAmount), "Unit_ARACH_Amt": to_str(ARACHAmount), "Unit_SMR_Amt": to_str(SMRAmount), "Unit_SNR_Amt": to_str(SNRAmount)})
            ProposalDictionary.update({"Unit_Thiolex_Amt": to_str(ThiolexAmount), "Unit_LoCat_Amt": to_str(LoCatAmount), "Unit_GasPlant_Amt": to_str(GasPlantAmount), "Unit_Caustic_Amt": to_str(CausticTreaterAmount), "Unit_Isom_Amt": to_str(IsomAmount), "Unit_Type": to_str(SRUType)})

            self.Proposal_Dictionary_List.update({MyRefinery.ProductsInstance.AssayInstance.AssayID: ProposalDictionary.copy()})
            print("LOOK HERE *************************" + str(self.Proposal_Dictionary_List))
            print("This is the Volume Unit:" + str(self.VolumeUnit))

        #After the for loop
        index = 0
        for MaxedUnitOp in Max_Cost_Unit_Op_List:
            print("This is the Total Flow of the unit operation " + to_str(MaxedUnitOp.Name))
            print(MaxedUnitOp.TotalFlow)
            if MaxedUnitOp.TotalFlow > 0:
                self.OverallRefinery.Specified_Unit_List.append(MaxedUnitOp)

        print(self.OverallRefinery.Specified_Unit_List)

        self.OverallRefinery.CalculateTotalUtilities()
        if self.IsDomestic:
            MaxTotalPrice = self.OverallRefinery.DomesticTotalPrice
            MaxAnnualSpares = round_up(0.02 * MaxTotalPrice) #Max Total Price
        elif self.IsInternational:
            MaxTotalPrice = self.OverallRefinery.OverseasTotalPrice
            MaxAnnualSpares = round_up(0.02 * MaxTotalPrice) #Max Total Price
        #These are the Utilities and Unit information that will be calculated beforehand to see which one is maximum so that we can do multiple assays
        ProposalDictionary.clear()
        
        ProposalDictionary.update({"Octane": to_str(zero_to_none(Octane)), "Steam_lb": special_format(zero_to_none(MaxSteam)), "Power_kwh": special_format(zero_to_none(MaxPower)), "Cooling_H2O": special_format(zero_to_none(MaxCoolingWater))})
        ProposalDictionary.update({"Cool_Twr": special_format(zero_to_none(MaxCoolingTowerPower)), "Desalt_H2O": special_format(zero_to_none(MaxDesalterWater)), "Fuel_MMbtu": special_format(zero_to_none(MaxFuel)), "Operators": to_str(zero_to_none(MaxOperators)), "Annual_Spares": special_format(zero_to_none(MaxAnnualSpares))})
        ProposalDictionary.update({"TAN#": None, "Kero_Mercaptan": to_str(zero_to_none(KeroMercaptan)), "Kero_Smoke_Point": to_str(zero_to_none(KeroSmokePoint)), "Hydrogen_scfb": special_format(zero_to_none(MaxSCFH2pBBL))})

        ProposalDictionary.update({"Desalter": to_str(DesalterUnit), "CDU": to_str(CrudeDistillationUnit), "VDU": to_str(VacuumDistillationUnit), "NHT": to_str(NaphthaHydrotreaterUnit), "Naph_Splitter": to_str(NaphthaSplitterUnit), "DHT": to_str(DieselHydrotreaterUnit)})
        ProposalDictionary.update({"Dist_Split": to_str(DistillateSplitterUnit), "CRU": to_str(CatalyticReformingUnit), "Stabilizer": to_str(NaphthaStabilizerUnit), "KHT": to_str(KeroseneHydrotreaterUnit), "ARA_HCU": '', "ARA_CH": '' })
        ProposalDictionary.update({"Additional_Alloy": '', "SMR": '', "SNR": '', "Merichem_Thiolex": to_str(MerichemThiolexUnit), "Merichem_LoCat": to_str(MerichemLoCatUnit), "Gas_Plant": '', "Caustic": to_str(CausticTreaterUnit)})
        ProposalDictionary.update({"Isom_Unit": to_str(Isomerization_Unit)})

        ProposalDictionary.update({"Desalt_Capacity": special_format(zero_to_none(MaxDesalterCapacity)), "CDU_Capacity": special_format(zero_to_none(MaxCDUCapacity)), "VDU_Capacity": special_format(zero_to_none(MaxVDUCapacity)), "NHT_Capacity": special_format(zero_to_none(MaxNHTCapacity)), "NapSp_Capacity": special_format(zero_to_none(MaxNaphthaSplitterCapacity)), "DHT_Capacity": special_format(zero_to_none(MaxDHTCapacity))})
        ProposalDictionary.update({"DistSp_Capacity": special_format(zero_to_none(MaxDistillateSplitterCapacity)), "CRU_Capacity": special_format(zero_to_none(MaxCRUCapacity)), "Stab_Capacity": special_format(zero_to_none(MaxStabilizerCapacity)), "KHT_Capacity": special_format(zero_to_none(MaxKHTCapacity)), "ARAHCU_Capacity": special_format(zero_to_none(ARAHCUCapacity)), "ARACH_Capacity": special_format(zero_to_none(ARACHCapacity))})
        ProposalDictionary.update({"Additional_Alloy_Cap": special_format(zero_to_none(MaxAdditionalAlloyCapacity)), "SMR_Capacity": special_format(zero_to_none(SMRCapacity)), "SNR_Capacity": special_format(zero_to_none(SNRCapacity)), "Thiolex_Capacity": to_str(zero_to_none(MaxThiolexCapacity)), "LoCat_Capacity": to_str(zero_to_none(MaxLoCatCapacity)), "GasPlant_Capacity": special_format(zero_to_none(GasPlantCapacity)), "Caustic_Capacity": special_format(zero_to_none(MaxCausticTreaterCapacity))})
        ProposalDictionary.update({"Isom_Capacity": special_format(zero_to_none(MaxIsomCapacity)), "Num_SideDraw": to_str(zero_to_none(MaxNumSideDraw)), "Desalt_Amt": to_str(zero_to_none(MaxDesalterAmount)), "CDU_Amt": to_str(zero_to_none(MaxCDUAmount)), "VDU_Amt": to_str(zero_to_none(MaxVDUAmount)), "NHT_Amt": to_str(zero_to_none(MaxNHTAmount)), "NapSp_Amt": to_str(zero_to_none(MaxNaphthaSplitterAmount)), "DHT_Amt": to_str(zero_to_none(MaxDHTAmount))})
        ProposalDictionary.update({"DistSp_Amt": to_str(zero_to_none(MaxDistillateSplitterAmount)), "CRU_Amt": to_str(zero_to_none(MaxCRUAmount)), "Stab_Amt": to_str(zero_to_none(MaxStabilizerAmount)), "KHT_Amt": to_str(zero_to_none(MaxKHTAmount)), "ARAHCU_Amt": to_str(zero_to_none(ARAHCUAmount)), "ARACH_Amt": to_str(ARACHAmount), "SMR_Amt": to_str(zero_to_none(SMRAmount)), "SNR_Amt": to_str(zero_to_none(SNRAmount))})
        ProposalDictionary.update({"Thiolex_Amt": to_str(zero_to_none(MaxThiolexAmount)), "LoCat_Amt": to_str(zero_to_none(MaxLoCatAmount)), "GasPlant_Amt": to_str(zero_to_none(GasPlantAmount)), "Caustic_Amt": to_str(zero_to_none(MaxCausticTreaterAmount)), "Isom_Amt": to_str(zero_to_none(MaxIsomAmount)), "Type": to_str(SRUType)})

        ProposalDictionary.update({"Total_Price": "{:,}".format(MaxTotalPrice), "Unit": Unit, "ClientName": self.clientName, "SENT_BY": SentBy })
        ProposalDictionary.update({"Capacity": to_str(Capacity), "ReferenceNum": to_str(ReferenceNum), "Year": to_str(Year)})
        ProposalDictionary.update({"Month": Month, "Day": to_str(Day), "Rev#": to_str(RevNumber), "Company": Company, "Def_Cost": to_str(self.DefinitionStudyCost), "Feas_Cost": to_str(self.FeasibilityStudyCost), "PDP_Cost": to_str(self.ProcessDesignStudyCost), "Commission_Rate": to_str(self.commissionValue)})

        self.Dictionary.update(ProposalDictionary.copy())
        #print(ProposalDictionary)
        return
    def FillOutCHEMEXUnitVariables(self):
        #This function is called after all of the products and refineries have been calculated. All it does is show a physical display of the unit pricing and info on the Process Units page
        if len(self.Refinery_List) == 0: #Error Checking
            return
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.IsCalculated != True:
            return
        self.dlg.txt_desalterAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].DesalterAmount))
        self.dlg.txt_desalterCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].DesalterCapacity))
        self.dlg.txt_desalterPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].DesalterPrice))
        self.dlg.txt_CDUAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].CDUAmount))
        self.dlg.txt_CDUCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].CDUCapacity))
        self.dlg.txt_CDUPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].CDUPrice))
        self.dlg.txt_sideDrawsAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].SideDrawAmount))
        self.dlg.txt_sideDrawsCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].SideDrawCapacity))
        self.dlg.txt_sideDrawsPriceCHEMEX.setText(special_format(special_multiply(self.Refinery_List[self.SelectedRefinery].SideDrawPrice, self.Refinery_List[self.SelectedRefinery].SideDrawAmount)))
        self.dlg.txt_BlindedSideDrawsAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].BlindedSideDrawAmount))
        self.dlg.txt_BlindedSideDrawsCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].BlindedSideDrawCapacity))
        self.dlg.txt_BlindedSideDrawsPriceCHEMEX.setText(special_format(special_multiply(self.Refinery_List[self.SelectedRefinery].BlindedSideDrawPrice, self.Refinery_List[self.SelectedRefinery].BlindedSideDrawAmount)))
        self.dlg.txt_VDUAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].VDUAmount))
        self.dlg.txt_VDUCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].VDUCapacity))
        self.dlg.txt_VDUPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].VDUPrice))
        self.dlg.txt_NHTAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].NHTAmount))
        self.dlg.txt_NHTCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].NHTCapacity))
        self.dlg.txt_NHTPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].NHTPrice))
        self.dlg.txt_NaphSpAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].NaphthaSplitterAmount))
        self.dlg.txt_NaphSpCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].NaphthaSplitterCapacity))
        self.dlg.txt_NaphSpPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].NaphthaSplitterPrice))
        self.dlg.txt_DHTAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].DHTAmount))
        self.dlg.txt_DHTCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].DHTCapacity))
        self.dlg.txt_DHTPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].DHTPrice))
        self.dlg.txt_distSpAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].DistillateSplitterAmount))
        self.dlg.txt_distSpCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].DistillateSplitterCapacity))
        self.dlg.txt_distSpPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].DistillateSplitterPrice))
        self.dlg.txt_CRUAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].CRUAmount))
        self.dlg.txt_CRUCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].CRUCapacity))
        self.dlg.txt_CRUPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].CRUPrice))
        self.dlg.txt_gasPlantAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].GasPlantAmount))
        self.dlg.txt_gasPlantCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].GasPlantCapacity))
        self.dlg.txt_gasPlantPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].GasPlantPrice))
        self.dlg.txt_isomAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].IsomAmount))
        self.dlg.txt_isomCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].IsomCapacity))
        self.dlg.txt_isomPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].IsomPrice))
        self.dlg.txt_stabilizerAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].StabilizerAmount))
        self.dlg.txt_stabilizerCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].StabilizerCapacity))
        self.dlg.txt_stabilizerPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].StabilizerPrice))
        self.dlg.txt_KHTAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].KHTAmount))
        self.dlg.txt_KHTCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].KHTCapacity))
        self.dlg.txt_KHTPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].KHTPrice))
        self.dlg.txt_ARAHCUAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].ARAHCUAmount))
        self.dlg.txt_ARAHCUCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].ARAHCUCapacity))
        self.dlg.txt_ARAHCUPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].ARAHCUPrice))
        self.dlg.txt_ARACHAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].ARACHAmount))
        self.dlg.txt_ARACHCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].ARACHCapacity))
        self.dlg.txt_ARACHPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].ARACHPrice))
        self.dlg.txt_alloyCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].AdditionalAlloyCapacity))
        self.dlg.txt_alloyPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].AdditionalAlloyPrice))
        self.dlg.txt_SMRAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].SMRAmount))
        self.dlg.txt_SMRCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].SMRCapacity))
        self.dlg.txt_SMRPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].SMRPrice))
        self.dlg.txt_ThiolexAmtCHEMEX.setText(to_str(self.Refinery_List[self.SelectedRefinery].ThiolexAmount))
        self.dlg.txt_ThiolexCapCHEMEX.setText(to_str(self.Refinery_List[self.SelectedRefinery].ThiolexCapacity))
        self.dlg.txt_ThiolexPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].ThiolexPrice))
        self.dlg.txt_LoCatAmtCHEMEX.setText(to_str(self.Refinery_List[self.SelectedRefinery].LoCatAmount))
        self.dlg.txt_LoCatCapCHEMEX.setText(to_str(self.Refinery_List[self.SelectedRefinery].LoCatCapacity))
        self.dlg.txt_LoCatPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].LoCatPrice))
        self.dlg.txt_SNRAmtCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].SNRAmount))
        self.dlg.txt_SNRCapCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].SNRCapacity))
        self.dlg.txt_SNRPriceCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].SNRPrice))
        self.dlg.txt_processUnitsCostCHEMEX.setText(special_format(self.Refinery_List[self.SelectedRefinery].TotalBPD))
        if self.IsDomestic:
            self.dlg.txt_processUnitsCostCHEMEX.setText("$ " + special_format(self.Refinery_List[self.SelectedRefinery].DomesticTotalPrice))
        elif self.IsInternational:
            self.dlg.txt_processUnitsCostCHEMEX.setText("$ " + special_format(self.Refinery_List[self.SelectedRefinery].OverseasTotalPrice))
        return True
    def ProductsBreakdownCreateUnit(self):
        #This is one of the function/buttons that is central to all of the calculations done by APE. It calls many other functions that are written in this class.
        #The inputs to this function are the products that are selected by the user (through the GUI) and the capacity that is desired
        #APE can handle multiple refineries, but this function only calculates for the currently selected refiner. ProductsBreakdownCreateAllUnits calculates every present refinery

        #sc = MLG.MyStaticMplCanvas(dlg, width=5, height=4, dpi=100)
        try:
            if len(self.Refinery_List) == 0: #Error checking
                return
            self.dlg.txt_SelectedRefinery.setText(self.SelectedRefinery)
        
            if self.ProductsCheckBoxCheck() == False: #Makes sure there are no illegal inputs
                return
            if self.ProductsChangesCheck() == True or self.RefineryChangesCheck() == True: #If you changed something, it will re initialize so the values will update on the screen after calculations
                self.Refinery_List[self.SelectedRefinery].ReInitialize()

            self.AssignDashboardVariables()
            AssignmentError = self.AssignDashboardVariables() #Only grabs the BPD, and checkboxes right now
            if AssignmentError == True:
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please Enter in a Capacity Value in BPD', QtWidgets.QMessageBox.Ok)
                return
            if self.Refinery_List[self.SelectedRefinery].IsCalculated == False:
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.ProductsBreakdown()
                self.Refinery_List[self.SelectedRefinery].CalculateRefinery()
                self.Refinery_List[self.SelectedRefinery].OptimizeHydrotreaters()
                self.Refinery_List[self.SelectedRefinery].AddSRU()
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.CalculateVolumeFlow() 
                self.Refinery_List[self.SelectedRefinery].UnitListBreakDown()
                self.Refinery_List[self.SelectedRefinery].CalculateTotalUtilities()
                self.FillOutDashboardVariables()
                self.CreateDictionary()
                self.UploadImage()
            #dlg.tabWidget.hide()
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Calculating Units. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
            raise
        return
    def ProductsBreakdownCreateAllUnits(self):
        #This is the function/button that is central to all of the calculations done by APE. It calls many other functions that are written in this class.
        #The inputs to this function are the products that are selected by the user (through the GUI) and the capacity that is desired
        #APE can handle multiple refineries, and this function calculates all of the refineries that were added prior to pushing this button

        try:
            if len(self.Refinery_List) == 0: #Error checking
                return
            #This section of code makes sure that all selected products information is correct and no errors occur
            if self.ProductsChangesCheck() == True or self.RefineryChangesCheck() == True: #If you changed something, it will re initialize so the values will update on the screen after calculations
                self.Refinery_List[self.SelectedRefinery].ReInitialize()
        
            self.AssignDashboardVariables()
            OriginalSelectedRefinery = self.SelectedRefinery
            for RefineryName in list(self.Refinery_List.keys()):
                self.SelectedRefinery = RefineryName
                self.dlg.txt_SelectedRefinery.setText(self.SelectedRefinery)
                self.FillOutDashboardVariables()
                if self.ProductsCheckBoxCheck() == False :
                    return

                self.AssignDashboardVariables()

            #This next section of code will go through and alculate each individual refinery

            for RefineryName in list(self.Refinery_List.keys()):
                print("THE CURRENT REFINERY NAME IS: " + self.SelectedRefinery)
                if self.Refinery_List[RefineryName].IsCalculated == False:
                    self.Refinery_List[RefineryName].ProductsInstance.ProductsBreakdown()
                    self.Refinery_List[RefineryName].CalculateRefinery()
                    self.Refinery_List[RefineryName].OptimizeHydrotreaters()
                    self.Refinery_List[RefineryName].AddSRU()  
                    self.Refinery_List[RefineryName].ProductsInstance.CalculateVolumeFlow() 
                    self.Refinery_List[RefineryName].UnitListBreakDown() 
                    self.Refinery_List[RefineryName].CalculateTotalUtilities()
            self.CreateDictionary()        
                #dlg.tabWidget.hide() 
            self.SelectedRefinery = OriginalSelectedRefinery
            self.FillOutDashboardVariables()
            self.UploadImage() 
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Creating Units. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def SaveProductInfo(self):
        #This function is connected to a button, and assigns the current info on the dashboard to a singular refinery (the selected one) if no errors are present
        try:
            if len(self.Refinery_List) == 0:
                return
            self.ProductsCheckBoxCheck()
            if self.ProductsChangesCheck() == True or self.RefineryChangesCheck() == True:
                self.Refinery_List[self.SelectedRefinery].ReInitialize()
                self.OverallRefinery.ReInitialize() #Method to Re initialize the refinery without re initializing the assay attached to that refinery
            self.AssignDashboardVariables() 
            self.FillOutDashboardVariables()
            self.UploadImage()

            QtWidgets.QMessageBox.information(self.dlg, 'Save', 'You have successfully saved the selected information', QtWidgets.QMessageBox.Ok)
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Saving Assay Data. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def SaveAllProductInfo(self):
        #This function is connected to a button and assigns the current info on the dashboard to all refineries that were added if no errors are present
        try:
            if len(self.Refinery_List) == 0:
                return
            OriginalSelectedRefinery = self.SelectedRefinery
            self.ProductsCheckBoxCheck()
            for RefineryName in list(self.Refinery_List.keys()):
                self.SelectedRefinery = RefineryName
                if self.ProductsChangesCheck() == True or self.RefineryChangesCheck() == True:
                    self.Refinery_List[RefineryName].ReInitialize()
                    self.OverallRefinery.ReInitialize() #Method to Re initialize the refinery without re initializing the assay attached to that refinery
                self.AssignDashboardVariables()
                self.FillOutDashboardVariables()
                self.UploadImage()
            QtWidgets.QMessageBox.information(self.dlg, 'Save', 'You have successfully saved all the selected information', QtWidgets.QMessageBox.Ok)
            self.SelectedRefinery = OriginalSelectedRefinery
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Saving Assay Info. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return

    # Assay Table on Assay Viewer Page
    def AddRow(self): #This function is connected to a button
        self.dlg.tbl_assay.insertRow(self.dlg.tbl_assay.rowCount())
        for i in range(0,self.dlg.tbl_assay.columnCount()):
            self.dlg.tbl_assay.setItem(self.dlg.tbl_assay.rowCount()-1,i,QTableWidgetItem(''))
            self.dlg.tbl_assay.setVerticalHeaderItem(self.dlg.tbl_assay.rowCount()-1,QTableWidgetItem(to_str(self.dlg.tbl_assay.rowCount()-2)))
        return
    def RemoveRow(self): #This function is connected to a button
        removedRow = False
        if self.dlg.tbl_assay.rowCount() > 2:
            try:
                print("The selected indeces of the row is:" + str(self.dlg.tbl_assay.selectedIndexes()[0].row())) 
                if self.dlg.tbl_assay.selectedIndexes()[0].row() > 2:
                    self.dlg.tbl_assay.removeRow(self.dlg.tbl_assay.selectedIndexes()[0].row())
                    removedRow = True
            except:
                pass
            if removedRow == False:
                self.dlg.tbl_assay.removeRow(self.dlg.tbl_assay.rowCount()-1)
        return
    def InputDataToDatabase(self):
        try:
            #This function inserts assay information into the database
            self.MyDatabase.ConnectToDatabase()
            RowCount = self.dlg.tbl_assay.rowCount()
            ColumnCount = self.dlg.tbl_assay.columnCount()

            try: #Ensures that all the data in the table are floating point values
                for i in range(1,RowCount):
                    for j in range(0,ColumnCount):
                        a = to_float(self.dlg.tbl_assay.item(i,j).text())
                        if a == 'None':
                            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Error in entered information. Please check to make sure all values are numbers', QtWidgets.QMessageBox.Ok)
                            return
            except:
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Error in entered information. Please check to make sure all values are numbers', QtWidgets.QMessageBox.Ok)
                return

            #The turn_null function only turns values null if a value does not exist (exists as '' in the interface)          
            Year = to_str(self.dlg.txt_year.text())
            AssayName = to_str(self.dlg.txt_assay.text())
            Source = to_str(self.dlg.txt_source.text())
            AssayRegion = to_str(self.dlg.txt_region.text())
            Country = to_str(self.dlg.txt_country.text())
            AssaySource = Source

            #Temporary PlaceHolders to compare later

            TempCountry = ''
            TempYear = ''

            #the following code appends year and country to the source name to create an overall assay source to be stored in the database (Will be revised later)

            if Country != '':
                TempCountry += " <" + Country + ">"
                AssaySource += TempCountry
            if Year != '':
                TempYear += " <" + Year + ">"
                AssaySource += TempYear

            #This next section checks the Source Name combo box to see if the data that is currently on the table will create a new entry in the database, or should prompt the user to select Replace/Add New
        
            #At this point we need to check if there are mutliple instances of the same assay
            #https://dba.stackexchange.com/questions/117609/looking-for-simple-contains-method-when-searching-text-in-postgresql

            self.MyDatabase.cursor.execute("SELECT assaysource FROM " + self.MyDatabase.tableName + " WHERE assaysource LIKE " + "'%" + TempCountry + "%' AND assaysource LIKE " + "'%" + TempYear + "%'" + 
                                           " AND assaysource LIKE " + "'%" + Source + "%'" + " AND assayname LIKE " + "'%" + AssayName + "%'" + " AND assayregion LIKE " + "'%" + AssayRegion + "%'" + ";")
            self.List_Of_Assay_Sources = self.MyDatabase.cursor.fetchall()
            #print(self.List_Of_Assay_Sources)

            if len(self.List_Of_Assay_Sources) > 0 and self.PopupCheck == False:
                self.MyDatabase.DisconnectFromDatabase()
                self.BuildPopup()
                return
            
            #Array Variable Initialization
            Initial_Temperature_Data = [None] * (RowCount - 1)
            Final_Temperature_Data = [None] * (RowCount - 1)
            Weight_Yield_Data = [None] * (RowCount - 1)
            Volume_Yield_Data = [None] * (RowCount - 1)
            Density_Data = [None] * (RowCount - 1)
            Sulfur_Data = [None] * (RowCount - 1)
            Mercaptan_Sulfur_Data = [None] * (RowCount - 1)
            Smoke_Point_Data = [None] * (RowCount - 1)
            Viscosity_T1_Data = [None] * (RowCount - 1)
            Viscosity_T2_Data = [None] * (RowCount - 1)
            TAN_Data = [None] * (RowCount - 1)
            Molecular_Weight_Data = [None] * (RowCount - 1)
            Salt_Content_Data = [None] * (RowCount - 1)
            Flash_Point_Data = [None] * (RowCount - 1)
            Freeze_Point_Data = [None] * (RowCount - 1)
            Cloud_Point_Data = [None] * (RowCount - 1)
            Pour_Point_Data = [None] * (RowCount - 1)
            RON_Data = [None] * (RowCount - 1)
            MON_Data = [None] * (RowCount - 1)
            Cetane_Data = [None] * (RowCount - 1)
            RVP_Data = [None] * (RowCount - 1)
            TVP_Data = [None] * (RowCount - 1)
            Aromatics_Data = [None] * (RowCount - 1)
            Olefins_Data = [None] * (RowCount - 1)
            Paraffins_Data = [None] * (RowCount - 1)
            Naphthenes_Data = [None] * (RowCount - 1)
            Vanadium_Data = [None] * (RowCount - 1)
            Nickel_Data = [None] * (RowCount - 1)
            Nitrogen_Data = [None] * (RowCount - 1)

            #Assign singular values
            AssaySource += " <" + str(self.SourceNumber) + ">"
            CumulativeCheck = self.dlg.chk_CumulativeCheck.isChecked()
            Methane = turn_null(self.dlg.txt_Methane.text())
            Ethane = turn_null(self.dlg.txt_Ethane.text())
            Propane = turn_null(self.dlg.txt_Propane.text())
            iButane = turn_null(self.dlg.txt_iButane.text())
            nButane = turn_null(self.dlg.txt_nButane.text())
            nPentane = turn_null(self.dlg.txt_nPentane.text())
            iPentane = turn_null(self.dlg.txt_iPentane.text())
            Cyclopentane = turn_null(self.dlg.txt_Cyclopentane.text())
            Cyclohexane = turn_null(self.dlg.txt_Cyclohexane.text())
            Benzene = turn_null(self.dlg.txt_Benzene.text())
            T1 = turn_null(self.dlg.txt_T1.text())
            T2 = turn_null(self.dlg.txt_T2.text())
            T3 = turn_null(self.dlg.txt_T3.text())
            T4 = turn_null(self.dlg.txt_T4.text())
            T5 = turn_null(self.dlg.txt_T5.text())
            IBP = turn_null(self.dlg.txt_IBP.text())
            FBP = turn_null(self.dlg.txt_FBP.text())

            #Assign Units from Singular values
            MethaneUnit = self.dlg.cmb_MethaneUnit.currentText()
            EthaneUnit = self.dlg.cmb_EthaneUnit.currentText()
            PropaneUnit = self.dlg.cmb_PropaneUnit.currentText()
            iButaneUnit = self.dlg.cmb_iButaneUnit.currentText()
            nButaneUnit = self.dlg.cmb_nButaneUnit.currentText()
            nPentaneUnit = self.dlg.cmb_nPentaneUnit.currentText()
            iPentaneUnit = self.dlg.cmb_iPentaneUnit.currentText()
            CyclopentaneUnit = self.dlg.cmb_CyclopentaneUnit.currentText()
            CyclohexaneUnit = self.dlg.cmb_CyclohexaneUnit.currentText()
            BenzeneUnit = self.dlg.cmb_BenzeneUnit.currentText()
            T1Unit = self.dlg.cmb_T1Unit.currentText()
            T2Unit = self.dlg.cmb_T2Unit.currentText()
            IBPUnit = self.dlg.cmb_IBPUnit.currentText()
            FBPUnit = self.dlg.cmb_FBPUnit.currentText()

#********************************************************************************************************************************************************************************************************      
            #Note that the first element will  go to Whole Crude Properties
            #NOTE: THESE MUST BE IN ORDER WITH THE UI FILE OR DATA WILL GO IN INCORRECTLY. PLEASE ENSURE THESE ARE IN ORDER BEFORE MAKING ANY CHANGES TO DATA
            #The following is a section of copied code. In the future, this should reference by column name, not index, so that if we were to make a data change, we will not have to worry about the order
            filelists = [Initial_Temperature_Data, Final_Temperature_Data, Weight_Yield_Data, Volume_Yield_Data, Density_Data, Sulfur_Data, Mercaptan_Sulfur_Data, Smoke_Point_Data, Viscosity_T1_Data, Viscosity_T2_Data, TAN_Data, Molecular_Weight_Data, Salt_Content_Data,
                         Flash_Point_Data, Freeze_Point_Data, Cloud_Point_Data, Pour_Point_Data, RON_Data, MON_Data, Cetane_Data, RVP_Data, TVP_Data, Aromatics_Data, Olefins_Data, Paraffins_Data, Naphthenes_Data, Vanadium_Data, Nickel_Data, Nitrogen_Data] 
            for x in range(1, RowCount):
                for i, filelist in enumerate(filelists):
                    print("Before if statement, item at " + str(x) + ', ' + str(i) + " = " + to_str(self.dlg.tbl_assay.item(x,i)))
                    if  self.dlg.tbl_assay.item(x,i) == None or to_str(self.dlg.tbl_assay.item(x,i).text()) == '' or to_str(self.dlg.tbl_assay.item(x,i).text()) == 'IBP' or to_str(self.dlg.tbl_assay.item(x,i).text()) == 'FBP':
                        filelist[x-1] = 'NULL'
                        #print("After if statement, item at x,i = " + to_str(self.dlg.tbl_assay.item(x,i)))
                        continue
                    filelist[x-1] = to_float(self.dlg.tbl_assay.item(x, i).text())

#********************************************************************************************************************************************************************************************************  

            Initial_Temperature_Data = str(Initial_Temperature_Data).replace("'","")
            Final_Temperature_Data = str(Final_Temperature_Data).replace("'","")
            Weight_Yield_Data = str(Weight_Yield_Data).replace("'","")
            Volume_Yield_Data = str(Volume_Yield_Data).replace("'","")
            Density_Data = str(Density_Data).replace("'","")
            Sulfur_Data = str(Sulfur_Data).replace("'","")
            Mercaptan_Sulfur_Data = str(Mercaptan_Sulfur_Data).replace("'","")
            Smoke_Point_Data = str(Smoke_Point_Data).replace("'","")
            Viscosity_T1_Data = str(Viscosity_T1_Data).replace("'","")
            Viscosity_T2_Data = str(Viscosity_T2_Data).replace("'","")
            TAN_Data = str(TAN_Data).replace("'","")
            Molecular_Weight_Data = str(Molecular_Weight_Data).replace("'","")
            Salt_Content_Data = str(Salt_Content_Data).replace("'","")
            Flash_Point_Data = str(Flash_Point_Data).replace("'","")
            Freeze_Point_Data = str(Freeze_Point_Data).replace("'","")
            Cloud_Point_Data = str(Cloud_Point_Data).replace("'","")
            Pour_Point_Data = str(Pour_Point_Data).replace("'","")
            RON_Data = str(RON_Data).replace("'","")
            MON_Data = str(RON_Data).replace("'","")
            Cetane_Data = str(Cetane_Data).replace("'","")
            RVP_Data = str(RVP_Data).replace("'","")
            TVP_Data = str(TVP_Data).replace("'","")
            Aromatics_Data = str(Aromatics_Data).replace("'","")
            Olefins_Data = str(Olefins_Data).replace("'","")
            Paraffins_Data = str(Paraffins_Data).replace("'","")
            Naphthenes_Data = str(Naphthenes_Data).replace("'","")
            Vanadium_Data = str(Vanadium_Data).replace("'","")
            Nickel_Data = str(Nickel_Data).replace("'","")
            Nitrogen_Data = str(Nitrogen_Data).replace("'","")

            self.PopupCheck = False
            self.SourceNumber = 1
            print(AssaySource)
       
            #NOTE: Because this is entering in to the database using postgres, the corresponding name must match with its string. No exceptions can ever be made here
            InsertQuery = ("INSERT INTO " + to_str(self.MyDatabase.tableName) + "(assayregion, assayname, assaysource, cumulativecheck, initial_temperature, initialtemperatureunit, "
                "final_temperature, finaltemperatureunit, weight_yield, weightyieldunit, volume_yield, volumeyieldunit, density_data, densitydataunit, sulfur_data, sulfurdataunit, mercaptan_data, mercaptandataunit, "
                "smoke_point, smokepointunit, viscosity1_data, viscosity1dataunit, t1_viscosity1, t1viscosity1unit, viscosity2_data, viscosity2dataunit,  t2_viscosity2, t2viscosity2unit, tan_data, tandataunit,  "
                "molecular_weight, molecularweightunit, salt_data, saltdataunit, flash_point, flashpointunit, freeze_point, freezepointunit, cloud_point, cloudpointunit, pour_point, pourpointunit, ron_data, mon_data, cetane_number, rvp_data, rvpdataunit, "
                "tvp_data, tvpdataunit, aromatics_data, aromaticsdataunit, olefins_data, olefinsdataunit, paraffins_data, paraffinsdataunit, naphthenes_data, naphthenesdataunit, vanadium_data, vanadiumdataunit, nickel_data, nickeldataunit, nitrogen_data, nitrogendataunit,  "
                "methane_data, methanedataunit, ethane_data, ethanedataunit, propane_data, propanedataunit, isobutane_data, isobutanedataunit, n_butane_data, nbutanedataunit) " )

            #In the value query, we do ARRAY _____ ::FLOAT[] so that we can store an array of NULL values just in case we are given no data for that specific input
            #EXAMPLE:   ARRAY [13,26,NULL,NULL,NULL]::FLOAT[]   would be the string corresponding to a float array postgres input
            ValueQuery = ("VALUES (" + add_quotes(AssayRegion) + ', ' + add_quotes(AssayName) + ', ' + add_quotes(AssaySource) + ', ' + to_str(CumulativeCheck) + ', ' 
                          + 'ARRAY ' + Initial_Temperature_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Initial Temp\n')).currentText()) + ', '
                          + 'ARRAY ' + Final_Temperature_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Final Temp\n')).currentText()) + ', '
                          + 'ARRAY ' + Weight_Yield_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Weight Yield\n')).currentText()) + ', '
                          + 'ARRAY ' + Volume_Yield_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Volume Yield\n')).currentText()) + ', '  
                          + 'ARRAY ' + Density_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Density\n')).currentText()) + ', '
                          + 'ARRAY ' + Sulfur_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Sulfur\n')).currentText()) + ', '
                          + 'ARRAY ' + Mercaptan_Sulfur_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Mercaptan Sulfur\n')).currentText()) + ', '
                          + 'ARRAY ' + Smoke_Point_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Smoke Point\n')).currentText()) + ', '
                          + 'ARRAY ' + Viscosity_T1_Data +  '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Viscosity @ T1\n')).currentText()) + ', ' + T1 + ', ' + add_quotes(T1Unit) + ', '
                          + 'ARRAY ' + Viscosity_T2_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Viscosity @ T2\n')).currentText()) + ', ' + T2 + ', ' + add_quotes(T2Unit) + ', '
                          + 'ARRAY ' + TAN_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('TAN\n')).currentText()) + ', '
                          + 'ARRAY ' + Molecular_Weight_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Molecular Weight\n')).currentText()) + ', '
                          + 'ARRAY ' + Salt_Content_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Salt Content\n')).currentText()) + ', '
                          + 'ARRAY ' + Flash_Point_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Flash Point\n')).currentText()) + ', '
                          + 'ARRAY ' + Freeze_Point_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Freeze Point\n')).currentText()) + ', '
                          + 'ARRAY ' + Cloud_Point_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Cloud Point\n')).currentText()) + ', '
                          + 'ARRAY ' + Pour_Point_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Pour Point\n')).currentText()) + ', '
                          + 'ARRAY ' + RON_Data + '::FLOAT[]' + ', '
                          + 'ARRAY ' + MON_Data + '::FLOAT[]' + ', '
                          + 'ARRAY ' + Cetane_Data + '::FLOAT[]' + ', '
                          + 'ARRAY ' + RVP_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('RVP\n')).currentText()) + ', ' 
                          + 'ARRAY ' + TVP_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('TVP\n')).currentText()) + ', '
                          + 'ARRAY ' + Aromatics_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Aromatics\n')).currentText()) + ', '
                          + 'ARRAY ' + Olefins_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Olefins\n')).currentText()) + ', '
                          + 'ARRAY ' + Paraffins_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Paraffins\n')).currentText()) + ', '
                          + 'ARRAY ' + Naphthenes_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Naphthenes\n')).currentText()) + ', '
                          + 'ARRAY ' + Vanadium_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Vanadium\n')).currentText()) + ', '
                          + 'ARRAY ' + Nickel_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Nickel\n')).currentText()) + ', '
                          + 'ARRAY ' + Nitrogen_Data + '::FLOAT[]' + ', ' + add_quotes(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Nitrogen\n')).currentText()) + ', '
                          + Methane + ', ' + add_quotes(MethaneUnit) + ', '
                          + Ethane + ', ' + add_quotes(EthaneUnit) + ', '
                          + Propane + ', ' + add_quotes(PropaneUnit) + ', '
                          + iButane + ', ' + add_quotes(iButaneUnit) + ', '
                          + nButane + ', ' + add_quotes(nButaneUnit) + ');')
            print(filelists)
            try:
                print(InsertQuery + ValueQuery)
                self.MyDatabase.cursor.execute(InsertQuery + ValueQuery)
                self.MyDatabase.connection.commit()
                

            except(Exception, psycopg2.DatabaseError) as error:
                print("Error uploading the data to the database", error)
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Error in entry process. Please make sure everything is correct', QtWidgets.QMessageBox.Ok)
                self.MyDatabase.DisconnectFromDatabase()
                return
            QtWidgets.QMessageBox.information(self.dlg, 'Success', 'Assay Successfully entered in to database', QtWidgets.QMessageBox.Ok)
            self.MyDatabase.DisconnectFromDatabase()
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Inputting Data to Database. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def DeleteDataFromDatabase(self):
        try:
            #This function inserts assay information into the database
            self.MyDatabase.ConnectToDatabase()

            #This next section checks the Source Name combo box to see if the data that is currently on the table will create a new entry in the database, or should prompt the user to select Replace/Add New
        
            #At this point we need to check if there are mutliple instances of the same assay
            #https://dba.stackexchange.com/questions/117609/looking-for-simple-contains-method-when-searching-text-in-postgresql

            SourceName = ''
            FullID = self.dlg.cmb_SourceName.currentText()
            for i in range(len(FullID)-1,-1,-1):
                if FullID[i] == ')':
                    EndOfSourceIndex = i
                    continue
                if FullID[i] == '(':
                    StartOfSourceIndex = i
                    break
            for i in range(StartOfSourceIndex+1,EndOfSourceIndex):
                SourceName += FullID[i]
        
        
            try:
                self.MyDatabase.cursor.execute("DELETE FROM " + self.MyDatabase.tableName + " WHERE assaysource = " + add_quotes(SourceName))
                self.MyDatabase.connection.commit()
                QtWidgets.QMessageBox.information(self.dlg, 'Success!', 'Assay Successfully Deleted from Database', QtWidgets.QMessageBox.Ok)
            except:
                QtWidgets.QMessageBox.information(self.dlg, 'Error!', 'Error in Assay Deletion', QtWidgets.QMessageBox.Ok)
            #self.List_Of_Assay_Sources = self.MyDatabase.cursor.fetchall()
            self.MyDatabase.DisconnectFromDatabase()
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Deleting Data From Database. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def GraphAssayData(self):
            
        if self.dlg.cmb_SourceName.currentText() == "Select Source Name":
            self.dlg.wdg_AssayDataPieChart.canvas.axes.clear()
            self.dlg.wdg_AssayDataGraph.canvas.axes.clear()
            self.dlg.wdg_AssayDataPieChart.canvas.draw()
            self.dlg.wdg_AssayDataGraph.canvas.draw()
            return
        if self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.IsCalculated != True:
            return
        if len(self.Refinery_List.values()) == 0:
            return
        #This next section of code graphs the TBP curve
        self.dlg.wdg_AssayDataGraph.canvas.axes.clear()

        colorArray = ['goldenrod','turquoise','springgreen','green','red','maroon','sienna','dimgray']
        #plt.subplot(1,2,1) 
        self.dlg.wdg_AssayDataGraph.canvas.axes.set_xlim(0,100) #set the chart x limit
        self.dlg.wdg_AssayDataGraph.canvas.axes.set_ylim(0,self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.AssayInstance.FBP) #set the chart y limit
        self.dlg.wdg_AssayDataGraph.canvas.axes.set_yticks(np.arange(0,800,step=50.0))
        labels = self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.Product_Names
        cuts = self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.Yield_List
        self.dlg.wdg_AssayDataGraph.canvas.axes.set_title("TBP Curve")
        self.dlg.wdg_AssayDataGraph.canvas.axes.set_ylabel("Temperature (C)")
        self.dlg.wdg_AssayDataGraph.canvas.axes.set_xlabel("Volume % Distilled")
        self.dlg.wdg_AssayDataGraph.canvas.axes.plot(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.AssayInstance.Volume_Plotting_Data, self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.AssayInstance.Temperature_Data, color = 'black') #This will eventually be replaced with self.MyAssay.Volume_Plotting_Data when we get the ability for options

        for i in range(0,len(remove_zeros(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents))):
            print(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.Original_Yield_List)
            print(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.New_Specified_Cut_Temps)
            self.dlg.wdg_AssayDataGraph.canvas.axes.plot([0,remove_zeros(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[i])],[self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.New_Specified_Cut_Temps[i][1],self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.New_Specified_Cut_Temps[i][1]], alpha = 0.3, c = colorArray[i])
            
        for i in range(0,len(remove_zeros(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.Original_Yield_List))): 
            self.dlg.wdg_AssayDataGraph.canvas.axes.plot([self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[i],self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[i]],[0,self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.New_Specified_Cut_Temps[i][1]], c = colorArray[i])



        # Fill in the plot to correspond with the pie chart
        tempVolumeData = []
        tempTemperatureData = []
        self.dlg.wdg_AssayDataGraph.canvas.axes.fill_between([0,0,self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[0],self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[0]],[0,self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.AssayInstance.IBP,self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.New_Specified_Cut_Temps[0][1],0], alpha = 0.65, color = colorArray[0])
        for i in range(0,len(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.Original_Yield_List)-1):
            for j in range(0,len(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.AssayInstance.Volume_Plotting_Data)):
                if self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.AssayInstance.Volume_Plotting_Data[j] > self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[i] and self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.AssayInstance.Volume_Plotting_Data[j] < self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[i+1]:
                    tempVolumeData.append(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.AssayInstance.Volume_Plotting_Data[j])
                    tempTemperatureData.append(self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.AssayInstance.Temperature_Data[j])
            tempx = [self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[i],self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[i], self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[i+1],self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.CumulativeVolumePercents[i+1]]
            tempy = [0,self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.New_Specified_Cut_Temps[i][1], self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.New_Specified_Cut_Temps[i+1][1],0]
            index = 2
            for k in range(0,len(tempVolumeData)):
                tempx.insert(index,tempVolumeData[k])
                tempy.insert(index,tempTemperatureData[k])
                index += 1
            self.dlg.wdg_AssayDataGraph.canvas.axes.fill_between(tempx,tempy, alpha = 0.5, color = colorArray[i+1])
            tempVolumeData = []
            tempTemperatureData = []

        self.dlg.wdg_AssayDataGraph.canvas.draw()
        

        #This next section of code plots the pie chart
        self.dlg.wdg_AssayDataPieChart.canvas.axes.clear()

        labels = self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.Product_Names
        cuts = self.Refinery_List[self.dlg.cmb_SourceName.currentText()].ProductsInstance.Yield_List
        self.dlg.wdg_AssayDataPieChart.canvas.axes.pie(cuts, labels=labels, shadow = True, autopct ='%1.1f%%', colors = colorArray,)
        self.dlg.wdg_AssayDataPieChart.canvas.axes.axis('equal')
        #plt.show()

        self.dlg.wdg_AssayDataPieChart.canvas.draw()
        return

    #Upload Method on Assay Viewer Page
    def AssignAssayVariables(self): #This function is not connected to a button.

        #FUNCTION: This function adds assay variables from the database to the selected Refinery in the refinery lsit

        #This is temporary for now
        if len(self.SelectedRefinery) == 0: #Error Checking
            return


        self.MyDatabase.ConnectToDatabase()
        self.MyDatabase.cursor.execute("SELECT * FROM " + self.MyDatabase.tableName + " WHERE assayregion = '" + self.dlg.cmb_AssayRegion.currentText() + "' AND assayname = '" + self.dlg.cmb_AssayName.currentText() + "' AND assaysource = '" + self.dlg.cmb_AssaySource.currentText() + "'")
        AssayDatabaseDictionary = self.MyDatabase.cursor.fetchall()
        print("The assay name is: " + AssayDatabaseDictionary[0]["assayname"])
        
        # THIS IS ASSIGNMENT OF VARIABLES
        
        #Need to add some sort of conversion
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CumulativeVolCheck = AssayDatabaseDictionary[0]["cumulativecheck"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CumulativeWeightCheck = AssayDatabaseDictionary[0]["cumulativecheck"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_Data = AssayDatabaseDictionary[0]["weight_yield"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data = AssayDatabaseDictionary[0]["volume_yield"]
        TemporaryTemperatureData = AssayDatabaseDictionary[0]["initial_temperature"]
        #We now assign the assay variables to the database values
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data = AssayDatabaseDictionary[0]["sulfur_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Mercaptan_Data = AssayDatabaseDictionary[0]["mercaptan_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Smoke_Point_Data = AssayDatabaseDictionary[0]["smoke_point"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.TAN_Data = AssayDatabaseDictionary[0]["tan_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Salt_Content_Data = AssayDatabaseDictionary[0]["salt_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Flash_Point_Data = AssayDatabaseDictionary[0]["flash_point"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Freeze_Point_Data = AssayDatabaseDictionary[0]["freeze_point"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Cloud_Point_Data = AssayDatabaseDictionary[0]["cloud_point"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Pour_Point_Data = AssayDatabaseDictionary[0]["pour_point"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.RON_Data = AssayDatabaseDictionary[0]["ron_data"] #RON does not have a unit
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Cetane_Data = AssayDatabaseDictionary[0]["cetane_number"] #Cetane number has no units
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.RVP_Data = AssayDatabaseDictionary[0]["rvp_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Aromatics_Data = AssayDatabaseDictionary[0]["aromatics_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Olefins_Data = AssayDatabaseDictionary[0]["olefins_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Paraffins_Data = AssayDatabaseDictionary[0]["paraffins_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Vanadium_Data = AssayDatabaseDictionary[0]["vanadium_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Nickel_Data = AssayDatabaseDictionary[0]["nickel_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Nitrogen_Data = AssayDatabaseDictionary[0]["nitrogen_data"]
        #Single Value items 
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.T1ForVisc1 = AssayDatabaseDictionary[0]["t1_viscosity1"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.T2ForVisc2 = AssayDatabaseDictionary[0]["t2_viscosity2"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Methane = AssayDatabaseDictionary[0]["methane_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Ethane = AssayDatabaseDictionary[0]["ethane_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Propane = AssayDatabaseDictionary[0]["propane_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NButane = AssayDatabaseDictionary[0]["n_butane_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IsoButane = AssayDatabaseDictionary[0]["isobutane_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IsoPentane = AssayDatabaseDictionary[0]["i_pentane_data"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NPentane = AssayDatabaseDictionary[0]['n_pentane_data']
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CycloPentane = AssayDatabaseDictionary[0]['cyclopentane_data']
        #We must take units into account
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SulfurDataUnit = AssayDatabaseDictionary[0]["sulfurdataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.MercaptanDataUnit = AssayDatabaseDictionary[0]["mercaptandataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SmokePointDataUnit = AssayDatabaseDictionary[0]["smokepointunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.TANDataUnit = AssayDatabaseDictionary[0]["tandataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SaltContentDataUnit = AssayDatabaseDictionary[0]["saltdataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FlashPointDataUnit = AssayDatabaseDictionary[0]["flashpointunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FreezePointDataUnit = AssayDatabaseDictionary[0]["freezepointunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CloudPointDataUnit = AssayDatabaseDictionary[0]["cloudpointunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.PourPointDataUnit = AssayDatabaseDictionary[0]["pourpointunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.RVPDataUnit = AssayDatabaseDictionary[0]["rvpdataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.AromaticsDataUnit = AssayDatabaseDictionary[0]["aromaticsdataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.OlefinsDataUnit = AssayDatabaseDictionary[0]["olefinsdataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.ParaffinsDataUnit = AssayDatabaseDictionary[0]["paraffinsdataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.VanadiumDataUnit = AssayDatabaseDictionary[0]["vanadiumdataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NickelDataUnit = AssayDatabaseDictionary[0]["nickeldataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NitrogenDataUnit = AssayDatabaseDictionary[0]["nitrogendataunit"]
        #Single Value Item Units
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.T1ForVisc1Unit = AssayDatabaseDictionary[0]["t1viscosity1unit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.T2ForVisc2Unit = AssayDatabaseDictionary[0]["t2viscosity2unit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.MethaneUnit = AssayDatabaseDictionary[0]["methanedataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.EthaneUnit = AssayDatabaseDictionary[0]["ethanedataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.PropaneUnit = AssayDatabaseDictionary[0]["propanedataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NButaneUnit = AssayDatabaseDictionary[0]["nbutanedataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IsoButaneUnit = AssayDatabaseDictionary[0]["isobutanedataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IsoPentaneUnit = AssayDatabaseDictionary[0]["ipentanedataunit"]
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NPentaneUnit = AssayDatabaseDictionary[0]['npentanedataunit']
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CycloPentaneUnit = AssayDatabaseDictionary[0]['cyclopentanedataunit']

        #The following code goes through the database and assigns all the important values to the Assay Instance for the selected refinery
        #For this first block of code we focus on density. Density is a weird case because it can be given in multiple forms including API, SG, or a density unit
        if AssayDatabaseDictionary[0]["densitydataunit"] == "API":
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API = AssayDatabaseDictionary[0]['density_data'].pop(0) #we pop SG/API data earlier because there is high probability we have overall SG, but not data for it
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API_Data = AssayDatabaseDictionary[0]["density_data"]
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SGDataCheck = False
            print("API IS: " + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API))
            #print("API IS: " + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API))
        elif AssayDatabaseDictionary[0]["densitydataunit"] == "SG":
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG = AssayDatabaseDictionary[0]['density_data'].pop(0) #we pop SG/API data earlier because there is high probability we have overall SG, but not data for it
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data = AssayDatabaseDictionary[0]["density_data"]
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.APIDataCheck = False
            print("SG IS: " + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG))
        elif AssayDatabaseDictionary[0]["densitydataunit"] == "lb/ft3":
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG = AssayDatabaseDictionary[0]['density_data'].pop(0)
            if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG != None:
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG = self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG/PoundsPerCubicFootWater
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data = AssayDatabaseDictionary[0]["density_data"] #We will convert the array of data correctly later after doing proper checks
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.APIDataCheck = False
        elif AssayDatabaseDictionary[0]["densitydataunit"] == "kg/m^3":
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG = AssayDatabaseDictionary[0]['density_data'].pop(0)
            if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG != None:
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG = self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG/KilogramsPerCubicMeterWater
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data = AssayDatabaseDictionary[0]["density_data"] #We will convert the array of data correctly later after doing proper checks
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.APIDataCheck = False
        elif AssayDatabaseDictionary[0]["densitydataunit"] == "kg/L":
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG = AssayDatabaseDictionary[0]['density_data'].pop(0)
            if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG != None:
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG = self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data = AssayDatabaseDictionary[0]["density_data"] #We will convert the array of data correctly later after doing proper checks
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.APIDataCheck = False
        #Temperatures are considered to be coupled, so if we are in fahrenheit, celsiuscheck must be set to false and vice versa
        if AssayDatabaseDictionary[0]["initialtemperatureunit"] == "Celsius" or  AssayDatabaseDictionary[0]["finaltemperatureunit"] == "Celsius":
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CelsiusCheck = True
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FahrenheitCheck = False
        elif AssayDatabaseDictionary[0]["initialtemperatureunit"] == "Fahrenheit" or  AssayDatabaseDictionary[0]["finaltemperatureunit"] == "Fahrenheit":
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FahrenheitCheck = True
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CelsiusCheck = False

        if AssayDatabaseDictionary[0]["sulfurdataunit"] == "wt%":
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur = self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data.pop(0)
        elif AssayDatabaseDictionary[0]["sulfurdataunit"] == "ppm":
            if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data[0] != None:
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur = self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data.pop(0) / PPMPerWeightPercent
            else: #This is in case of a dumb storage error ex: 'Wt%'
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur = self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data.pop(0)

        if AssayDatabaseDictionary[0]["saltdataunit"] == "ptb": #We do not have any additional if statements, because salt data is always given in pounds per thousand barrels
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SaltContent = self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Salt_Content_Data.pop(0)
        

        #End of Assignment of Variables

        #This is the first check we do to see if the data is present. If there is an error with input data (aka the lengths are not the same) we will consider the data false and will not calculate units based on that data
        VolumeDataCheck = True
        WeightDataCheck = True
        SGDataCheck = True
        APIDataCheck = True

        #There is clear separation between the above variable Checks and these. For the following variables, we will assume there is no data (Hence the variables are initialized to False) until proven otherwise. 
        #For the above pieces of data, we assume it is fully correct, unless we observe faults or data is missing.
        SulfurDataCheck = False
        SmokePointDataCheck = False
        MercaptanDataCheck = False
        TANDataCheck = False

        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data.pop(0) #This should equal 100 but there is no guarantee. 
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_Data.pop(0) #This should equal 100 but there is no guarantee.
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Smoke_Point_Data.pop(0) #The smoke point of the whole crude is not valuable
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Mercaptan_Data.pop(0) #The Mercaptans of the whole crude is not valuable
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.TAN_Data.pop(0) #The TAN of the whole crude is not valuable

        #The following are print statements to observe our untouched assay data. Feel free to comment the next part out if you want to increase speed and decrease command window clutter
        print("WeightDataCheck before:" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.WeightDataCheck))
        print("The following is weight data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_Data))
        print("The following is volume data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data))
        print("The following is SG data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data))
        print("The following is API data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API_Data))
        print("The following is sulfur data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data))
        print("The following is smoke point data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Smoke_Point_Data))
        print("The following is TAN data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.TAN_Data))
        print("The following is mercaptan data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Mercaptan_Data))
        
        #These for loops check to make sure everything is true
        for i in range(0, len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data)): #Checks to make sure all values in volume yield array are floats
            if not (isinstance(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data[i],float)):
                VolumeDataCheck = False
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.VolumeDataCheck = VolumeDataCheck

        if len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_Data) == len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data): #Makes sure that data entered is consistent
            for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_Data)): #Checks to make sure all values in the weight yield array are floats
                if not(isinstance(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_Data[i],float)):
                    WeightDataCheck = False
                    break
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.WeightDataCheck  = WeightDataCheck        

        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data != None and len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data) == len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data): # Length - 1 because we popped the SG earlier in the code
            for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data)): #Checks to make sure all values in the SG/density array are floats
                if not(isinstance(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data[i],float)):
                    SGDataCheck = False
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API_Data = [None] * len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data)
                    break
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SGDataCheck = SGDataCheck

        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API_Data != None and len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API_Data) == len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data): #Length - 1 because we popped the SG earlier in the code
            for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API_Data)): #Checks to make sure all values in the API/density array are floats
                if not(isinstance(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API_Data[i],float)):
                    APIDataCheck = False
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data = [None] * len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API_Data)
                    break
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.APIDataCheck = APIDataCheck
        
        #For theses next three checks, it is ok if some data is missing, because we have a way of interpolating/estimating later. However, We must find out if data exists to interpolate, and handle the case where there is no data for some of these categories
        if len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data) == len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data):
            for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data)): #Checks to make sure all values in the sulfur yield array are floats
                if (isinstance(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data[i],float)):
                    SulfurDataCheck = True 
                    break
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SulfurDataCheck = SulfurDataCheck

        if len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Smoke_Point_Data) == len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data):
            for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Smoke_Point_Data)): #Checks to make sure all values in the smoke point array are floats
                if (isinstance(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Smoke_Point_Data[i],float)):
                    SmokePointDataCheck = True
                    break
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SmokePointDataCheck = SmokePointDataCheck
            

        if len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Mercaptan_Data) == len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data):
            for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Mercaptan_Data)): #Checks to make sure all values in the weight yield array are floats
                if (isinstance(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Mercaptan_Data[i],float)):
                    MercaptanDataCheck = True
                    break
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.MercaptanDataCheck = MercaptanDataCheck
            

        if len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.TAN_Data) == len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data):
            for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.TAN_Data)): #Checks to make sure all values in the weight yield array are floats
                if (isinstance(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.TAN_Data[i],float)):
                    TANDataCheck = True
                    break
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.TANDataCheck = TANDataCheck
            

        # The following checks if we were given fuel gas data in the form of methane/ethane/propane/butane
        if (self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Methane != None and self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Ethane != None
        and self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Propane != None
        and (self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IsoButane != None or self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NButane != None)):
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.LightEndsDataCheck = True


        # The following are unit conversions
        if SGDataCheck == True:
            if AssayDatabaseDictionary[0]["densitydataunit"] == "lb/ft3":
                for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data)):
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data[i] = self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data[i]/PoundsPerCubicFootWater
            elif AssayDatabaseDictionary[0]["densitydataunit"] == "kg/m^3":
                for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data)):
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data[i] = self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data[i]/KilogramsPerCubicMeterWater


        if SulfurDataCheck == True:
            if AssayDatabaseDictionary[0]["sulfurdataunit"] == "ppm":
                for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data)):
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data[i] = special_multiply(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data[i], 1/PPMPerWeightPercent) #See Useful_Functions.py to look at special_multiply


        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Methane != None:
            if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.MethaneUnit == "wt%":
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Methane /= SGLiquidMethane
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.MethaneUnit = "vol%"
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Ethane != None:
            if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.EthaneUnit == "wt%":
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Ethane /= SGLiquidEthane
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.EthaneUnit = "vol%"
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Propane != None:
            if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.PropaneUnit == "wt%":
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Propane /= SGLiquidPropane
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.PropaneUnit = "vol%"
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IsoButane != None:
            if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IsoButaneUnit == "wt%":
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IsoButane /= SGLiquidButane
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IsoButaneUnit = "vol%"
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NButane != None:
            if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NButaneUnit == "wt%":
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NButane /= SGLiquidButane
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.NButaneUnit = "vol%"

        
        #We convert from fahrenheit to celsius because that is what our calculations are done in

        print("WeightDataCheck after:" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.WeightDataCheck))
        print("The following is weight data after: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_Data))
        print("The following is volume data after: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data))
        print("The following is SG data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data))
        print("The following is API data before: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.API_Data))
        print("The following is sulfur data after: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Sulfur_Data))
        print("The following is smoke point data after: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Smoke_Point_Data))
        print("The following is TAN data after: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.TAN_Data))
        print("The following is mercaptan data after: " + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Mercaptan_Data))
        #The following code turns Initial Temperature and final temperature data into a usable list of temperature values
        #Check to see if we have an IBP
        if TemporaryTemperatureData[1] == None:
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.InitialDataPointCheck = False
        else:
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.InitialDataPointCheck = True
            #if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CelsiusCheck == True:
            #Check this to make sure this makes sense
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBP = TemporaryTemperatureData[1]
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBPCheck = True
            #self.self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CalculateIBP()

            #elif self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FahrenheitCheck == True:
                #self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBP = (TemporaryTemperatureData[1]-32)* (5/9) #Another thing that needs to be changed. Id prefer FahrenheitToCelsius() method take care of this for me
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.InitialDataPointCheck == True:
            pass
        elif AssayDatabaseDictionary[0]['initial_boiling_point'] != None:
            if AssayDatabaseDictionary[0]['initial_boiling_point'] == 'C5' or AssayDatabaseDictionary[0]['initial_boiling_point'] == 'c5':
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBP = BPPentaneCelsius
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBPCheck = True
            elif AssayDatabaseDictionary[0]['initial_boiling_point'] == 'C4' or AssayDatabaseDictionary[0]['initial_boiling_point'] == 'c4':
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBP = BPButaneCelsius
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBPCheck = True
            else:
                try:
                    AssayDatabaseDictionary[0]['initial_boiling_point'] = float(AssayDatabaseDictionary[0]['initial_boiling_point'])
                except:
                    pass
                if isinstance( AssayDatabaseDictionary[0]['initial_boiling_point'],float):
                    print("========== This is the TemporaryTemperature Data=========" + to_str(AssayDatabaseDictionary[0]['initial_boiling_point']))
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBP = AssayDatabaseDictionary[0]['initial_boiling_point']
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBPCheck = True
                    
                else:
                    self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.IBPCheck = False
                
        
        TemporaryTemperatureData.extend(AssayDatabaseDictionary[0]["final_temperature"])
        #Check to see if we have a final data point. If one doesnt exist, then we will have to estimate it using baryecentric extrapolation
        if TemporaryTemperatureData[-1] == None:
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FinalDataPointCheck = False
        else:
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FinalDataPointCheck = True
        if AssayDatabaseDictionary[0]['final_boiling_point'] == None:
            pass
        else:
            try:
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FBP = to_float(AssayDatabaseDictionary[0]['final_boiling_point'])
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FBPCheck = True
            except:
                self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FBP = None
                #QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Something went wrong with the FBP data', QtWidgets.QMessageBox.Ok)

        print("Does Assay have initial Data Point?: " + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.InitialDataPointCheck))
        print("Does Assay have final Data Point?: " + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FinalDataPointCheck))

        print("This is the Temporary temperature data : " + to_str(TemporaryTemperatureData))
        TemporaryTemperatureData = remove_none_from_list(TemporaryTemperatureData)
        print(TemporaryTemperatureData)
        TemporaryTemperatureData.sort()
        index = 1
        #The following removes duplicates from the list
        while index < len(TemporaryTemperatureData):
            if TemporaryTemperatureData[index] == TemporaryTemperatureData[index-1]:
                del TemporaryTemperatureData[index]
            else:
                index += 1
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Temperature_Data = TemporaryTemperatureData
        print("After sorting the new temeprature data values are as follows:" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Temperature_Data))


        #This is for cases in which the crude finishes evaporating before the final temperature
        for i in range(0,len(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data)):
            pass

        for i in range(0,2): #Now we resolve and complete any incompleted Specific Gravity/Weight Percent Data. Iterate twice in case we have API Data
            print("i = " + str(i))
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.GetWeightData()
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.GetSGData()
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.GetAPIData()
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_to_CumulWeight()
        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Vol_to_CumulVol()
        #self.MyAssay.Fahrenheit_to_Celsius()
        
        print("This is the Volume data before end:" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data))
        #if IBPCheck == False: @@@@@@@@@@@@@@@@@@@@@
        print("This is the calculated IBP:" + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FindIBP()))
        #if FBPCheck == False: @@@@@@@@@@@@@@@@@@@
        print("This is the calculated FBP:" + str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.FindFBP()))
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CelsiusCheck == False:
            self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Fahrenheit_to_Celsius()

        # MyAssay.Celsius_to_Fahrenheit()
        print("This is the temp data before CompleteCurves():" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Temperature_Data))
        print("This is the SG Data before CompleteCurves()" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data))
        print("This is the Volume Data before CompleteCurves()" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data))
        print("This is the Weight Data before CompleteCurves()" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_Data))
        #print(self.MyAssay.FindIBP())
        #print(self.MyAssay.FindFBP())


        self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.CompleteCurves()
        print("This is the SG Data after Complete Curves for: " + to_str(self.SelectedRefinery) + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data))
        print("This is the temp data after CompleteCurves():" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Temperature_Data))
        print("This is the SG Data after CompleteCurves()" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.SG_Data))
        print("This is the Volume Data after CompleteCurves()" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Volume_Data))
        print("This is the Weight Data after CompleteCurves()" + to_str(self.Refinery_List[self.SelectedRefinery].ProductsInstance.AssayInstance.Weight_Data))


        self.MyDatabase.DisconnectFromDatabase()

        return
    def OcrSelected(self): #This function is connected to a button
        self.dlg.lbl_SourceName.hide()
        self.dlg.cmb_SourceName.hide() 
        self.dlg.cmd_UploadUsingOCR.show()
        self.dlg.cmd_GenerateCHEMEXTemplate.show()
        #self.dlg.spin_rowCount.hide() #Where is this coming from??
        return
    def UploadFromDatabaseSelected(self):
        self.dlg.lbl_SourceName.show()
        self.dlg.cmb_SourceName.show() 
        self.dlg.cmd_UploadUsingOCR.hide()
        self.dlg.cmd_GenerateCHEMEXTemplate.hide()
        return
    def UploadUsingOCR(self):
        try:
            self.scanner = OCRScanner(self)
            self.scanner.dlg.show()
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Using OCR Scanner. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def ColumnFromLabel(self,label): #https://stackoverflow.com/questions/22306899/pyqt-get-cell-from-tableview-with-header-label
        #This function returns the index of a column based on the label of that column in a table
        i = 0
        model = self.dlg.tbl_assay.horizontalHeader().model()
        for column in range(model.columnCount()):
            if model.headerData(column,QtCore.Qt.Horizontal) == label:
                return i
            i += 1
        return 
    def AddAssayInfoToTable(self):
        #Adds all of the information from the database to the Assay Viewer Table

        self.MyDatabase.ConnectToDatabase()
        self.ClearAssayTable()
        for i in range(self.dlg.tbl_assay.rowCount()-1,0,-1):
            self.RemoveRow()
        AssayRegion = ''
        AssayName = ''
        AssaySource = ''
        AssayYear = ''
        AssayCountry = ''

        SortCount = 0
        if self.dlg.cmb_SourceName.currentText() == 'Select Source Name':
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please Select Source Name to View Assay in Table', QtWidgets.QMessageBox.Ok)
            return

        SelectedAssay = self.dlg.cmb_SourceName.currentText()
        #The following line is regex. The link will take you to the stackoverflow code. This is a popular way to find and replace specific strings
        WholeAssay = re.findall('\((.*?)\)',SelectedAssay) #https://stackoverflow.com/questions/4894069/regular-expression-to-return-text-between-parenthesis
        AssayRegion = WholeAssay[0]
        AssayName = WholeAssay[1]
        AssayWholeSource = WholeAssay[2] #we will break this down in to its respective components next
        TempAssayWholeSource = AssayWholeSource #This is a temporary value, because we will use the split function on it
        AssayWholeSourceComponents = re.findall('\<(.*?)\>',SelectedAssay) #this regex expression returns everything between <> in a list

        for i in range(0,len(AssayWholeSourceComponents)-1): #This should never get above 2 iterations
            if AssayWholeSourceComponents[i].isdigit():
                AssayYear = AssayWholeSourceComponents[i]
            else:
                AssayCountry = AssayWholeSourceComponents[i]
            pass
        AssaySource = TempAssayWholeSource.split(' <')[0]
        self.MyDatabase.cursor.execute("SELECT * FROM " + self.MyDatabase.tableName + " WHERE assayregion = '" + AssayRegion + "' AND assayname = '" + AssayName + "' AND assaysource = '" + AssayWholeSource + "'")
        PostgresTableData = self.MyDatabase.cursor.fetchone()
        print(PostgresTableData)
        #assign all floats/single entries
        MethaneData = PostgresTableData['methane_data']
        EthaneData = PostgresTableData['ethane_data']
        PropaneData = PostgresTableData['propane_data']
        IsobutaneData = PostgresTableData['isobutane_data']
        NButaneData = PostgresTableData['n_butane_data']
        NPentaneData = PostgresTableData['n_pentane_data']
        IsopentaneData = PostgresTableData['i_pentane_data']
        CyclopentaneData = PostgresTableData['cyclopentane_data']
        CyclohexaneData = PostgresTableData['cyclohexane_data']
        BenzeneData = PostgresTableData['benzene_data']
        T1Viscosity1 = PostgresTableData['t1_viscosity1']
        T2Viscosity2 = PostgresTableData['t2_viscosity2']
        IBP = PostgresTableData['initial_boiling_point']
        FBP = PostgresTableData['final_boiling_point']
        #assign all arrays
        Initial_Temperature_Data = PostgresTableData['initial_temperature']
        Final_Temperature_Data = PostgresTableData['final_temperature']
        Weight_Yield_Data = PostgresTableData['weight_yield']
        Volume_Yield_Data = PostgresTableData['volume_yield']
        Density_Data = PostgresTableData['density_data']
        Sulfur_Data = PostgresTableData['sulfur_data']
        Mercaptan_Sulfur_Data = PostgresTableData['mercaptan_data']
        Smoke_Point_Data = PostgresTableData['smoke_point']
        Viscosity_T1_Data = PostgresTableData['viscosity1_data']
        Viscosity_T2_Data = PostgresTableData['viscosity2_data']
        TAN_Data = PostgresTableData['tan_data']
        Molecular_Weight_Data = PostgresTableData['molecular_weight']
        Salt_Content_Data = PostgresTableData['salt_data']
        Flash_Point_Data = PostgresTableData['flash_point']
        Freeze_Point_Data = PostgresTableData['freeze_point']
        Cloud_Point_Data = PostgresTableData['cloud_point']
        Pour_Point_Data = PostgresTableData['pour_point']
        RON_Data = PostgresTableData['ron_data']
        MON_Data = PostgresTableData['mon_data']
        Cetane_Data = PostgresTableData['cetane_number']
        RVP_Data = PostgresTableData['rvp_data']
        TVP_Data = PostgresTableData['tvp_data']
        Aromatics_Data = PostgresTableData['aromatics_data']
        Olefins_Data = PostgresTableData['olefins_data']
        Paraffins_Data = PostgresTableData['paraffins_data']
        Naphthenes_Data = PostgresTableData['naphthenes_data']
        Vanadium_Data = PostgresTableData['vanadium_data']
        Nickel_Data = PostgresTableData['nickel_data']
        Nitrogen_Data = PostgresTableData['nitrogen_data']
        #assign all units
        InitialTempDataUnit = PostgresTableData['initialtemperatureunit']
        FinalTempDataUnit = PostgresTableData['finaltemperatureunit']
        WeightYieldDataUnit = PostgresTableData['weightyieldunit']
        VolumeYieldDataUnit = PostgresTableData['volumeyieldunit']
        DensityDataUnit = PostgresTableData['densitydataunit']
        SulfurDataUnit = PostgresTableData['sulfurdataunit']
        MercaptanSulfurDataUnit = PostgresTableData['mercaptandataunit']
        SmokePointDataUnit = PostgresTableData['smokepointunit']
        ViscosityT1DataUnit = PostgresTableData['viscosity1dataunit']
        ViscosityT2DataUnit = PostgresTableData['viscosity2dataunit']
        TANDataUnit = PostgresTableData['tandataunit']
        MolecularWeightDataUnit = PostgresTableData['molecularweightunit']
        SaltContentDataUnit = PostgresTableData['saltdataunit']
        FlashPointDataUnit = PostgresTableData['flashpointunit']
        FreezePointDataUnit = PostgresTableData['freezepointunit']
        CloudPointDataUnit = PostgresTableData['cloudpointunit']
        PourPointDataUnit = PostgresTableData['pourpointunit']
        RVPDataUnit = PostgresTableData['rvpdataunit']    
        TVPDataUnit = PostgresTableData['tvpdataunit']
        AromaticsDataUnit = PostgresTableData['aromaticsdataunit']
        OlefinsDataUnit = PostgresTableData['olefinsdataunit']
        ParaffinsDataUnit = PostgresTableData['paraffinsdataunit']
        NaphthenesDataUnit = PostgresTableData['naphthenesdataunit']
        VanadiumDataUnit = PostgresTableData['vanadiumdataunit']
        NickelDataUnit = PostgresTableData['nickeldataunit']
        NitrogenDataUnit = PostgresTableData['nitrogendataunit']
        #Assign all Single Value Units
        T1Viscosity1Unit = PostgresTableData['t1viscosity1unit']
        T2Viscosity2Unit = PostgresTableData['t2viscosity2unit']
        MethaneDataUnit = PostgresTableData['methanedataunit']
        EthaneDataUnit = PostgresTableData['ethanedataunit']
        PropaneDataUnit = PostgresTableData['propanedataunit']
        IsobutaneDataUnit = PostgresTableData['isobutanedataunit']
        NButaneDataUnit = PostgresTableData['nbutanedataunit']
        NPentaneDataUnit = PostgresTableData['npentanedataunit']
        IsopentaneDataUnit = PostgresTableData['ipentanedataunit']
        CyclopentaneDataUnit = PostgresTableData['cyclopentanedataunit']
        CyclohexaneDataUnit = PostgresTableData['cyclohexanedataunit']
        BenzeneDataUnit = PostgresTableData['benzenedataunit']
        IBPDataUnit = PostgresTableData['initialtemperatureunit'] #NOTICE: the argument to PostgresTableData should be initialboilingpointunit, but because data does not currently exist for it, we will assume the same units as the start temps
        FBPDataUnit = PostgresTableData['finaltemperatureunit'] #NOTICE: the argument to PostgresTableData should be finalboilingpointunit, but because data does not currently exist for it, we will assume the same units as the end temps
        #Assign Cumulative Checked
        CumulativeChecked = PostgresTableData['cumulativecheck']

        #Next, we populate the Assay Table
        TableLength = len(Initial_Temperature_Data) #We couldve used any array in the database since they are all the same length
        #Clear Current Entries From the Table First
        for i in range(1,TableLength):
            self.AddRow()
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Initial Temp\n')).setText(to_str(Initial_Temperature_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Final Temp\n')).setText(to_str(Final_Temperature_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Weight Yield\n')).setText(to_str(Weight_Yield_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Volume Yield\n')).setText(to_str(Volume_Yield_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Density\n')).setText(to_str(Density_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Sulfur\n')).setText(to_str(Sulfur_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Mercaptan Sulfur\n')).setText(to_str(Mercaptan_Sulfur_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Smoke Point\n')).setText(to_str(Smoke_Point_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Viscosity @ T1\n')).setText(to_str(Viscosity_T1_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Viscosity @ T2\n')).setText(to_str(Viscosity_T2_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('TAN\n')).setText(to_str(TAN_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Molecular Weight\n')).setText(to_str(Molecular_Weight_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Salt Content\n')).setText(to_str(Salt_Content_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Flash Point\n')).setText(to_str(Flash_Point_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Freeze Point\n')).setText(to_str(Freeze_Point_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Cloud Point\n')).setText(to_str(Cloud_Point_Data[i]))
        if Pour_Point_Data != None: #TEMPORARY
            for i in range(0,len(Initial_Temperature_Data)):
                self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Pour Point\n')).setText(to_str(Pour_Point_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('RON\n')).setText(to_str(RON_Data[i]))
        if MON_Data != None: #TEMPORARY
            for i in range(0,len(Initial_Temperature_Data)):
                self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('MON\n')).setText(to_str(MON_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Cetane\n')).setText(to_str(Cetane_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('RVP\n')).setText(to_str(RVP_Data[i]))
        if TVP_Data != None: #TEMPORARY
            for i in range(0,len(Initial_Temperature_Data)):
                self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('TVP\n')).setText(to_str(TVP_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Aromatics\n')).setText(to_str(Aromatics_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Olefins\n')).setText(to_str(Olefins_Data[i]))
        if Paraffins_Data != None: #TEMPORARY
            for i in range(0,len(Initial_Temperature_Data)):
                self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Paraffins\n')).setText(to_str(Paraffins_Data[i]))
        if Naphthenes_Data != None: #TEMPORARY
            for i in range(0,len(Initial_Temperature_Data)):
                self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Naphthenes\n')).setText(to_str(Naphthenes_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Vanadium\n')).setText(to_str(Vanadium_Data[i]))
        for i in range(0,len(Initial_Temperature_Data)):
            self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Nickel\n')).setText(to_str(Nickel_Data[i]))
        if Nitrogen_Data != None: #TEMPORARY
            for i in range(0,len(Initial_Temperature_Data)):
                self.dlg.tbl_assay.item(i+1,self.ColumnFromLabel('Nitrogen\n')).setText(to_str(Nitrogen_Data[i]))
        

        if len(Initial_Temperature_Data) != len(Final_Temperature_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF FINAL TEMP DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Weight_Yield_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF WEIGHT YIELD DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Volume_Yield_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF VOLUME YIELD DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Density_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF DENSITY DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Sulfur_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF Sulfur DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Mercaptan_Sulfur_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF MERCAPTAN SULFUR DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Smoke_Point_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF SMOKE POINT DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Viscosity_T1_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF VISC @ T1 DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Viscosity_T2_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF VISC @ T2 DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(TAN_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF TAN DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Molecular_Weight_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF MOLECULAR WEIGHT DATA +++++++++++++++++++++++")
        if len(Initial_Temperature_Data) != len(Salt_Content_Data): print("+++++++++++++++++++ ERROR IN LENGTH OF MOLECULAR WEIGHT DATA +++++++++++++++++++++++")



        #Populate Table Units (Can be done with a for loop)
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Initial Temp\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Initial Temp\n')).findText(InitialTempDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Final Temp\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Final Temp\n')).findText(FinalTempDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Weight Yield\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Weight Yield\n')).findText(WeightYieldDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Volume Yield\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Volume Yield\n')).findText(VolumeYieldDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Density\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Density\n')).findText(DensityDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Sulfur\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Sulfur\n')).findText(SulfurDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Mercaptan Sulfur\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Mercaptan Sulfur\n')).findText(MercaptanSulfurDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Smoke Point\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Smoke Point\n')).findText(SmokePointDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Viscosity @ T1\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Viscosity @ T1\n')).findText(ViscosityT1DataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Viscosity @ T2\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Viscosity @ T2\n')).findText(ViscosityT2DataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('TAN\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('TAN\n')).findText(TANDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Molecular Weight\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Molecular Weight\n')).findText(MolecularWeightDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Salt Content\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Salt Content\n')).findText(SaltContentDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Flash Point\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Flash Point\n')).findText(FlashPointDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Freeze Point\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Freeze Point\n')).findText(FreezePointDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Cloud Point\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Cloud Point\n')).findText(CloudPointDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Pour Point\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Pour Point\n')).findText(PourPointDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('RVP\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('RVP\n')).findText(RVPDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('TVP\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('TVP\n')).findText(TVPDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Aromatics\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Aromatics\n')).findText(AromaticsDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Olefins\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Olefins\n')).findText(OlefinsDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Paraffins\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Paraffins\n')).findText(ParaffinsDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Naphthenes\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Naphthenes\n')).findText(NaphthenesDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Vanadium\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Vanadium\n')).findText(VanadiumDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Nickel\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Nickel\n')).findText(NickelDataUnit))
        self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Nitrogen\n')).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel('Nitrogen\n')).findText(NitrogenDataUnit))

        #Populate Units (can be done with a for loop)
        self.dlg.cmb_MethaneUnit.setCurrentIndex(self.dlg.cmb_MethaneUnit.findText(MethaneDataUnit))
        self.dlg.cmb_EthaneUnit.setCurrentIndex(self.dlg.cmb_EthaneUnit.findText(EthaneDataUnit))
        self.dlg.cmb_PropaneUnit.setCurrentIndex(self.dlg.cmb_PropaneUnit.findText(PropaneDataUnit))
        self.dlg.cmb_iButaneUnit.setCurrentIndex(self.dlg.cmb_iButaneUnit.findText(IsobutaneDataUnit))
        self.dlg.cmb_nButaneUnit.setCurrentIndex(self.dlg.cmb_nButaneUnit.findText(NButaneDataUnit))
        self.dlg.cmb_nPentaneUnit.setCurrentIndex(self.dlg.cmb_nPentaneUnit.findText(NPentaneDataUnit))
        self.dlg.cmb_iPentaneUnit.setCurrentIndex(self.dlg.cmb_iPentaneUnit.findText(IsopentaneDataUnit))
        self.dlg.cmb_CyclopentaneUnit.setCurrentIndex(self.dlg.cmb_CyclopentaneUnit.findText(CyclopentaneDataUnit))
        self.dlg.cmb_CyclohexaneUnit.setCurrentIndex(self.dlg.cmb_CyclohexaneUnit.findText(CyclohexaneDataUnit))
        self.dlg.cmb_BenzeneUnit.setCurrentIndex(self.dlg.cmb_BenzeneUnit.findText(BenzeneDataUnit))
        self.dlg.cmb_IBPUnit.setCurrentIndex(self.dlg.cmb_IBPUnit.findText(IBPDataUnit))
        self.dlg.cmb_FBPUnit.setCurrentIndex(self.dlg.cmb_FBPUnit.findText(FBPDataUnit))

        #Add Temperature viscosity Data and Other Assay Data
        self.dlg.txt_Methane.setText(to_str(MethaneData))
        self.dlg.txt_Ethane.setText(to_str(EthaneData))
        self.dlg.txt_Propane.setText(to_str(PropaneData))
        self.dlg.txt_iButane.setText(to_str(IsobutaneData))
        self.dlg.txt_nButane.setText(to_str(NButaneData))
        self.dlg.txt_nPentane.setText(to_str(NPentaneData))
        self.dlg.txt_iPentane.setText(to_str(IsopentaneData))
        self.dlg.txt_Cyclopentane.setText(to_str(CyclopentaneData))
        self.dlg.txt_Cyclohexane.setText(to_str(CyclohexaneData))
        self.dlg.txt_Benzene.setText(to_str(BenzeneData))
        self.dlg.txt_T1.setText(to_str(T1Viscosity1))
        self.dlg.txt_T2.setText(to_str(T2Viscosity2))
        self.dlg.txt_IBP.setText(to_str(IBP))
        self.dlg.txt_FBP.setText(to_str(FBP))
        self.dlg.cmb_T1Unit.setCurrentIndex(self.dlg.cmb_T1Unit.findText(to_str(T1Viscosity1Unit)))
        self.dlg.cmb_T2Unit.setCurrentIndex(self.dlg.cmb_T2Unit.findText(to_str(T2Viscosity2Unit)))
        self.dlg.txt_region.setText(AssayRegion)
        self.dlg.txt_assay.setText(AssayName)
        self.dlg.txt_source.setText(AssaySource)
        self.dlg.txt_year.setText(to_str(AssayYear))
        self.dlg.txt_country.setText(AssayCountry)

        #Set cumulative Checked
        self.dlg.chk_CumulativeCheck.setChecked(CumulativeChecked)

        self.MyDatabase.DisconnectFromDatabase()
        return
    def UpdateAssayViewer(self):
        try:
            if self.dlg.cmb_SourceName.currentText() == 'Select Source':
                self.ClearAssayTable()
                self.GraphAssayData()
                return
            self.AddAssayInfoToTable()
            self.GraphAssayData()
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Updating Assay Viewer. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
            raise Exception
        return
    def ClearAssayTable(self):
        for i in range(1,self.dlg.tbl_assay.rowCount()):
            for j in range(0,self.dlg.tbl_assay.columnCount()):
                self.dlg.tbl_assay.item(i,j).setText('')
        self.dlg.txt_Methane.setText('')
        self.dlg.txt_Ethane.setText('')
        self.dlg.txt_Propane.setText('')
        self.dlg.txt_iButane.setText('')
        self.dlg.txt_nButane.setText('')
        self.dlg.txt_nPentane.setText('')
        self.dlg.txt_iPentane.setText('')
        self.dlg.txt_Cyclopentane.setText('')
        self.dlg.txt_Cyclohexane.setText('')
        self.dlg.txt_Benzene.setText('')
        self.dlg.txt_T1.setText('')
        self.dlg.txt_T2.setText('')
        self.dlg.cmb_T1Unit.setCurrentIndex(self.dlg.cmb_T1Unit.findText(to_str('')))
        self.dlg.cmb_T2Unit.setCurrentIndex(self.dlg.cmb_T2Unit.findText(to_str('')))
        self.dlg.txt_region.setText('')
        self.dlg.txt_assay.setText('')
        self.dlg.txt_source.setText('')
        self.dlg.txt_year.setText(to_str(''))
        self.dlg.txt_country.setText('')
        return
    def UploadInfoFromOCRScanner(self):

        AssayTemplate = pe.get_array(file_name="assayscanner.xlsx", sheet_name = "INTERIM")
        #The Following are the all the possible strings that can exist to indicated that we have a specific kind of data
        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        Whole_Crude_List = ['crude', 'Crude', 'Whole crude', 'Whole Crude']
        IBP_List = ['IBP']
        C4_C7_List = ['c4','C4','c5','C5','c6','C6','c7','C7']
        FBP_List = ['FBP','+']
        Fahrenheit_List = ['F']
        Celsius_List = ['C']
        UnitsColumn = False
        if self.scanner!= None:
            UnitsColumn = self.scanner.dlg.chkUnits.isChecked()

        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        #This section is a little complicated. The "Indicators" are values in the units present (wether they are part of the property name or have their own column) that correspond to a specific unit
        # that will be used when filling out the Assay section of the dashboard. You will notice that the lengths of the Indicators are the same length as the units lists used in the comboboxes.
        # The indicators MUST be the same length, because later when we iterate through, we will look for the value in the indicator, then match it to its corresponding lists's index and that will be what
        # is filled out in the assay viewer combobox Unit.

        Temperature_Units_Indicator = [['F'], ['C'], [],[]]  # self.Temperature_Units_List = ['Fahrenheit', 'Celsius', 'Kelvin', 'Rankine']
        TAN_Units_Indicator = [['KOH']]     # self.TAN_Units_List = ['mg KOH/g']
        Weight_Units_Indicator = [['ppm','mg/kg'],['wt','WT'],['mg']]    # self.Weight_Percent_Units_List = ['ppm', 'wt%', 'mg/g']
        Percentage_Units_Indicator = [['vol','Vol','LV', 'VOL'],['wt','Wt','Wgt','WT']] # self.Percentage_Units_List = ['vol%', 'wt%']
        Density_Units_Indicator = [['SG'],['API'],['L'],['m^3', 'm3'],['gal', 'GAL'],['bbl','BBL'],['ft^3','ft3']] # self.Density_Units_List = ['SG', 'API', 'kg/L','kg/m^3', 'lb/gal', 'lb/bbl', 'lb/ft3']
        Distance_Units_Indicator = [['mm'],['ft'], ['cm'], ['m']] #self.Distance_Units_List = ['mm', 'ft' ,'cm', 'm']
        Viscosity_Units_Indicator = [['cst']] #self.Viscosity_Units_List = ['cSt']
        Molecular_Weight_Units_Indicator = [['g'],['lb']] #self.Molecular_Weight_Units_List = ['g/mol', 'lb/mol']
        Pressure_Units_Indicator = [['kPa'],['psi','Psi']] #self.Pressure_Units_List = ['kPa', 'psi']
        Molecular_Weight_Units_Indicator = [['g'], ['lb']] #self.Molecular_Weight_Units_List = ['g/mol', 'lb/mol']

        
        #NOTE: The following 2 sections of variables are meant to be in added to. If you add properties, or get a weird assay, the variables between the ~~~~ can be edited
        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        #This list represents the values we are looking for in the excel file that gives us an indication that we have a specific kind of data. If it finds a match, it will record the index where it found and eventually
        # it will record the data. These will need to be added to, as some assays have some very weird ways of displaying property names
        Weight_Percent_List = ['Cut Weight', 'Weight %', 'Cut Yield','CutYield']
        Volume_Percent_List = ['Cut volume', 'Cut Volume', 'Volume %', 'Volume percent', 'Volume Percent', 'Cut Yield','CutYield']
        SG_List = ['Specific Gravity', 'SG']
        API_List = ['API Gravity', 'API']
        Density_List = ['Density']
        Sulfur_List = ['Sulfur']
        Mercaptan_List = ['Mercaptans', 'Mercaptan']
        Smoke_Point_List = ['Smoke','smoke']
        TAN_List = ['TAN', 'Neutralization number', 'TotalAcidNumber','Total Acid Number']
        Nickel_List = ['Nickel']
        Vanadium_List = ['Vanadium']
        RON_List = ['Research Octane Number', 'RON']
        Salt_List = ['Salt content', 'Salt']
        Aromatics_List = ['Aromatics', 'Arom','AROM']
        Olefins_List = ['Olefins']
        Cloud_Point_List = ['Cloud']
        Freeze_Point_List = ['Freeze']
        RVP_List = ['RVP', 'Reid']
        TVP_List = ['TVP', 'TrueVaporPressure', 'True Vapor Pressure']
        Cetane_Number_List = ['Cetane Index', 'Cetane']
        Flash_Point_List = ['Flash Point', 'Flash point', 'Flashpoint','FlashPoint']
        Paraffins_List = ['paraffins','Paraffins']
        Pour_Point_List = ['pour','Pour']
        Nitrogen_List = ['nitrogen','Nitrogen','N2']
        TVP_List = ['TVP', 'TotalVaporPressure','Total Vapor Pressure']
        MON_List = ['MON','MotorOctane','Motor Octane']
        Molecular_Weight_List = ['MW','MolecularW','Molecular W']
        Naphthenes_List = ['Naphthene']
        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~  

        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        #The following are the column labels we will use to find where in the assay table we will put our values
        VolumePercentDataLabel = 'Volume Yield\n'
        WeightPercentDataLabel  = 'Weight Yield\n'
        SGDataLabel  = 'Density\n'
        APIDataLabel  = 'Density\n'
        SulfurDataLabel  = 'Sulfur\n'
        MercaptanDataLabel  = 'Mercaptan Sulfur\n'
        SmokePointDataLabel  = 'Smoke Point\n'
        TANDataLabel  = 'TAN\n'
        NickelDataLabel  = 'Nickel\n'
        VanadiumDataLabel  = 'Vanadium\n'
        RONDataLabel  = 'RON\n'
        MONDataLabel = 'MON\n'
        SaltDataLabel  = 'Salt Content\n'
        AromaticsDataLabel  = 'Aromatics\n'
        OlefinsDataLabel  = 'Olefins\n'
        CloudPointDataLabel  = 'Cloud Point\n'
        FreezePointDataLabel  = 'Freeze Point\n'
        RVPDataLabel  = 'RVP\n'
        TVPDataLabel = 'TVP\n'
        CetaneDataLabel  = 'Cetane\n'
        FlashPointDataLabel  = 'Flash Point\n'
        ParaffinsDataLabel = 'Paraffins\n'
        PourPointDataLabel = 'Pour Point\n'
        NitrogenDataLabel = 'Nitrogen\n'
        MolecularWeightDataLabel = 'Molecular Weight\n'
        TVPDataLabel = 'TVP\n'
        NaphthenesDataLabel = 'Naphthenes\n'

        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        # Volume% and Weight % I would like to fix. This is an instance of something I would like to change. It works for now, but could be better
        Data_List = [OCRDatapoint("Volume Percent",Volume_Percent_List, VolumePercentDataLabel,self.Percentage_Units_List,Percentage_Units_Indicator), OCRDatapoint("Weight Percent",Weight_Percent_List, WeightPercentDataLabel, self.Percentage_Units_List,Percentage_Units_Indicator), OCRDatapoint("SG",SG_List, SGDataLabel),
                     OCRDatapoint("API",API_List, APIDataLabel), OCRDatapoint("Sulfur",Sulfur_List, SulfurDataLabel, self.Weight_Percent_Units_List, Weight_Units_Indicator), OCRDatapoint("Mercaptan Sulfur",Mercaptan_List, MercaptanDataLabel, self.Weight_Percent_Units_List, Weight_Units_Indicator),
                     OCRDatapoint("Smoke Point",Smoke_Point_List, SmokePointDataLabel, self.Distance_Units_List, Distance_Units_Indicator), OCRDatapoint("TAN",TAN_List, TANDataLabel), OCRDatapoint("Nickel",Nickel_List, NickelDataLabel, self.Weight_Percent_Units_List, Weight_Units_Indicator),
                     OCRDatapoint("Vanadium",Vanadium_List, VanadiumDataLabel, self.Weight_Percent_Units_List, Weight_Units_Indicator), OCRDatapoint("RON",RON_List, RONDataLabel), OCRDatapoint("Salt Content",Salt_List, SaltDataLabel),
                     OCRDatapoint("Aromatics", Aromatics_List, AromaticsDataLabel, self.Percentage_Units_List, Percentage_Units_Indicator), OCRDatapoint("Olefins",Olefins_List, OlefinsDataLabel, self.Percentage_Units_List, Percentage_Units_Indicator), OCRDatapoint("Cloud Point",Cloud_Point_List, CloudPointDataLabel, self.Temperature_Units_List, Temperature_Units_Indicator),
                     OCRDatapoint("Freeze Point",Freeze_Point_List, FreezePointDataLabel,self.Temperature_Units_List, Temperature_Units_Indicator), OCRDatapoint("RVP",RVP_List, RVPDataLabel, self.Pressure_Units_List, Pressure_Units_Indicator), OCRDatapoint("Cetane",Cetane_Number_List, CetaneDataLabel),
                     OCRDatapoint("Flash Point",Flash_Point_List, FlashPointDataLabel, self.Temperature_Units_List, Temperature_Units_Indicator), OCRDatapoint("Density",Density_List, SGDataLabel, self.Density_Units_List, Density_Units_Indicator),
                     OCRDatapoint("Paraffins",Paraffins_List, ParaffinsDataLabel, self.Percentage_Units_List, Percentage_Units_Indicator), OCRDatapoint("Pour Point",Pour_Point_List, PourPointDataLabel, self.Temperature_Units_List, Temperature_Units_Indicator),
                     OCRDatapoint("Nitrogen",Nitrogen_List, NitrogenDataLabel, self.Weight_Percent_Units_List, Weight_Units_Indicator), OCRDatapoint("TVP",TVP_List, TVPDataLabel, self.Pressure_Units_List, Pressure_Units_Indicator),
                     OCRDatapoint("Naphthenes", Naphthenes_List, NaphthenesDataLabel, self.Percentage_Units_List, Percentage_Units_Indicator), OCRDatapoint("MON",MON_List, MONDataLabel),
                     OCRDatapoint("Molecular Weight", Molecular_Weight_List, MolecularWeightDataLabel, self.Molecular_Weight_Units_List, Molecular_Weight_Units_Indicator)]

        #Variables we are going to record
        IBP = None
        IndexOfIBPRange = None
        FBP = None
        IndexOfFBPRange = None
        IndexOfWholeCrude = None
        Recorded_Initial_Temperatures = []
        Recorded_Final_Temperatures = []
        TemperatureUnit = None
        WholeCrudeIndexOffset = 0
        UnitsColumnIndex = None

        Rows = len(AssayTemplate)
        Columns = len(AssayTemplate[0])
        TemporaryInfoArray = [[None]*Columns] * Rows

        print(AssayTemplate)
        print(TemporaryInfoArray)

        if UnitsColumn == True:
            for TempUnit in Fahrenheit_List:
                for i in range(1,len(AssayTemplate[0])):
                    if TempUnit in AssayTemplate[0][i]:
                        TemperatureUnit = 'Fahrenheit'
                        UnitsColumnIndex = i
            for TempUnit in Celsius_List:
                for i in range(1,len(AssayTemplate[0])):
                    if TempUnit in AssayTemplate[0][i] and TemperatureUnit == None:
                        TemperatureUnit = 'Celsius'
                        UnitsColumnIndex = i
                                

        #This gets us our temperature list so we can extract IBP and FBPs
        Temperatures = AssayTemplate[0]
        print("OG TEMPERATURES:" + str(Temperatures))
        i = 0
        while i < len(Temperatures):
            BreakCheck = False
            for WholeCrude in Whole_Crude_List:
                if WholeCrude in Temperatures[i]:
                    IndexOfWholeCrudeInTemperaturesList = i  #Only Applies in the Temps index
                    IndexOfWholeCrude = i + WholeCrudeIndexOffset
                    BreakCheck = True
                    i+=1
                    break
            for j in range(0,len(Temperatures[i])): #This checks to see if we have any weird data before our actual temperatures. We check this by seeing if any numbers are present
                if j == len(Temperatures[i]) - 1 and BreakCheck == False: #If we got through the whole string with no numbers, it means it is a useless piece of data and we will delete it from our temeprature array
                    if i == len(Temperatures)-1: #This means we are on the last index of the final point and there has been no numbers so we will check if it just says 'FBP'
                        for FBPValue in FBP_List:
                            if FBPValue in Temperatures[i]:
                                BreakCheck = True
                                i+=1
                                break
                    if BreakCheck == False:
                        del Temperatures[i]
                        WholeCrudeIndexOffset += 1 #Applies to the data sheet
                        BreakCheck = True
                        break
                try:
                    print(int(Temperatures[i][j]))
                    print("WE BROKE OUT ON INDEX" + str(i))
                    break
                except:
                    continue
            if BreakCheck == True:
                continue
            #if there are any blanks we watnt ot delete them
            if Temperatures[i] == '':
                del Temperatures[i]
                print("DELETED TEMPERATURE")
                WholeCrudeIndexOffset += 1                   
            #remove spaces to make the cuts easier to work with
            else:
                Temperatures[i] = Temperatures[i].replace(" ", "")
                i+=1
        print(Temperatures)


        # STEP 1
        #Now we will iterate through temperatures again to get useful information for IBP and FBPs
        i = 0
        while i < len(Temperatures):
            #We will now check each index to see if C4-C7 is one of our IBPs
            for C4_C7 in C4_C7_List:
                #If we find a match, and we have not already assigned an IBP, we will assign IBP to that C Number and record the index in the list it is at
                if C4_C7 in Temperatures[i] and IBP == None:
                    IBP = C4_C7
                    IndexOfIBPRange = i
                    break
            #We will now check each index to see if we have an FBP. If we get an FBP that matches our criteria, it is basically saying that FinalDataPointCheck is False
            for FBPValue in FBP_List:
                #We will aslso record the index of the FBP
                if FBPValue in Temperatures[i]  and FBP == None:
                    IndexOfFBPRange = i
            i+= 1
        
        #STEP 2
        #If the assay does not start at a specific Carbon number (C4-C7) we will check to see if it specifically says "IBP anywhere in the numbers
        if IBP == None:
            for i in range(0,len(Temperatures)):
                for IBPValue in IBP_List:
                    if IBPValue in Temperatures[i]:
                        IBP = 'IBP'
                        IndexOfIBPRange = i
                        print("This is the IBP" + str(IBP))
                        break
        
        #STEP 3
        #Now we will find the temperature unit used 
        for i in range(0,len(Temperatures)):
            for FahrenheitIndicator in Fahrenheit_List:
                if FahrenheitIndicator == Temperatures[i][-1] and TemperatureUnit == None:
                    TemperatureUnit = 'Fahrenheit'
                    print("Temperature Unit was found to be Fahrenheit")
            for CelsiusIndicator in Celsius_List:
                if CelsiusIndicator == Temperatures[i][-1] and TemperatureUnit == None:
                    TemperatureUnit = 'Celsius'
                    print("Temperature Unit was found to be Celsius")
        #Quick check to see if we foudn a unit. If we did not find a unit we cannot continue
        if TemperatureUnit == None:
            print("ERROR NO TEMPERATURE UNIT WAS FOUND")
            #return

        #STEP 4
        #We are now going to attempt to assign values to the Initial and Final Temperatures for the boiling point ranges
        BreakCheck = False
        IndexOfTheStartOfTemperatureValue = 0
        IndexOfTheEndOfTemperatureValue = 0
        IntegerStringStart = False #This is a Boolean that tells us when we have "Hit" a number and we can start extracting the data to figure out which numbers are the starting and ending temperatures for our cuts
        RangeCheck = False #This boolean determines if we are given a singular value for the temperature ranges for the different cuts or if it is just a singular value
        if IndexOfIBPRange == None:
            IndexOfIBPRange = 0
            if IndexOfWholeCrude != None:
                IndexOfIBPRange+=1
            print("PLEASE BEWARE OF DATA. AN IBP WAS NOT FOUND MANUALLY SO THE INDEX OF THE FIRST IBP IS ASSUMED TO BE 1")
        if IndexOfFBPRange == None:
            IndexOfFBPRange = len(Temperatures) - 1
            print("PLEASE BEWARE OF DATA. AN FBP WAS NOT FOUND MANUALLY SO THE INDEX OF THE FIRST IBP IS ASSUMED TO BE THE LAST INDEX OF THE TEMPERATURE VALUES")
        for i in range(IndexOfIBPRange, IndexOfFBPRange+1):
            BreakCheck = False
            TemperatureString = Temperatures[i]
            for j in range(0,len(TemperatureString)):
                # This is a subset of code for specifically the first index. This index is weird because it commonly contains 'C4-C7" or 'IBP'. In this case, we do not want to attempt to find a "starting" and "ending" point,
                # Our strategy is to iterate through the strings and find the intergers and append them to find starting and ending points to form coeherent values for Initial and Final Boiling Point Ranges
                # If the first value had "C5-180F" , the algorithm would pick up 5 as the Initial Boiling temperature and 180 as the final boiling temperature. So in reality, for the first point only, we just need the
                # Final boiling temperature. NOTE: This code will not likely work if a singular value is given for the IBP for example 20
                if i == IndexOfIBPRange:
                    for C4_C7 in C4_C7_List:
                        # If they have specified an IBP or a C5 starting point then we will continue with the code, otherwise they most likely just gave us something like 20-50 for the first point and the normal method will work
                        if C4_C7 in Temperatures[i] and '-' in Temperatures[i] or 'IBP' in Temperatures[i]: #I do not like this. I would like a way to iterate the IBP list as well. Will require some restructuring but will work for now
                            #If we go down this path, we confirmed have the special case, and we will iterate backwards through the strign to find the final boiling temperature for the first cut
                            for k in range(len(TemperatureString)-1,-1,-1):
                                try: 
                                    int(TemperatureString[k])
                                    if IntegerStringStart == False:
                                        #If this is our first number found, it is the beginning of the final boiling temperature
                                        IndexOfTheEndOfTemperatureValue = k
                                        print(IndexOfTheEndOfTemperatureValue)
                                        IntegerStringStart = True
                                except:
                                    if IntegerStringStart == False:
                                        continue
                                    if IntegerStringStart == True:
                                        #If this code executes, that means that this is our first occurence of a "fail" after we found and integer, meaning we are at the end of finding our final boiling temperature
                                        IndexOfTheStartOfTemperatureValue = k+1
                                        print(IndexOfTheStartOfTemperatureValue)
                                        IntegerStringStart = False
                                        Recorded_Initial_Temperatures.append(IBP)
                                        Recorded_Final_Temperatures.append(TemperatureString[IndexOfTheStartOfTemperatureValue:IndexOfTheEndOfTemperatureValue+1])
                                        print("The following is the Final Boiling Point Tmeperature Recorded for the IBP range of the Assay:" + str(Recorded_Final_Temperatures[0]))
                                        IndexOfTheStartOfTemperatureValue = 0
                                        IndexOfTheEndOfTemperatureValue = 0
                                        BreakCheck = True
                                        break
                            if BreakCheck == True:
                                break
                    if len(Recorded_Final_Temperatures) == 1:
                        break
                    elif len(Recorded_Final_Temperatures) == 0: #This is the case where either there was no number or it was a singular number, Ex 1: 60. Ex 2: C5 (The C5/C4 case only should not exist
                        #Need case to deal with C4 only, as well as a singular value.
                        # We are going to find and record the first number that we find, and input the IBP as None
                        IBPUpperRange = ''
                        for k in range(0,len(TemperatureString)):
                            if TemperatureString[k].isdigit():
                                IBPUpperRange+=(TemperatureString[k])
                        if IBPUpperRange == '':
                            print("THERE IS NO IBP DATA PRESENT")
                            return
                        else:
                            Recorded_Final_Temperatures.append(int(IBPUpperRange))
                            Recorded_Initial_Temperatures.append('IBP')

                #This is the normal Try Block
                try:
                    int(TemperatureString[j])
                    if IntegerStringStart == False:
                        IndexOfTheStartOfTemperatureValue = j
                        IntegerStringStart = True
                    if j == len(TemperatureString)-1:
                        Recorded_Initial_Temperatures.append(None)
                        Recorded_Final_Temperatures.append(TemperatureString)
                        IndexOfTheStartOfTemperatureValue = 0
                except:
                    if i == (IndexOfFBPRange+1) -1 and BreakCheck == False:
                        for FBPValue in FBP_List:
                            if FBPValue in Temperatures[i]:
                                Recorded_Initial_Temperatures.append(Recorded_Final_Temperatures[-1])
                                Recorded_Final_Temperatures.append('FBP')
                                BreakCheck = True
                                IntegerStringStart = False
                                break
                    if IntegerStringStart == False:
                        continue
                    if IntegerStringStart == True:
                        #If this code executes, that means that this is our first occurence of a "fail" after we found and integer, meaning we are at the end of finding our final boiling temperature
                        IndexOfTheEndOfTemperatureValue = j-1
                        IntegerStringStart = False
                        #We will now venture off into the string to determine if we are given a range or a singular value
                        #If we are given a range, we will assign the number to the Recorded Initial temperature but if not, we will assign it to the final recorded temperature
                        #Example:  starting bp -> 65-120 <-final bp  VS  120 <- final bp
                        for k in range(j,len(TemperatureString)):
                            if TemperatureString[k] == '-':
                                RangeCheck = True
                        if RangeCheck == True and len(Recorded_Initial_Temperatures) == len(Recorded_Final_Temperatures):
                            Recorded_Initial_Temperatures.append(TemperatureString[IndexOfTheStartOfTemperatureValue:IndexOfTheEndOfTemperatureValue+1])
                            continue
                        elif RangeCheck == True and len(Recorded_Initial_Temperatures) != len(Recorded_Final_Temperatures):
                            Recorded_Final_Temperatures.append(TemperatureString[IndexOfTheStartOfTemperatureValue:IndexOfTheEndOfTemperatureValue+1])
                            RangeCheck = False
                            IndexOfTheStartOfTemperatureValue = 0
                            IndexOfTheEndOfTemperatureValue = 0
                            print("Initial Recorded Temperatures: " + str(Recorded_Initial_Temperatures))
                            print("Final Recorded Temperatures: " + str(Recorded_Final_Temperatures))
                            break
                        if RangeCheck == False:
                            if i == IndexOfFBPRange:
                                Recorded_Initial_Temperatures.append(TemperatureString[IndexOfTheStartOfTemperatureValue:IndexOfTheEndOfTemperatureValue+1])
                                Recorded_Final_Temperatures.append('FBP')
                            else:
                                Recorded_Initial_Temperatures.append(None)
                                Recorded_Final_Temperatures.append(TemperatureString[IndexOfTheStartOfTemperatureValue:IndexOfTheEndOfTemperatureValue+1])
                            RangeCheck = False
                            IndexOfTheStartOfTemperatureValue = 0
                            IndexOfTheEndOfTemperatureValue = 0
                            print("Initial Recorded Temperatures: " + str(Recorded_Initial_Temperatures))
                            print("Final Recorded Temperatures: " + str(Recorded_Final_Temperatures))
                            break
        for i in range(1,len(Recorded_Initial_Temperatures)):
            if Recorded_Initial_Temperatures[i] == None:
                Recorded_Initial_Temperatures[i] = Recorded_Final_Temperatures[i-1]
        print("FINAL VALUES: Initial Recorded Temperatures: " + str(Recorded_Initial_Temperatures))
        print("FINAL VALUES: Final Recorded Temperatures: " + str(Recorded_Final_Temperatures))

        #STEP 5
        #Now we will clean up the IBP and FBP ranges so that they are valid numbers 
        for C4_C7 in C4_C7_List:
            if C4_C7 in Recorded_Initial_Temperatures[0] or 'IBP' in Recorded_Initial_Temperatures[0]:
                Recorded_Initial_Temperatures[0] = None
                break
        for FBPValue in FBP_List:
            if FBPValue in Recorded_Final_Temperatures[-1]:
                Recorded_Final_Temperatures[-1] = None
                break

        #STEP 6
        #Now we will iterate through the array and find what data we have and what we dont. We will also record the index we find the particular data in so we may enter it directly in to our table
        DensityCheck = False
        Property_Index_List = [] #The purpose of this is another check to ensure that one of our indeces have not already been used
        for OCRData in Data_List:
            for i in range(0, len(AssayTemplate)):
                BreakCheck = False
                for j in range(0,len(OCRData.Property_Identifier_List)):
                    if OCRData.PropertyIndex == None and OCRData.Property_Identifier_List[j] in AssayTemplate[i][0]:
                        for k in range(0,len(Property_Index_List)): #Ensures that this property has not already been added
                            if i == Property_Index_List[k]:
                                BreakCheck = True
                                break
                        if BreakCheck == True:
                            break
                        Property_Index_List.append(i)
                        OCRData.PropertyIndex = i
                        OCRData.HasProperty = True
                        print(to_str(OCRData.Name) + ": " + to_str(OCRData.PropertyIndex))
                        if (OCRData.Name == "API" or OCRData.Name == "SG" or OCRData.Name == "Density"): #This is here for the purpose of catching densities before untis are assigned. I would definetely like to change the structure of densities and their corresponding units
                            if DensityCheck == False:
                                DensityCheck = True
                                if(OCRData.Name == "API" or OCRData.Name == "SG"):
                                    OCRData.Unit = OCRData.Name
                                    print("LOOK HERE FOR THE DENSITY RELATED THINGS:" + OCRData.Unit)
                                    pass
                        elif OCRData.Unit_List_Indicator != None:
                            for k in range(0,len(OCRData.Unit_List_Indicator)):
                                for l in range(0,len(OCRData.Unit_List_Indicator[k])):
                                    if UnitsColumn == False:
                                        UnitStartingIndex = AssayTemplate[i][0].find(',')
                                        if UnitStartingIndex == -1:
                                            UnitStartingIndex = AssayTemplate[i][0].find('(') #Many times the units are kept in parenthesis ex: (Fahrenheit)
                                        print(AssayTemplate[i][0][UnitStartingIndex+1:])
                                        if OCRData.Unit == None and UnitStartingIndex != -1 and OCRData.Unit_List_Indicator[k][l] in AssayTemplate[i][0][UnitStartingIndex+1:]:
                                            OCRData.UnitTableName = OCRData.Unit_List[k]
                                            OCRData.Unit = OCRData.Unit_List[k]
                                            print("Unit: " + OCRData.Unit)
                                            break
                                    elif UnitsColumn == True:
                                        if OCRData.Unit == None and OCRData.Unit_List_Indicator[k][l] in AssayTemplate[i][UnitsColumnIndex]:
                                            OCRData.UnitTableName = OCRData.Unit_List[k]
                                            OCRData.Unit = OCRData.Unit_List[k]
                                            print("Unit: " + OCRData.Unit)
                                            break
                                        
                                       
        #STEP 7
        #Now that we know what data we have, and which row it is in the array, we can officially populate the data point with the actual data, based on our index range of our boiling point temperatures
        for OCRData in Data_List:
            for i in range(IndexOfIBPRange+WholeCrudeIndexOffset, IndexOfFBPRange+WholeCrudeIndexOffset+1): #We add +1 and +2 because we must remember our first row is the property name itself
                if OCRData.HasProperty:
                    OCRData.Data_List.append(AssayTemplate[OCRData.PropertyIndex][i])
            if IndexOfWholeCrude != None and OCRData.HasProperty:
                OCRData.Data_List.insert(0,to_str(AssayTemplate[OCRData.PropertyIndex][IndexOfWholeCrude]))
            print(OCRData.Data_List)
            print(OCRData.PropertyIndex)
                
        #STEP 8
        # Now our data is organized and we can insert it in to our table by referencing the Data label for each data point
        self.ClearAssayTable()
        for i in range(self.dlg.tbl_assay.rowCount()-1,0,-1):
            self.RemoveRow()
        for i in range(0,len(Data_List)):
            if Data_List[i].HasProperty:
                TableLength = len(Data_List[i].Data_List) #We couldve used any array in the database since they are all the same length
                break
        #Add Rows to the table
        for i in range(1,TableLength):
            self.AddRow()
        StartRowIndex = 2
        #Here we want to ensure the data is high quality before entering it in to the table
        for OCRData in Data_List:
            if OCRData.HasProperty:
                for i in range(0,len(OCRData.Data_List)): #We are going to check each piece of data
                    j = 0
                    OCRData.Data_List[i] = OCRData.Data_List[i].replace(" ","")
                    while j < len(OCRData.Data_List[i]):
                        try:
                            int(OCRData.Data_List[i][j])
                            j += 1
                        except:
                            if OCRData.Data_List[i][j] == ')' and j == len(OCRData.Data_List[i])-1: #This is a result of Erin's code such that it will put negative numbers in parenthasis, and we will correct it later
                                j+=1
                                continue
                            elif j == len(OCRData.Data_List[i])-1: #This means we reached the end of the string and its not an integer, or a ), we delete it
                                OCRData.Data_List[i] = OCRData.Data_List[i].replace(OCRData.Data_List[i][j],'',1)
                                continue
                            if OCRData.Data_List[i][j] == '-' or OCRData.Data_List[i][j] == '.' or '(' in OCRData.Data_List[i]: 
                                if OCRData.Data_List[i][j+1].isdigit(): #This could mean we have negative number, a decimal, or be in the parenthasis meaning there was a negative number
                                    j+=1
                                    continue
                            else:
                                OCRData.Data_List[i] = OCRData.Data_List[i].replace(OCRData.Data_List[i][j],'',1)
                                continue
                            OCRData.Data_List[i] = OCRData.Data_List[i].replace("-","")
                    if '(' in OCRData.Data_List[i] and ')' in OCRData.Data_List[i]: #this means we have a negative number
                        OCRData.Data_List[i] = OCRData.Data_List[i].replace("(","")
                        OCRData.Data_List[i] = OCRData.Data_List[i].replace(")","")
                        OCRData.Data_List[i] = '-' + OCRData.Data_List[i]

        if IndexOfWholeCrude != None:
            StartRowIndex += -1
        for OCRData in Data_List:
            for i in range(0,len(OCRData.Data_List)):
                print(i+StartRowIndex)
                print(self.dlg.tbl_assay.item(i+StartRowIndex,self.ColumnFromLabel(OCRData.AssayTableIdentifier)).text())
                if self.dlg.tbl_assay.item(i+StartRowIndex,self.ColumnFromLabel(OCRData.AssayTableIdentifier)).text() == '':
                    self.dlg.tbl_assay.item(i+StartRowIndex,self.ColumnFromLabel(OCRData.AssayTableIdentifier)).setText(to_str(OCRData.Data_List[i]))
            if OCRData.Unit != None:
                self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel(OCRData.AssayTableIdentifier)).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel(OCRData.AssayTableIdentifier)).findText(OCRData.Unit))
        for i in range(0,len(Recorded_Initial_Temperatures)):
            self.dlg.tbl_assay.item(i+2,self.ColumnFromLabel("Initial Temp\n")).setText(to_str(Recorded_Initial_Temperatures[i]))
            self.dlg.tbl_assay.item(i+2,self.ColumnFromLabel("Final Temp\n")).setText(to_str(Recorded_Final_Temperatures[i]))
        if TemperatureUnit != None:
            self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel("Initial Temp\n")).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel("Initial Temp\n")).findText(TemperatureUnit))
            self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel("Final Temp\n")).setCurrentIndex(self.dlg.tbl_assay.cellWidget(0,self.ColumnFromLabel("Final Temp\n")).findText(TemperatureUnit))
        self.dlg.txt_assay.setText(self.scanner.dlg.txtAssay.text())
        self.dlg.txt_source.setText(self.scanner.dlg.txtSource.text())

        return
    def GenerateCHEMEXTemplate(self):
        #This function takes the assay that is on the Assay Viewer Page, and fills out the CHEMEX Custom Assay excel file with the correct data and units. NOTE this function does not pull from the Assay Viewer
        # interface, it pulls directly fromt he Postgres Database
        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        CHEMEXAssayTemplateWorkbook = load_workbook(os.getcwd() + "\CHEMEXBrandedAssayTemplate.xlsx")
        Sheets = CHEMEXAssayTemplateWorkbook.sheetnames
        CHEMEXAssayTemplate = CHEMEXAssayTemplateWorkbook[Sheets[0]]
        MaxRows = 94 #This is simply from observing the excel file. IF CHANGES ARE MADE PLEASE UPDATE THIS PARAMETER
        MaxColumns = 34 #This is simply from observing the excel file. IF CHANGES ARE MADE PLEASE UPDATE THIS PARAMETER
        TopLeftIndexingKey = "Start Temperature"
        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        BreakCheck = False
        CHEMEXAssayTemplateWorkbook.save(os.getcwd() + "\CHEMEXBrandedAssayTemplate.xlsx")
        for i in range(1,MaxRows): #This for loop is just to find where we start, based on what our top left indicator is
            if BreakCheck:
                break
            for j in range(1,MaxColumns):
                if TopLeftIndexingKey in str(CHEMEXAssayTemplate.cell(row = i, column = j).value):
                    TopLeftIndexingRow = i
                    TopLeftIndexingColumn = j
                    BreakCheck = True
                    break
        
        #This section of code connects to the database, and assigns the values to a dictionary called PostgresTableData
        self.MyDatabase.ConnectToDatabase()
        self.ClearAssayTable()
        for i in range(self.dlg.tbl_assay.rowCount()-1,0,-1):
            self.RemoveRow()
        AssayRegion = ''
        AssayName = ''
        AssaySource = ''
        AssayYear = ''
        AssayCountry = ''

        SortCount = 0
        if self.dlg.cmb_SourceName.currentText() == 'Select Source Name':
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please Select Source Name to View Assay in Table', QtWidgets.QMessageBox.Ok)
            return

        SelectedAssay = self.dlg.cmb_SourceName.currentText()
        #The following line is regex. The link will take you to the stackoverflow code. This is a popular way to find and replace specific strings
        WholeAssay = re.findall('\((.*?)\)',SelectedAssay) #https://stackoverflow.com/questions/4894069/regular-expression-to-return-text-between-parenthesis
        AssayRegion = WholeAssay[0]
        AssayName = WholeAssay[1]
        AssayWholeSource = WholeAssay[2] #we will break this down in to its respective components next
        TempAssayWholeSource = AssayWholeSource #This is a temporary value, because we will use the split function on it
        AssayWholeSourceComponents = re.findall('\<(.*?)\>',SelectedAssay) #this regex expression returns everything between <> in a list

        for i in range(0,len(AssayWholeSourceComponents)-1): #This should never get above 2 iterations
            if AssayWholeSourceComponents[i].isdigit():
                AssayYear = AssayWholeSourceComponents[i]
            else:
                AssayCountry = AssayWholeSourceComponents[i]
            pass
        AssaySource = TempAssayWholeSource.split(' <')[0]
        self.MyDatabase.cursor.execute("SELECT * FROM " + self.MyDatabase.tableName + " WHERE assayregion = '" + AssayRegion + "' AND assayname = '" + AssayName + "' AND assaysource = '" + AssayWholeSource + "'")
        PostgresTableData = self.MyDatabase.cursor.fetchone()
        if PostgresTableData == None:
            return

        #First lets fill out the Assay Name, Source, and Year
        BreakCheck = False #Having a breakcheck here allows us to break out of both for loops when we reach our desired result and saves time 
        for i in range(1,MaxColumns):
            if BreakCheck == False:
                for j in range(1,MaxRows):
                    if CHEMEXAssayTemplate.cell(row = i, column = j).value != None and "ASSAYNAME" in CHEMEXAssayTemplate.cell(row = i, column = j).value:
                        CHEMEXAssayTemplate.cell(row = i, column = j).value = CHEMEXAssayTemplate.cell(row = i, column = j).value.replace("ASSAYNAME",AssayName)
                    if CHEMEXAssayTemplate.cell(row = i, column = j).value != None and "Name:" in CHEMEXAssayTemplate.cell(row = i, column = j).value:
                        CHEMEXAssayTemplate.cell(row = i, column = j).value = CHEMEXAssayTemplate.cell(row = i, column = j).value + ' ' + AssayName
                    if CHEMEXAssayTemplate.cell(row = i, column = j).value != None and "Origin:" in CHEMEXAssayTemplate.cell(row = i, column = j).value:
                        CHEMEXAssayTemplate.cell(row = i, column = j).value = CHEMEXAssayTemplate.cell(row = i, column = j).value + ' ' + AssaySource
                    if  CHEMEXAssayTemplate.cell(row = i, column = j).value != None and "Year:" in CHEMEXAssayTemplate.cell(row = i, column = j).value:
                        CHEMEXAssayTemplate.cell(row = i, column = j).value = CHEMEXAssayTemplate.cell(row = i, column = j).value + ' ' + AssayYear
                        BreakCheck = True
                        break
            else:
                break
        BreakCheck = False




        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        # Here we are creating a dictionary. The first value is the property name that will be referenced in the CHEMEXBrandedAssayTemplate.xlsx and the second is a Datapoint instance with 3 values. The first value is
        # the data that is held in the database associated with the Excel Property name, and the second value is the corresponding unit. The third is the row index, which we will fill in after we iterate through the excel sheet.
        # Please reference the Sub_Classes class to see the datapoint class structure 
        CHEMEXBrandedAssayDictionary = dict({"Start Temperature": CHEMEXAssayTemplateDatapoint(PostgresTableData['initial_temperature'],PostgresTableData['initialtemperatureunit']), "End Temperature": CHEMEXAssayTemplateDatapoint(PostgresTableData['final_temperature'], PostgresTableData['finaltemperatureunit'])})
        CHEMEXBrandedAssayDictionary.update({"Volume % yield on crude": CHEMEXAssayTemplateDatapoint(PostgresTableData['volume_yield'], PostgresTableData['volumeyieldunit']), "Weight % yield on crude": CHEMEXAssayTemplateDatapoint(PostgresTableData['weight_yield'], PostgresTableData['weightyieldunit'])})
        CHEMEXBrandedAssayDictionary.update({"Density": CHEMEXAssayTemplateDatapoint(PostgresTableData['density_data'], PostgresTableData['densitydataunit']), "API Gravity": CHEMEXAssayTemplateDatapoint(PostgresTableData['density_data'], PostgresTableData['densitydataunit'])}) #This will be corrected later
        CHEMEXBrandedAssayDictionary.update({"Molecular Weight": CHEMEXAssayTemplateDatapoint(PostgresTableData['molecular_weight'], PostgresTableData['molecularweightunit']), "Sulfur": CHEMEXAssayTemplateDatapoint(PostgresTableData['sulfur_data'], PostgresTableData['sulfurdataunit'])})
        CHEMEXBrandedAssayDictionary.update({"Mercaptan Sulfur": CHEMEXAssayTemplateDatapoint(PostgresTableData['mercaptan_data'], PostgresTableData['mercaptandataunit']), "Nitrogen": CHEMEXAssayTemplateDatapoint(PostgresTableData['nitrogen_data'], PostgresTableData['nitrogendataunit'])})
        CHEMEXBrandedAssayDictionary.update({"Basic Nitrogen": CHEMEXAssayTemplateDatapoint(PostgresTableData['basic_nitrogen'], PostgresTableData['basicnitrogenunit']), "TotalAcidNumber TAN": CHEMEXAssayTemplateDatapoint(PostgresTableData['tan_data'], PostgresTableData['tandataunit'])})
        CHEMEXBrandedAssayDictionary.update({"Paraffins": CHEMEXAssayTemplateDatapoint(PostgresTableData['paraffins_data'], PostgresTableData['paraffinsdataunit']), "Naphthenes": CHEMEXAssayTemplateDatapoint(PostgresTableData['naphthenes_data'], PostgresTableData['naphthenesdataunit'])})
        CHEMEXBrandedAssayDictionary.update({"Aromatics": CHEMEXAssayTemplateDatapoint(PostgresTableData['aromatics_data'], PostgresTableData['aromaticsdataunit']), "Asphaltenes": CHEMEXAssayTemplateDatapoint(PostgresTableData['asphaltenes_data'], PostgresTableData['asphaltenesdataunit'])})
        CHEMEXBrandedAssayDictionary.update({"Research Octane Number": CHEMEXAssayTemplateDatapoint(PostgresTableData['ron_data'], None), "Smoke Point": CHEMEXAssayTemplateDatapoint(PostgresTableData['smoke_point'], PostgresTableData['smokepointunit'])})
        CHEMEXBrandedAssayDictionary.update({"Pour Point": CHEMEXAssayTemplateDatapoint(PostgresTableData['pour_point'], PostgresTableData['pourpointunit']), "Cloud Point": CHEMEXAssayTemplateDatapoint(PostgresTableData['cloud_point'], PostgresTableData['cloudpointunit'])})
        CHEMEXBrandedAssayDictionary.update({"Freeze Point": CHEMEXAssayTemplateDatapoint(PostgresTableData['freeze_point'], PostgresTableData['freezepointunit']), "Flash Point": CHEMEXAssayTemplateDatapoint(PostgresTableData['flash_point'], PostgresTableData['flashpointunit'])})
        CHEMEXBrandedAssayDictionary.update({"Cetane Index": CHEMEXAssayTemplateDatapoint(PostgresTableData['cetane_number'], None), "Refractive Index": CHEMEXAssayTemplateDatapoint(PostgresTableData['refractive_index'], None)})
        CHEMEXBrandedAssayDictionary.update({"Wax": CHEMEXAssayTemplateDatapoint(PostgresTableData['wax_percent'], None), "Aniline Point": CHEMEXAssayTemplateDatapoint(PostgresTableData['aniline_point'], PostgresTableData['anilinepointunit'])})
        CHEMEXBrandedAssayDictionary.update({"Conradson Carbon": CHEMEXAssayTemplateDatapoint(PostgresTableData['conradson_carbon'], PostgresTableData['conradsoncarbonunit']), "Gross Heating Value": CHEMEXAssayTemplateDatapoint(PostgresTableData['gross_heating_value'], PostgresTableData['grossheatingvalueunit'])})
        CHEMEXBrandedAssayDictionary.update({"Net Heating Value": CHEMEXAssayTemplateDatapoint(PostgresTableData['netheatingvalue'], PostgresTableData['netheatingvalueunit']), "Watson K": CHEMEXAssayTemplateDatapoint(PostgresTableData['watson_k_factor'], None)}) #Note netheatingvalue needs to be changed
        CHEMEXBrandedAssayDictionary.update({"CtoHRatio": CHEMEXAssayTemplateDatapoint(PostgresTableData['c_to_h_ratio'], None), "Mole %": CHEMEXAssayTemplateDatapoint(PostgresTableData['mole_percent'], None)})
        CHEMEXBrandedAssayDictionary.update({"MON": CHEMEXAssayTemplateDatapoint(PostgresTableData['mon_data'], None), "Reid Vapor Pressure RVP": CHEMEXAssayTemplateDatapoint(PostgresTableData['rvp_data'], PostgresTableData['rvpdataunit'])})
        CHEMEXBrandedAssayDictionary.update({"TVP": CHEMEXAssayTemplateDatapoint(PostgresTableData['tvp_data'], PostgresTableData['tvpdataunit']), "Luminometer Number (%)": CHEMEXAssayTemplateDatapoint(PostgresTableData['luminometer_number_percent'], None)})
        CHEMEXBrandedAssayDictionary.update({"Centroid Boiling Temperature": CHEMEXAssayTemplateDatapoint(PostgresTableData['centroid_boiling_temperature'], PostgresTableData['centroidboilingtemperatureunit']), "Vanadium": CHEMEXAssayTemplateDatapoint(PostgresTableData['vanadium_data'], PostgresTableData['vanadiumdataunit'])})
        CHEMEXBrandedAssayDictionary.update({"Nickel": CHEMEXAssayTemplateDatapoint(PostgresTableData['nickel_data'], PostgresTableData['nickeldataunit']), "Iron": CHEMEXAssayTemplateDatapoint(PostgresTableData['iron_data'], PostgresTableData['irondataunit'])})
        CHEMEXBrandedAssayDictionary.update({"Hydrogen": CHEMEXAssayTemplateDatapoint(PostgresTableData['hydrogen_percent'], None), "NaCl/Salt": CHEMEXAssayTemplateDatapoint(PostgresTableData['salt_data'], PostgresTableData['saltdataunit'])})
        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

        for i in range(TopLeftIndexingRow,MaxRows):
            for PropertyName in CHEMEXBrandedAssayDictionary.keys():
                if PropertyName == to_str(CHEMEXAssayTemplate.cell(row = i, column = TopLeftIndexingColumn).value) and CHEMEXBrandedAssayDictionary[PropertyName].Data != None:
                    if CHEMEXBrandedAssayDictionary[PropertyName].SpreadsheetRowIndex == None:
                        CHEMEXBrandedAssayDictionary[PropertyName].SpreadsheetRowIndex = i
                    else:
                        continue
                    for j in range(TopLeftIndexingColumn+1, TopLeftIndexingColumn + len(CHEMEXBrandedAssayDictionary[PropertyName].Data) + 2): # we will iterate from the units column to whenever our data stops
                        
                        if j == TopLeftIndexingColumn+1:
                            CHEMEXAssayTemplate.cell(row = i, column = j).value = CHEMEXBrandedAssayDictionary[PropertyName].Unit
                        else:
                            CHEMEXAssayTemplate.cell(row = i, column = j).value = CHEMEXBrandedAssayDictionary[PropertyName].Data[j-(TopLeftIndexingColumn+2)]
                    break
        
        #pe.save_as(sheet_stream = CHEMEXAssayTemplate, dest_file_name = os.getcwd() + "// ChemexBrandedTemplateCOMPLETED.xlsx")

        #Now we have to take some steps to clean up the data so it looks presentable on the CHEMEX Assay.
        #First we will get the API/density issue fixed
        if PostgresTableData['densitydataunit'] == 'API':
            CHEMEXBrandedAssayDictionary['Density'].Unit = 'SG'
            CHEMEXAssayTemplate.cell(row = CHEMEXBrandedAssayDictionary['Density'].SpreadsheetRowIndex, column = TopLeftIndexingColumn+1).value = CHEMEXBrandedAssayDictionary['Density'].Unit
            if CHEMEXBrandedAssayDictionary["API Gravity"].Data != None:
                CHEMEXBrandedAssayDictionary["Density"].Data = []
                for i in range(0,len(CHEMEXBrandedAssayDictionary["API Gravity"].Data)):
                    CHEMEXBrandedAssayDictionary["Density"].Data.append(Assay.api_to_sg(CHEMEXBrandedAssayDictionary["API Gravity"].Data[i]))
            for i in range(0,len(CHEMEXBrandedAssayDictionary["Density"].Data)):
                CHEMEXAssayTemplate.cell(row = CHEMEXBrandedAssayDictionary['Density'].SpreadsheetRowIndex, column = TopLeftIndexingColumn+2+i).value = CHEMEXBrandedAssayDictionary['Density'].Data[i]
        
        DensityMultiplier = None #The density multiplier is used to turn a density with a specific unit, into as SG, so that we can quickly convert it to API and properly fill out the CHEMEX template
        if PostgresTableData['densitydataunit'] == 'SG':
            DensityMultiplier = 1
        elif PostgresTableData['densitydataunit'] == 'kg/L':
            DensityMultiplier = 1
        elif PostgresTableData['densitydataunit'] == 'kg/m^3':
            DensityMultiplier = 1/KilogramsPerCubicMeterWater
        elif PostgresTableData['densitydataunit'] == 'lb/gal':
            DensityMultiplier = 1/PoundsPerGallonWater
        elif PostgresTableData['densitydataunit'] == 'lb/bbl':
            DensityMultiplier = 1/(PoundsPerGallonWater * GallonsPerBarrel)
        elif PostgresTableData['densitydataunit'] == 'lb/ft^3':
            DensityMultiplier = 1/(PoundsPerCubicFootWater)
    
        if DensityMultiplier != None: #Now we will take the case that we are given a different density and now we have to get the API data and report it on the asssay
            CHEMEXBrandedAssayDictionary['API Gravity'].Unit = 'API'
            CHEMEXAssayTemplate.cell(row = CHEMEXBrandedAssayDictionary['API Gravity'].SpreadsheetRowIndex, column = TopLeftIndexingColumn+1).value = CHEMEXBrandedAssayDictionary['API Gravity'].Unit
            if CHEMEXBrandedAssayDictionary["Density"].Data != None:
                CHEMEXBrandedAssayDictionary["API Gravity"].Data = []
                for i in range(0,len(CHEMEXBrandedAssayDictionary["Density"].Data)):
                    if CHEMEXBrandedAssayDictionary["Density"].Data[i] != None:
                        CHEMEXBrandedAssayDictionary["Density"].Data[i] = CHEMEXBrandedAssayDictionary["Density"].Data[i] * DensityMultiplier
                    CHEMEXBrandedAssayDictionary["API Gravity"].Data.append(Assay.sg_to_api(CHEMEXBrandedAssayDictionary["Density"].Data[i]))
            for i in range(0,len(CHEMEXBrandedAssayDictionary["API Gravity"].Data)):
                CHEMEXAssayTemplate.cell(row = CHEMEXBrandedAssayDictionary['API Gravity'].SpreadsheetRowIndex, column = TopLeftIndexingColumn+2+i).value = CHEMEXBrandedAssayDictionary['API Gravity'].Data[i]

        #Next we will fix the start and end temperatures. Unfortunately, they have the exact same key "They say the same thing on the excel sheet". so unlike APi and density, we will nto be able to 
        # differentiate the two indeces. We will have to rely an raw indexing , and understand that the Celsius temeprature index will be 2 above the fahrenheit index, and vice versa. So just know
        # when you see the +2, it is simply because of the template. If the template is changed, then we have to change this are of the code too.
        StartTemperatureFahrenheitRowIndex = None
        StartTemperatureCelsiusRowIndex = None
        EndTemperatureFahrenheitRowIndex = None
        EndTemperatureCelsiusRowIndex = None

        if PostgresTableData['finaltemperatureunit'] == 'Celsius':
            StartTemperatureCelsiusRowIndex = CHEMEXBrandedAssayDictionary["Start Temperature"].SpreadsheetRowIndex
            StartTemperatureFahrenheitRowIndex = CHEMEXBrandedAssayDictionary["Start Temperature"].SpreadsheetRowIndex + 2
            EndTemperatureCelsiusRowIndex = CHEMEXBrandedAssayDictionary["End Temperature"].SpreadsheetRowIndex
            EndTemperatureFahrenheitRowIndex = CHEMEXBrandedAssayDictionary["End Temperature"].SpreadsheetRowIndex + 2
            if CHEMEXBrandedAssayDictionary["Start Temperature"].Data != None: #Start Temperature and End temperature will have the same dimensions
                CHEMEXAssayTemplate.cell(row = StartTemperatureCelsiusRowIndex, column = TopLeftIndexingColumn+1).value = CHEMEXBrandedAssayDictionary['Start Temperature'].Unit
                CHEMEXAssayTemplate.cell(row = EndTemperatureCelsiusRowIndex, column = TopLeftIndexingColumn+1).value = CHEMEXBrandedAssayDictionary['End Temperature'].Unit
                CHEMEXAssayTemplate.cell(row = StartTemperatureFahrenheitRowIndex, column = TopLeftIndexingColumn+1).value = 'Fahrenheit'
                CHEMEXAssayTemplate.cell(row = EndTemperatureFahrenheitRowIndex, column = TopLeftIndexingColumn+1).value = 'Fahrenheit'
                for i in range(0,len(CHEMEXBrandedAssayDictionary["Start Temperature"].Data)):
                    CHEMEXAssayTemplate.cell(row = StartTemperatureCelsiusRowIndex, column = TopLeftIndexingColumn+2+i).value = CHEMEXBrandedAssayDictionary['Start Temperature'].Data[i]
                    CHEMEXAssayTemplate.cell(row = StartTemperatureFahrenheitRowIndex, column = TopLeftIndexingColumn+2+i).value = Assay.celsius_to_fahrenheit(CHEMEXBrandedAssayDictionary['Start Temperature'].Data[i])
                    CHEMEXAssayTemplate.cell(row = EndTemperatureCelsiusRowIndex, column = TopLeftIndexingColumn+2+i).value = CHEMEXBrandedAssayDictionary['End Temperature'].Data[i]
                    CHEMEXAssayTemplate.cell(row = EndTemperatureFahrenheitRowIndex, column = TopLeftIndexingColumn+2+i).value = Assay.celsius_to_fahrenheit(CHEMEXBrandedAssayDictionary['End Temperature'].Data[i])

        elif PostgresTableData['finaltemperatureunit'] == 'Fahrenheit':
            StartTemperatureCelsiusRowIndex = CHEMEXBrandedAssayDictionary["Start Temperature"].SpreadsheetRowIndex + 2 
            StartTemperatureFahrenheitRowIndex = CHEMEXBrandedAssayDictionary["Start Temperature"].SpreadsheetRowIndex
            EndTemperatureCelsiusRowIndex = CHEMEXBrandedAssayDictionary["End Temperature"].SpreadsheetRowIndex +2
            EndTemperatureFahrenheitRowIndex = CHEMEXBrandedAssayDictionary["End Temperature"].SpreadsheetRowIndex
            if CHEMEXBrandedAssayDictionary["Start Temperature"].Data != None: #Start Temperature and End temperature will have the same dimensions
                CHEMEXAssayTemplate.cell(row = StartTemperatureCelsiusRowIndex, column = TopLeftIndexingColumn+1).value = 'Celsius'
                CHEMEXAssayTemplate.cell(row = EndTemperatureCelsiusRowIndex, column = TopLeftIndexingColumn+1).value = 'Celsius'
                CHEMEXAssayTemplate.cell(row = StartTemperatureFahrenheitRowIndex, column = TopLeftIndexingColumn+1).value = CHEMEXBrandedAssayDictionary['Start Temperature'].Unit
                CHEMEXAssayTemplate.cell(row = EndTemperatureFahrenheitRowIndex, column = TopLeftIndexingColumn+1).value = CHEMEXBrandedAssayDictionary['End Temperature'].Unit
                for i in range(0,len(CHEMEXBrandedAssayDictionary["Start Temperature"].Data)):
                    CHEMEXAssayTemplate.cell(row = StartTemperatureFahrenheitRowIndex, column = TopLeftIndexingColumn+2+i).value = CHEMEXBrandedAssayDictionary['Start Temperature'].Data[i]
                    CHEMEXAssayTemplate.cell(row = StartTemperatureCelsiusRowIndex, column = TopLeftIndexingColumn+2+i).value = Assay.fahrenheit_to_celsius(CHEMEXBrandedAssayDictionary['Start Temperature'].Data[i])
                    CHEMEXAssayTemplate.cell(row = EndTemperatureFahrenheitRowIndex, column = TopLeftIndexingColumn+2+i).value = CHEMEXBrandedAssayDictionary['End Temperature'].Data[i]
                    CHEMEXAssayTemplate.cell(row = EndTemperatureCelsiusRowIndex, column = TopLeftIndexingColumn+2+i).value = Assay.fahrenheit_to_celsius(CHEMEXBrandedAssayDictionary['End Temperature'].Data[i])

        #Clean up the data so it looks how we want it to look
        for i in range (TopLeftIndexingRow,MaxRows):
            if CHEMEXAssayTemplate.cell(row = i, column = TopLeftIndexingColumn + 1).value == 'zzzz':
                CHEMEXAssayTemplate.cell(row = i, column = TopLeftIndexingColumn + 1).value = ''
            if CHEMEXAssayTemplate.cell(row = i, column = TopLeftIndexingColumn + 1).value == 'Fahrenheit':
                CHEMEXAssayTemplate.cell(row = i, column = TopLeftIndexingColumn + 1).value = '°F'
            if CHEMEXAssayTemplate.cell(row = i, column = TopLeftIndexingColumn + 1).value == 'Celsius':
                CHEMEXAssayTemplate.cell(row = i, column = TopLeftIndexingColumn + 1).value = '°C'

        CHEMEXAssayTemplateWorkbook.save(os.getcwd() + "\CHEMEXBrandedAssayTemplateFINSHED.xlsx")

        return

    #Client Information on Dashboard
    def DomesticSelected(self):
        if len(self.Refinery_List.values()) == 0:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please make sure a Refinery is present', QtWidgets.QMessageBox.Ok)
            return
        self.dlg.btn_fixed.setChecked(True)
        self.dlg.txt_DefinitionStudyCost.setText('75000')
        self.dlg.txt_FeasibilityStudyCost.setText('212750')
        self.dlg.txt_ProcessDesignStudyCost.setText('125000')
        self.dlg.spin_commission.setValue(15)
        self.FixedSelected()
        self.IsDomestic = True
        self.IsInternational = False
        self.CreateDictionary()
        self.dlg.txt_ProductsBreakdownCost.setText("$ " + special_format(self.Refinery_List[self.SelectedRefinery].DomesticTotalPrice))
        return
    def InternationalSelected(self):
        if len(self.Refinery_List.values()) == 0:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please make sure a Refinery is present', QtWidgets.QMessageBox.Ok)
            return
        self.dlg.btn_budgetary.setChecked(True)
        self.dlg.txt_budgetary.setText('40')
        self.dlg.txt_DefinitionStudyCost.setText('125000')
        self.dlg.txt_FeasibilityStudyCost.setText('250000')
        self.dlg.txt_ProcessDesignStudyCost.setText('125000')
        self.BudgetarySelected()
        self.dlg.spin_commission.setValue(15)
        self.IsInternational = True
        self.IsDomestic = False
        self.CreateDictionary()
        self.dlg.txt_ProductsBreakdownCost.setText("$ " + special_format(self.Refinery_List[self.SelectedRefinery].OverseasTotalPrice))
        return
    def File_Select(self):
        self.DirectoryFolderName = str(QFileDialog.getExistingDirectory(self.dlg, "Select Directory"))
        self.dlg.txt_folder.setText(self.DirectoryFolderName)
        return

    #Pricing on Dashboard
    def BudgetarySelected(self):
        self.dlg.txt_budgetary.show()
        self.Dictionary.update({'Budget_Fixed': 'budgetary', '+or-': '± %' + dlg.txt_budgetary.text(), "Pricing": 'Budgetary'})
    def FixedSelected(self):
        self.dlg.txt_budgetary.hide()
        self.Dictionary.update({'Budget_Fixed': 'fixed', '+or-': '', "Pricing": 'Fixed'})

    # Proposal Information on Dashboard
    def OtherContact(self):
        if self.dlg.cmb_contact.currentText() == 'Other':
                self.dlg.lbl_name.show()
                self.dlg.txt_name.show()
                self.dlg.lbl_title.show()
                self.dlg.txt_title.show()
                self.dlg.lbl_email.show()
                self.dlg.txt_email.show()
                self.dlg.lbl_phone.show()
                self.dlg.txt_phone.show()
        else: 
            self.dlg.lbl_name.hide()
            self.dlg.txt_name.hide()
            self.dlg.lbl_title.hide()
            self.dlg.txt_title.hide()
            self.dlg.lbl_email.hide()
            self.dlg.txt_email.hide()
            self.dlg.lbl_phone.hide()
            self.dlg.txt_phone.hide()

    # Unit information on Dashboard
    def ShowRegionOptions(self):
        try:
            if self.IsRegionSearch:
                self.MyDatabase.ConnectToDatabase()
                self.MyDatabase.cursor.execute("SELECT assayregion FROM " + self.MyDatabase.tableName + ' ORDER BY assayregion ASC;')
                Assay_Region_List = self.MyDatabase.cursor.fetchall()
                self.dlg.cmb_AssayRegion.clear()
                self.dlg.cmb_AssayRegion.addItem('Select Region')
                try:
                    self.dlg.cmb_AssayRegion.addItem(Assay_Region_List[0][0])
                except:
                    print("No Data Exists in the database")
                for i in range (0, len(Assay_Region_List)):
                    if i >= 1 and Assay_Region_List[i][0] != Assay_Region_List[i-1][0]:
                        self.dlg.cmb_AssayRegion.addItem(Assay_Region_List[i][0])
                self.MyDatabase.DisconnectFromDatabase()
                return
            if self.IsAssaySearch:
                self.MyDatabase.ConnectToDatabase()
                self.dlg.cmb_AssayRegion.clear()
                self.MyDatabase.cursor.execute("SELECT assayregion FROM " + self.MyDatabase.tableName + " WHERE assayname = '" + self.dlg.cmb_AssayName.currentText() +"' ORDER BY assayregion ASC;")
                Assay_Region_List = self.MyDatabase.cursor.fetchall()
                self.dlg.cmb_AssayRegion.addItem('Select Region')
                try:
                    self.dlg.cmb_AssayRegion.addItem(Assay_Region_List[0][0])
                except:
                    print("No Data Exists in the database")
                for i in range (0, len(Assay_Region_List)):
                    if i >= 1 and Assay_Region_List[i][0] != Assay_Region_List[i-1][0]:
                        self.dlg.cmb_AssayRegion.addItem(Assay_Region_List[i][0])
                self.dlg.cmb_AssayRegion.setCurrentIndex(1) #There should only be Select Region and just one region name
                self.ShowSourceOptions()
                self.MyDatabase.DisconnectFromDatabase()
                return
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Adding Assay. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def ShowAssayOptions(self):
        # displays the different source options for the assay selected
        try:
            if self.dlg.cmb_AssayRegion.currentText() == "Select Region" and self.dlg.cmb_AssayName.currentText() != "Select Assay Name": #If the user selects Region options to reset
                self.dlg.cmb_AssayName.setCurrentIndex(0) 
                self.ShowSourceOptions()
                return

            if self.dlg.cmb_AssayRegion.currentText() == "Select Region":
                self.IsRegionSearch = True
                self.IsAssaySearch = False
                self.dlg.cmb_AssayName.clear()
                self.dlg.cmb_AssaySource.clear()
                self.dlg.cmb_AssayName.addItem('Select Assay Name')
                self.MyDatabase.ConnectToDatabase()
                self.MyDatabase.cursor.execute("SELECT assayname FROM " + self.MyDatabase.tableName + ' ORDER BY assayname ASC;')
                Assay_Name_List = self.MyDatabase.cursor.fetchall()
                try:
                    self.dlg.cmb_AssayName.addItem(Assay_Name_List[0][0])
                except:
                    print("No Data Exists in the database")
                for i in range (0, len(Assay_Name_List)):
                    if i >= 1 and Assay_Name_List[i][0] != Assay_Name_List[i-1][0]:
                        self.dlg.cmb_AssayName.addItem(Assay_Name_List[i][0])
                self.MyDatabase.DisconnectFromDatabase()
                return
        
            self.MyDatabase.ConnectToDatabase()

            if self.IsRegionSearch:
                self.dlg.cmb_AssayName.clear()
                self.dlg.cmb_AssaySource.clear()
                self.dlg.cmb_AssayName.addItem('Select Assay Name')
                self.MyDatabase.cursor.execute("SELECT assayname FROM " + self.MyDatabase.tableName + " WHERE assayregion = '" + self.dlg.cmb_AssayRegion.currentText() + "' ORDER BY assayname ASC;" )
                Assay_Name_List = self.MyDatabase.cursor.fetchall()
                try:
                    self.dlg.cmb_AssayName.addItem(Assay_Name_List[0][0])
                except:
                    print("No Data Exists in the database")
                for i in range (0, len(Assay_Name_List)):
                    if i >= 1 and Assay_Name_List[i][0] != Assay_Name_List[i-1][0]:
                        self.dlg.cmb_AssayName.addItem(Assay_Name_List[i][0])
                self.MyDatabase.DisconnectFromDatabase()
                #self.ShowSourceOptions()
                return
        

            if self.IsAssaySearch:
                self.MyDatabase.cursor.execute("SELECT assaysource FROM " + self.MyDatabase.tableName + " WHERE assayregion = '" + self.dlg.cmb_AssayRegion.currentText() + "' AND assayname = '" +  self.dlg.cmb_AssayName.currentText() + "' ORDER BY assaysource ASC;")
                Assay_Source_List = self.MyDatabase.cursor.fetchall()
                try:
                    self.dlg.cmb_AssaySource.addItem(str(Assay_Source_List[0][0]))
                except:
                    print("No Data Exists in the database")
                for i in range (0, len(Assay_Source_List)):
                    if i >= 1 and Assay_Source_List[i][0] != Assay_Source_List[i-1][0]:
                        self.dlg.cmb_AssaySource.addItem(str(Assay_Source_List[i][0]))
                self.MyDatabase.DisconnectFromDatabase()
                return
            self.MyDatabase.DisconnectFromDatabase()
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Adding Assay. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def ShowSourceOptions(self):
        try:
            if self.dlg.cmb_AssayName.currentText() == "Select Assay Name": 
                self.dlg.cmb_AssaySource.clear()
                self.IsAssaySearch = True
                self.IsRegionSearch = True
                self.MyDatabase.ConnectToDatabase()
                self.MyDatabase.cursor.execute("SELECT assayname FROM " + self.MyDatabase.tableName + ' ORDER BY assayname ASC;')
                Assay_Name_List = self.MyDatabase.cursor.fetchall()
                try:
                    self.dlg.cmb_AssayName.addItem(Assay_Name_List[0][0])
                except:
                    print("No Data Exists in the database")
                for i in range (0, len(Assay_Name_List)):
                    if i >= 1 and Assay_Name_List[i][0] != Assay_Name_List[i-1][0]:
                        self.dlg.cmb_AssayName.addItem(Assay_Name_List[i][0])
                self.MyDatabase.DisconnectFromDatabase()
                self.ShowRegionOptions()
                return
            if self.dlg.cmb_AssayRegion.currentText() == "Select Region":
                self.IsAssaySearch = True
                self.IsRegionSearch = False
                self.ShowRegionOptions()
                return

            # displays the different dates for the assay and source selected
            self.dlg.cmb_AssaySource.clear()
            self.dlg.cmb_AssaySource.addItem('Select Source') 
            self.MyDatabase.ConnectToDatabase()
            self.MyDatabase.cursor.execute("SELECT assaysource FROM " + self.MyDatabase.tableName + " WHERE assayregion = '" + self.dlg.cmb_AssayRegion.currentText() + "' AND assayname = '" +  self.dlg.cmb_AssayName.currentText() + "' ORDER BY assaysource ASC;")
            Assay_Source_List = self.MyDatabase.cursor.fetchall()
            try:
                self.dlg.cmb_AssaySource.addItem(str(Assay_Source_List[0][0]))
            except:
                print("No Data Exists in the database")
            for i in range (0, len(Assay_Source_List)):
                if i >= 1 and Assay_Source_List[i][0] != Assay_Source_List[i-1][0]:
                    self.dlg.cmb_AssaySource.addItem(str(Assay_Source_List[i][0]))
            self.MyDatabase.DisconnectFromDatabase()

            if self.dlg.cmb_AssaySource.count() == 1: #This means the query did not find any sources because the region and Assay didn't match
                if self.IsAssaySearch:
                    self.ShowRegionOptions()
            
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Adding Assay. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    ###Joe: Begin creating Price Index Date Dropdown and Fill-in Price Index Values
    def ShowPriceIndexOptions(self):
        if self.IsPriceIndexDateSearch:
            self.MyPriceIndexDatabase.ConnectToDatabase()
            self.MyPriceIndexDatabase.cursor.execute("SELECT date FROM " + self.MyPriceIndexDatabase.tableName + ' ORDER BY date DESC;')
            Price_Index_Date_List = self.MyPriceIndexDatabase.cursor.fetchall()
            #self.dlg.cmb_IndexDate.clear()
            try:
                self.dlg.cmb_IndexDate.addItem(str(Price_Index_Date_List[0][0]))
            except:
                print("No date Data Exists in the database")
            for i in range (0, len(Price_Index_Date_List)):
                if i >= 1 and Price_Index_Date_List[i][0] != Price_Index_Date_List[i-1][0]:
                    self.dlg.cmb_IndexDate.addItem(str(Price_Index_Date_List[i][0]))
            self.dlg.txt_NelsonFarrarInflationIndex.clear()
            self.dlg.txt_LaborComponent.clear()
            self.dlg.txt_MaterialsComponent.clear()
            self.dlg.txt_MiscEquipment.clear()
            self.dlg.txt_ENRConstructionCostIndex.clear()
            self.dlg.txt_ENRBuildingCostIndex.clear()
            self.dlg.txt_WPU101_Index.clear()
            self.dlg.txt_WPU13_Index.clear()
            self.dlg.txt_WPU1141_Index.clear()
            self.dlg.txt_WPU117_Index.clear()
            self.dlg.txt_WPU1194_Index.clear()
            self.dlg.txt_CUSR0_Index.clear()
            self.dlg.txt_PCU332_Index.clear()
            ###
            self.MyPriceIndexDatabase.cursor.execute("SELECT ENR_Construction_Cost_Index FROM " + self.MyPriceIndexDatabase.tableName + " WHERE date = '" + self.dlg.cmb_IndexDate.currentText() + "';")
            ENR_Construction_Cost_Index_List = self.MyPriceIndexDatabase.cursor.fetchall()
            try:
                self.dlg.txt_ENRConstructionCostIndex.setText(str(ENR_Construction_Cost_Index_List[0][0]))
            except:
                print("No ENRConstructionCostIndex Data Exists in the database")
            for i in range (0, len(ENR_Construction_Cost_Index_List)):
                if i >= 1 and ENR_Construction_Cost_Index_List[i][0] != ENR_Construction_Cost_Index_List[i-1][0]:
                    self.dlg.txt_ENRConstructionCostIndex.setText(str(ENR_Construction_Cost_Index_List[i][0]))
            self.ENR_Construction_Cost_Index_Value = self.dlg.txt_ENRConstructionCostIndex.text()
            ###
            self.MyPriceIndexDatabase.cursor.execute("SELECT ENR_Building_Cost_Index FROM " + self.MyPriceIndexDatabase.tableName + " WHERE date = '" + self.dlg.cmb_IndexDate.currentText() + "';")
            ENR_Building_Cost_Index_List = self.MyPriceIndexDatabase.cursor.fetchall()
            try:
                self.dlg.txt_ENRBuildingCostIndex.setText(str(ENR_Building_Cost_Index_List[0][0]))
            except:
                print("No ENRBuildingCostIndex Data Exists in the database")
            for i in range (0, len(ENR_Building_Cost_Index_List)):
                if i >= 1 and ENR_Building_Cost_Index_List[i][0] != ENR_Building_Cost_Index_List[i-1][0]:
                    self.dlg.txt_ENRBuildingCostIndex.setText(str(ENR_Building_Cost_Index_List[i][0]))
            self.ENR_Building_Cost_Index_Value = self.dlg.txt_ENRBuildingCostIndex.text()
            ###
            self.MyPriceIndexDatabase.cursor.execute("SELECT WPU101_Index FROM " + self.MyPriceIndexDatabase.tableName + " WHERE date = '" + self.dlg.cmb_IndexDate.currentText() + "';")
            WPU101_Index_List = self.MyPriceIndexDatabase.cursor.fetchall()
            try:
                self.dlg.txt_WPU101_Index.setText(str(WPU101_Index_List[0][0]))
            except:
                print("No WPU101_Index Data Exists in the database")
            for i in range (0, len(WPU101_Index_List)):
                if i >= 1 and WPU101_Index_List[i][0] != WPU101_Index_List[i-1][0]:
                    self.dlg.txt_WPU101_Index.setText(str(WPU101_Index_List[i][0]))
            self.WPU101_Index_Value = self.dlg.txt_WPU101_Index.text()
            ###
            self.MyPriceIndexDatabase.cursor.execute("SELECT WPU13_Index FROM " + self.MyPriceIndexDatabase.tableName + " WHERE date = '" + self.dlg.cmb_IndexDate.currentText() + "';")
            WPU13_Index_List = self.MyPriceIndexDatabase.cursor.fetchall()
            try:
                self.dlg.txt_WPU13_Index.setText(str(WPU13_Index_List[0][0]))
            except:
                print("No WPU13_Index Data Exists in the database")
            for i in range (0, len(WPU13_Index_List)):
                if i >= 1 and WPU13_Index_List[i][0] != WPU13_Index_List[i-1][0]:
                    self.dlg.txt_WPU13_Index.setText(str(WPU13_Index_List[i][0]))
            self.WPU13_Index_Value = self.dlg.txt_WPU13_Index.text()
            ###
            self.MyPriceIndexDatabase.cursor.execute("SELECT WPU1141_Index FROM " + self.MyPriceIndexDatabase.tableName + " WHERE date = '" + self.dlg.cmb_IndexDate.currentText() + "';")
            WPU1141_Index_List = self.MyPriceIndexDatabase.cursor.fetchall()
            try:
                self.dlg.txt_WPU1141_Index.setText(str(WPU1141_Index_List[0][0]))
            except:
                print("No WPU1141_Index Data Exists in the database")
            for i in range (0, len(WPU1141_Index_List)):
                if i >= 1 and WPU1141_Index_List[i][0] != WPU1141_Index_List[i-1][0]:
                    self.dlg.txt_WPU1141_Index.setText(str(WPU1141_Index_List[i][0]))
            self.WPU1141_Index_Value = self.dlg.txt_WPU1141_Index.text()
            ###
            self.MyPriceIndexDatabase.cursor.execute("SELECT WPU117_Index FROM " + self.MyPriceIndexDatabase.tableName + " WHERE date = '" + self.dlg.cmb_IndexDate.currentText() + "';")
            WPU117_Index_List = self.MyPriceIndexDatabase.cursor.fetchall()
            try:
                self.dlg.txt_WPU117_Index.setText(str(WPU117_Index_List[0][0]))
            except:
                print("No WPU117_Index Data Exists in the database")
            for i in range (0, len(WPU117_Index_List)):
                if i >= 1 and WPU117_Index_List[i][0] != WPU117_Index_List[i-1][0]:
                    self.dlg.txt_WPU117_Index.setText(str(WPU117_Index_List[i][0]))
            self.WPU117_Index_Value = self.dlg.txt_WPU117_Index.text()
            ###
            self.MyPriceIndexDatabase.cursor.execute("SELECT WPU1194_Index FROM " + self.MyPriceIndexDatabase.tableName + " WHERE date = '" + self.dlg.cmb_IndexDate.currentText() + "';")
            WPU1194_Index_List = self.MyPriceIndexDatabase.cursor.fetchall()
            try:
                self.dlg.txt_WPU1194_Index.setText(str(WPU1194_Index_List[0][0]))
            except:
                print("No WPU1194_Index Data Exists in the database")
            for i in range (0, len(WPU1194_Index_List)):
                if i >= 1 and WPU1194_Index_List[i][0] != WPU1194_Index_List[i-1][0]:
                    self.dlg.txt_WPU1194_Index.setText(str(WPU1194_Index_List[i][0]))
            self.WPU1194_Index_Value = self.dlg.txt_WPU1194_Index.text()
            ###
            self.MyPriceIndexDatabase.cursor.execute("SELECT CUSR0000SEEE_Index FROM " + self.MyPriceIndexDatabase.tableName + " WHERE date = '" + self.dlg.cmb_IndexDate.currentText() + "';")
            CUSR0000SEEE_Index_List = self.MyPriceIndexDatabase.cursor.fetchall()
            try:
                self.dlg.txt_CUSR0_Index.setText(str(CUSR0000SEEE_Index_List[0][0]))
            except:
                print("No CUSR0000SEEE_Index Data Exists in the database")
            for i in range (0, len(CUSR0000SEEE_Index_List)):
                if i >= 1 and CUSR0000SEEE_Index_List[i][0] != CUSR0000SEEE_Index_List[i-1][0]:
                    self.dlg.txt_CUSR0_Index.setText(str(CUSR0000SEEE_Index_List[i][0]))
            self.CUSR0000SEEE_Index_Value = self.dlg.txt_CUSR0_Index.text()
            ###
            self.MyPriceIndexDatabase.cursor.execute("SELECT PCU332410332410_Index FROM " + self.MyPriceIndexDatabase.tableName + " WHERE date = '" + self.dlg.cmb_IndexDate.currentText() + "';")
            PCU332410332410_Index_List = self.MyPriceIndexDatabase.cursor.fetchall()
            try:
                self.dlg.txt_PCU332_Index.setText(str(PCU332410332410_Index_List[0][0]))
            except:
                print("No PCU332410332410_Index Data Exists in the database")
            for i in range (0, len(PCU332410332410_Index_List)):
                if i >= 1 and PCU332410332410_Index_List[i][0] != PCU332410332410_Index_List[i-1][0]:
                    self.dlg.txt_PCU332_Index.setText(str(PCU332410332410_Index_List[i][0]))
            self.PCU332410332410_Index_Value = self.dlg.txt_PCU332_Index.text()
            ###
            self.MyPriceIndexDatabase.DisconnectFromDatabase()
            if self.ENR_Construction_Cost_Index_Value == "None" or self.ENR_Building_Cost_Index_Value == "None" or self.WPU101_Index_Value == "None" or self.WPU13_Index_Value == "None" or self.WPU1141_Index_Value == "None" or self.WPU117_Index_Value == "None" or self.WPU1194_Index_Value == "None" or self.CUSR0000SEEE_Index_Value == "None" or self.PCU332410332410_Index_Value == "None":
                pass
            else:
                LaborComponent = ((((float(self.ENR_Construction_Cost_Index_Value))/1.033)*100)*0.35)+((((float(self.ENR_Building_Cost_Index_Value))/1.8)*100)*0.65)
                self.dlg.txt_LaborComponent.setText(str(LaborComponent))
                MiscEquipment = ((float(self.WPU1141_Index_Value))+(float(self.WPU117_Index_Value))+(float(self.WPU1194_Index_Value))+(float(self.CUSR0000SEEE_Index_Value)*212)+(float(self.PCU332410332410_Index_Value)*9))/5
                self.dlg.txt_MiscEquipment.setText(str(MiscEquipment))
                MaterialsComponent = (float(self.WPU101_Index_Value)*0.5)+(float(self.WPU13_Index_Value)*0.2)+(float(MiscEquipment)*0.3)
                self.dlg.txt_MaterialsComponent.setText(str(MaterialsComponent))
                NelsonFarrarInflationIndex = (float(LaborComponent)*0.6)+(float(MaterialsComponent)*0.4)
                self.dlg.txt_NelsonFarrarInflationIndex.setText(str(NelsonFarrarInflationIndex))
            ###Joe: End creating Price Index Date Dropdown and Fill-in Price Index Values
        return
    def AddAssay(self):

        try:

            if self.dlg.cmb_AssayRegion.currentText() == 'Select Region':
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please select assay region.', QtWidgets.QMessageBox.Ok)
                return
            if self.dlg.cmb_AssayName.currentText() == 'Select Assay Name': #if assay hasnt been picked we do not add anything
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please select assay name.', QtWidgets.QMessageBox.Ok)
                return
            if self.dlg.cmb_AssaySource.currentText() == 'Select Source':
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please select assay source.', QtWidgets.QMessageBox.Ok)
                return


            assayID = '(' + self.dlg.cmb_AssayRegion.currentText() + ') ' + '(' + self.dlg.cmb_AssayName.currentText() + ') ' + '(' + self.dlg.cmb_AssaySource.currentText() + ')'
            self.SelectedRefinery = assayID
            self.Refinery_List.update({assayID: Refinery()})
            self.Refinery_List[assayID].ProductsInstance.InitializeAssay()
            self.Refinery_List[assayID].ProductsInstance.AssayInstance.AssayID = assayID
            self.Refinery_List[assayID].ProductsInstance.AssayInstance.AssayName = self.dlg.cmb_AssayName.currentText()
            self.Refinery_List[assayID].ProductsInstance.AssayInstance.AssayRegion = self.dlg.cmb_AssayRegion.currentText()
            self.Refinery_List[assayID].ProductsInstance.AssayInstance.AssaySource = self.dlg.cmb_AssaySource.currentText()
            self.AssignAssayVariables()
            self.Refinery_List[assayID].ProductsInstance.InitializeDefaultCutTemps()

            #This will be at the very end of the code
            self.dlg.list_assays.addItem(assayID)
            self.dlg.txt_SelectedRefinery.setText(self.SelectedRefinery)
            #Ensures there are no duplicates
            if self.dlg.list_assays.count() >= 1:
                for i in range(1, self.dlg.list_assays.count()):
                    if dlg.list_assays.item(i).text() ==  dlg.list_assays.item(i-1).text():
                        QtWidgets.QMessageBox.information(self.dlg, 'Error', 'You have selected more than one of the same Assay', QtWidgets.QMessageBox.Ok)
                        self.dlg.list_assays.takeItem(self.dlg.list_assays.count()-1)
                        return
            self.dlg.cmb_CurrentlySelectedRefinery.addItems([assayID])
            self.dlg.cmb_SourceName.addItems([assayID])
            self.dlg.cmb_SettingsSelectedRefinery.addItems([assayID])
            self.dlg.cmb_SettingsSelectedRefinery.setCurrentIndex(self.dlg.list_assays.count())
            self.AssignDashboardVariables()
            self.FillOutDashboardVariables() #This is done so we can see the default cut temperatures on the settings page
            #This has to do with blending and is not useful at the moment
            ################################################################################
            if self.dlg.list_assays.count() > 1: 
                self.dlg.chk_blend.show()
                if self.dlg.list_assays.count() == 2: 
                    self.dlg.txt_blendPercent1.show()
                    self.dlg.lbl_blendPercent1.setText('%  ' + dlg.list_assays.item(0).text())
                    self.dlg.lbl_blendPercent2.setText('%  ' + dlg.list_assays.item(1).text())
                    self.dlg.lbl_blendPercent1.show()
                    self.dlg.txt_blendPercent2.show()
                    self.dlg.lbl_blendPercent2.show()
                elif self.dlg.list_assays.count() == 3: 
                    self.dlg.lbl_blendPercent3.setText('%  ' + self.dlg.list_assays.item(2).text())
                    self.dlg.lbl_blendPercent3.show()
                    self.dlg.txt_blendPercent3.show()
                elif self.dlg.list_assays.count() == 4: 
                    self.dlg.lbl_blendPercent4.setText('%  ' + self.dlg.list_assays.item(3).text())
                    self.dlg.lbl_blendPercent4.show()
                    self.dlg.txt_blendPercent4.show()
                elif self.dlg.list_assays.count() == 5: 
                    self.dlg.lbl_blendPercent5.setText('%  ' + self.dlg.list_assays.item(4).text())
                    self.dlg.lbl_blendPercent5.show()
                    self.dlg.txt_blendPercent5.show()
                elif self.dlg.list_assays.count() == 6: 
                    self.dlg.lbl_blendPercent6.setText('%  ' + self.dlg.list_assays.item(5).text())
                    self.dlg.lbl_blendPercent6.show()
                    self.dlg.txt_blendPercent6.show()
                elif self.dlg.list_assays.count() == 7: 
                    self.dlg.lbl_blendPercent7.setText('%  ' + dlg.list_assays.item(6).text())
                    self.dlg.lbl_blendPercent7.show()
                    self.dlg.txt_blendPercent7.show()
            else: self.dlg.chk_blend.hide()
            return
        except:
            raise
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Adding Assay. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def ClearAssays(self):
        try:
            if len(self.Refinery_List.values()) < 1:
                return
            self.dlg.list_assays.clear()
            self.dlg.cmb_UnitOperation.clear()
            self.dlg.cmb_CurrentlySelectedRefinery.clear()
            self.dlg.cmb_CurrentlySelectedRefinery.addItems(["Select Refinery"])
            self.SelectedRefinery = ''
            self.dlg.txt_SelectedRefinery.setText(self.SelectedRefinery)
            self.dlg.cmb_SourceName.clear()
            self.dlg.cmb_SourceName.addItems(['Select Source Name'])
            self.dlg.cmb_SettingsSelectedRefinery.clear()
            self.dlg.cmb_SettingsSelectedRefinery.addItems(['Select Source Name'])
            self.Refinery_List.clear()
            self.OverallRefinery.ReInitialize()
            self.ClearDashboardVariables()
            self.commissionValue = self.dlg.spin_commission.value()
        
            self.UploadImage()
            self.Proposal_Dictionary_List.clear()
            self.Dictionary.clear()
            self.BudgetarySelected()
            self.ClearChemexCalculatedProcessUnitsUtilities()
            self.ClearDesiredProcessUnitsUtilities()
            self.ClearChemexUnit()
            self.ClearBlindedSideDraws()
            self.ClearProducts()

            self.dlg.chk_blend.hide()
            self.dlg.txt_blendPercent1.hide()
            self.dlg.lbl_blendPercent1.hide()
            self.dlg.txt_blendPercent2.hide()
            self.dlg.txt_blendPercent3.hide()
            self.dlg.txt_blendPercent4.hide()
            self.dlg.txt_blendPercent5.hide()
            self.dlg.txt_blendPercent6.hide()
            self.dlg.txt_blendPercent7.hide()
            self.dlg.lbl_blendPercent2.hide()
            self.dlg.lbl_blendPercent3.hide()
            self.dlg.lbl_blendPercent4.hide()
            self.dlg.lbl_blendPercent5.hide()
            self.dlg.lbl_blendPercent6.hide()
            self.dlg.lbl_blendPercent7.hide()
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Clearing Assays. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def ClearRefinery(self):
        try:
            print(self.dlg.list_assays.selectedItems())
            if len(self.dlg.list_assays.selectedItems()) == 0: #Error Checking
                QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please select an assay to clear', QtWidgets.QMessageBox.Ok)
                return
            self.Refinery_List.pop(self.dlg.list_assays.selectedItems()[0].text(),None)
            self.dlg.cmb_CurrentlySelectedRefinery.removeItem(self.dlg.list_assays.currentRow()+1)
            self.dlg.cmb_SourceName.removeItem(self.dlg.list_assays.currentRow()+1)
            self.dlg.cmb_SettingsSelectedRefinery.removeItem(self.dlg.list_assays.currentRow()+1)
            self.dlg.list_assays.takeItem(self.dlg.list_assays.currentRow())
            if len(self.Refinery_List) == 0: #Error checking
                self.SelectedRefinery = ''
                return
            self.SelectedRefinery = self.dlg.list_assays.item(0).text()
            self.dlg.txt_SelectedRefinery.setText(self.SelectedRefinery)
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Clearing Assay. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def SelectRefinery(self):
        try:
            if len(self.dlg.list_assays.selectedItems()) == 0: #Error Checking
                QtWidgets.QMessageBox.information(self, 'Error', 'Please select an assay to replace', QtWidgets.QMessageBox.Ok)
                return
            if self.ProductsChangesCheck() == True or self.RefineryChangesCheck() == True:
                self.Refinery_List[self.SelectedRefinery].ReInitialize()
                self.OverallRefinery.ReInitialize()
            self.AssignDashboardVariables()
            self.SelectedRefinery = self.dlg.list_assays.selectedItems()[0].text()
            self.dlg.txt_SelectedRefinery.setText(self.SelectedRefinery)
            self.dlg.cmb_SettingsSelectedRefinery.setCurrentIndex(self.dlg.list_assays.currentRow()+1)
            self.FillOutDashboardVariables()
            self.UploadImage()
        
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Adding Assay. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def RefineryChangesCheck(self):
        if len(self.Refinery_List) == 0: #Checking to ensure we have a refinery
            return True
        if self.dlg.spin_commission.value() != self.Refinery_List[self.SelectedRefinery].CommissionRate:
            return True
        return
    def ProductsChangesCheck(self):

        if len(self.Refinery_List) == 0: #Error checking
            return True

        if self.dlg.chk_FuelGas.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasFuelGas:
            return True
        if self.dlg.chk_LPG.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLPG:
            return True
        if self.dlg.chk_AGO.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasAGO: 
            return True
        if self.dlg.chk_ATB.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasATB: 
            return True
        if self.dlg.chk_LightNaphtha.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLightNaphtha:
            return True
        if self.dlg.chk_HeavyNaphtha.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasHeavyNaphtha: 
            return True
        if self.dlg.chk_SRNaphtha.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasSRNaphtha: 
            return True
        if self.dlg.chk_Gasoline.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasGasoline:
            return True
        if self.dlg.chk_Kerosene.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasKerosene:
            return True
        if self.dlg.chk_JetA.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasJetA:
            return True
        if self.dlg.chk_Diesel.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasDiesel:
            return True
        if self.dlg.chk_HSDiesel.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasHSDiesel: 
            return True
        if self.dlg.chk_ULSDiesel.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasULSDiesel:
            return True
        if self.dlg.chk_MDO.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasMDO: 
            return True
        if self.dlg.chk_LSMDO.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLSMDO: 
            return True
        if self.dlg.chk_LVGO.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasLVGO: 
            return True
        if self.dlg.chk_HVGO.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasHVGO: 
            return True
        if self.dlg.chk_VTB.isChecked() != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HasVTB:
            return True

        #If the user has added blinded side draws
        if (len(self.dlg.list_BlindedSideDraws.findItems('Kerosene',QtCore.Qt.MatchExactly)) >= 1) != self.Refinery_List[self.SelectedRefinery].ProductsInstance.KeroseneBlindedSideDrawCheck: return True
        if (len(self.dlg.list_BlindedSideDraws.findItems('Jet A',QtCore.Qt.MatchExactly)) >= 1) != self.Refinery_List[self.SelectedRefinery].ProductsInstance.JetABlindedSideDrawCheck: return  True
        if (len(self.dlg.list_BlindedSideDraws.findItems('Diesel',QtCore.Qt.MatchExactly)) >= 1) != self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselBlindedSideDrawCheck: return  True
        if (len(self.dlg.list_BlindedSideDraws.findItems('ULS Diesel',QtCore.Qt.MatchExactly)) >= 1) != self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselBlindedSideDrawCheck: return  True
        if (len(self.dlg.list_BlindedSideDraws.findItems('HS Diesel',QtCore.Qt.MatchExactly)) >= 1) != self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselBlindedSideDrawCheck: return  True
        if (len(self.dlg.list_BlindedSideDraws.findItems('LSMDO',QtCore.Qt.MatchExactly)) >= 1) != self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOBlindedSideDrawCheck: return  True
        if (len(self.dlg.list_BlindedSideDraws.findItems('MDO',QtCore.Qt.MatchExactly)) >= 1) != self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOBlindedSideDrawCheck: return  True
        if (len(self.dlg.list_BlindedSideDraws.findItems('AGO',QtCore.Qt.MatchExactly)) >= 1) != self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOBlindedSideDrawCheck: return  True

        for i in range(0,len(self.Temperature_Cut_Point_List)): #We will insert error checking to ensure the input is correct
            if self.Temperature_Cut_Point_List[i].text() != '' and math.fabs(float(self.Temperature_Cut_Point_List[i].text()) - self.Refinery_List[self.SelectedRefinery].ProductsInstance.Specified_Cut_Temps[i][1]) >= 0.01:
                return True

        try:
            if int(self.dlg.txt_totalCapacity.text()) != self.Refinery_List[self.SelectedRefinery].ProductsInstance.CapacityValue:
                return True
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please Enter a valid capacity', QtWidgets.QMessageBox.Ok)
            return False


        return False

    #Desired Products on Dashboard
    #Maximization Adding Buttons
    def FuelGasMax(self):
        self.dlg.list_MaximizedProducts.addItem('Fuel Gas')
        self.ShowMaximizationHierarchy()
        return
    def LPGMax(self):
        self.dlg.list_MaximizedProducts.addItem('LPG')
        self.ShowMaximizationHierarchy()
        return
    def LightNaphthaMax(self):
        self.dlg.list_MaximizedProducts.addItem('Light Naphtha')
        self.ShowMaximizationHierarchy()
        return
    def HeavyNaphthaMax(self):
        self.dlg.list_MaximizedProducts.addItem('Heavy Naphtha')
        self.ShowMaximizationHierarchy()
        return
    def SRNaphthaMax(self):
        self.dlg.list_MaximizedProducts.addItem('SR Naphtha')
        self.ShowMaximizationHierarchy()
        return
    def GasolineMax(self):
        self.dlg.list_MaximizedProducts.addItem('Gasoline')
        self.ShowMaximizationHierarchy()
        return
    def KeroseneMax(self):
        self.dlg.list_MaximizedProducts.addItem('Kerosene')
        self.ShowMaximizationHierarchy()
        return
    def JetAMax(self):
        self.dlg.list_MaximizedProducts.addItem('Jet A')
        self.ShowMaximizationHierarchy()
        return
    def DieselMax(self):
        self.dlg.list_MaximizedProducts.addItem('Diesel')
        self.ShowMaximizationHierarchy()
        return
    def ULSDieselMax(self):
        self.dlg.list_MaximizedProducts.addItem('ULS Diesel')
        self.ShowMaximizationHierarchy()
        return
    def HSDieselMax(self):
        self.dlg.list_MaximizedProducts.addItem('HS Diesel')
        self.ShowMaximizationHierarchy()
        return
    def MDOMax(self):
        self.dlg.list_MaximizedProducts.addItem('MDO')
        self.ShowMaximizationHierarchy()
        return
    def LSMDOMax(self):
        self.dlg.list_MaximizedProducts.addItem('LS MDO')
        self.ShowMaximizationHierarchy()
        return
    def AGOMax(self):
        self.dlg.list_MaximizedProducts.addItem('AGO')
        self.ShowMaximizationHierarchy()
        return
    def ATBMax(self):
        self.dlg.list_MaximizedProducts.addItem('ATB')
        self.ShowMaximizationHierarchy()
        return
    def LVGOMax(self):
        self.dlg.list_MaximizedProducts.addItem('LVGO')
        self.ShowMaximizationHierarchy()
        return
    def HVGOMax(self):
        self.dlg.list_MaximizedProducts.addItem('HVGO')
        self.ShowMaximizationHierarchy()
        return
    def VTBMax(self):
        self.dlg.list_MaximizedProducts.addItem('VTB')
        self.ShowMaximizationHierarchy()
        return
    #Blinded Side Draws adding buttons
    def KeroseneBlindedSideDraw(self):
        self.dlg.list_BlindedSideDraws.addItem('Kerosene')
        self.ShowBlindedSideDraws()
        if len(self.dlg.list_BlindedSideDraws.findItems('Kerosene',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('Jet A',QtCore.Qt.MatchExactly)) > 1:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please only select one kerosene-range product for a blinded side draw', QtWidgets.QMessageBox.Ok)
            self.dlg.list_BlindedSideDraws.takeItem(self.dlg.list_BlindedSideDraws.count()-1)
            return
        if len(self.Refinery_List) == 0: return
        #if self.Refinery_List[self.SelectedRefinery].ProductsInstance.JetABlindedSideDrawCheck == True: return #Redundancy
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.KeroseneBlindedSideDrawCheck = True
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.NumBlindedSideDraws += 1
        return
    def JetABlindedSideDraw(self):
        self.dlg.list_BlindedSideDraws.addItem('Jet A')
        self.ShowBlindedSideDraws()
        if len(self.dlg.list_BlindedSideDraws.findItems('Kerosene',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('Jet A',QtCore.Qt.MatchExactly)) > 1:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please only select one kerosene-range product for a blinded side draw', QtWidgets.QMessageBox.Ok)
            self.dlg.list_BlindedSideDraws.takeItem(self.dlg.list_BlindedSideDraws.count()-1)
            return
        if len(self.Refinery_List) == 0: return
        #if self.Refinery_List[self.SelectedRefinery].ProductsInstance.KeroseneBlindedSideDrawCheck == True: return #Redundancy
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.JetABlindedSideDrawCheck = True   
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.NumBlindedSideDraws += 1
        return
    def DieselBlindedSideDraw(self):
        self.dlg.list_BlindedSideDraws.addItem('Diesel')
        self.ShowBlindedSideDraws()
        if len(self.dlg.list_BlindedSideDraws.findItems('Diesel',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('HS Diesel',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('ULS Diesel',QtCore.Qt.MatchExactly)) > 1:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please only select one diesel-range product for a blinded side draw', QtWidgets.QMessageBox.Ok)
            self.dlg.list_BlindedSideDraws.takeItem(self.dlg.list_BlindedSideDraws.count()-1)
            return
        if len(self.Refinery_List) == 0: return
        #if self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselBlindedSideDrawCheck == True or self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselBlindedSideDrawCheck == True: return #Redundancy
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselBlindedSideDrawCheck = True
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.NumBlindedSideDraws += 1
        return
    def ULSDieselBlindedSideDraw(self):
        self.dlg.list_BlindedSideDraws.addItem('ULS Diesel')
        self.ShowBlindedSideDraws()
        if len(self.dlg.list_BlindedSideDraws.findItems('Diesel',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('HS Diesel',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('ULS Diesel',QtCore.Qt.MatchExactly)) > 1:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please only select one diesel-range product for a blinded side draw', QtWidgets.QMessageBox.Ok)
            self.dlg.list_BlindedSideDraws.takeItem(self.dlg.list_BlindedSideDraws.count()-1)
            return
        if len(self.Refinery_List) == 0: return
        #if self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselBlindedSideDrawCheck == True or self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselBlindedSideDrawCheck == True: return #Redundancy
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselBlindedSideDrawCheck = True
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.NumBlindedSideDraws += 1
        return
    def HSDieselBlindedSideDraw(self):
        self.dlg.list_BlindedSideDraws.addItem('HS Diesel')
        self.ShowBlindedSideDraws()
        if len(self.dlg.list_BlindedSideDraws.findItems('Diesel',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('HS Diesel',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('ULS Diesel',QtCore.Qt.MatchExactly)) > 1:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please only select one diesel-range product for a blinded side draw', QtWidgets.QMessageBox.Ok)
            self.dlg.list_BlindedSideDraws.takeItem(self.dlg.list_BlindedSideDraws.count()-1)
            return
        if len(self.Refinery_List) == 0: return
        #if self.Refinery_List[self.SelectedRefinery].ProductsInstance.ULSDieselBlindedSideDrawCheck == True or self.Refinery_List[self.SelectedRefinery].ProductsInstance.DieselBlindedSideDrawCheck == True: return #Redundancy
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.HSDieselBlindedSideDrawCheck = True
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.NumBlindedSideDraws += 1
        return
    def LSMDOBlindedSideDraw(self):
        self.dlg.list_BlindedSideDraws.addItem('LSMDO')
        self.ShowBlindedSideDraws()
        if len(self.dlg.list_BlindedSideDraws.findItems('LSMDO',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('MDO',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('AGO',QtCore.Qt.MatchExactly)) > 1:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please only select one gasoil-range product for a blinded side draw', QtWidgets.QMessageBox.Ok)
            self.dlg.list_BlindedSideDraws.takeItem(self.dlg.list_BlindedSideDraws.count()-1)
            return
        if len(self.Refinery_List) == 0: return
        #if self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOBlindedSideDrawCheck == True or self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOBlindedSideDrawCheck == True: return #Redundancy
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOBlindedSideDrawCheck = True
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.NumBlindedSideDraws += 1
        return
    def MDOBlindedSideDraw(self):
        self.dlg.list_BlindedSideDraws.addItem('MDO')
        self.ShowBlindedSideDraws()
        if len(self.dlg.list_BlindedSideDraws.findItems('LSMDO',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('MDO',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('AGO',QtCore.Qt.MatchExactly)) > 1:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please only select one gasoil-range product for a blinded side draw', QtWidgets.QMessageBox.Ok)
            self.dlg.list_BlindedSideDraws.takeItem(self.dlg.list_BlindedSideDraws.count()-1)
            return
        if len(self.Refinery_List) == 0: return
        #if self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOBlindedSideDrawCheck == True or self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOBlindedSideDrawCheck == True: return #Redundancy
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOBlindedSideDrawCheck = True
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.NumBlindedSideDraws += 1
        return
    def AGOBlindedSideDraw(self):
        self.dlg.list_BlindedSideDraws.addItem('AGO')
        self.ShowBlindedSideDraws()
        if len(self.dlg.list_BlindedSideDraws.findItems('LSMDO',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('MDO',QtCore.Qt.MatchExactly)) + len(self.dlg.list_BlindedSideDraws.findItems('AGO',QtCore.Qt.MatchExactly)) > 1:
            QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please only select one gasoil-range product for a blinded side draw', QtWidgets.QMessageBox.Ok)
            self.dlg.list_BlindedSideDraws.takeItem(self.dlg.list_BlindedSideDraws.count()-1)
            return
        if len(self.Refinery_List) == 0: return
        #if self.Refinery_List[self.SelectedRefinery].ProductsInstance.MDOBlindedSideDrawCheck == True or self.Refinery_List[self.SelectedRefinery].ProductsInstance.LSMDOBlindedSideDrawCheck == True: return #Redundancy
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.AGOBlindedSideDrawCheck = True
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.NumBlindedSideDraws += 1
        return
    # Maximization commands
    def MoveUp(self):
        currentRow = self.dlg.list_MaximizedProducts.currentRow()
        currentItem = self.dlg.list_MaximizedProducts.takeItem(currentRow)
        self.dlg.list_MaximizedProducts.insertItem(currentRow - 1, currentItem)
        self.dlg.list_MaximizedProducts.setCurrentItem(self.dlg.list_MaximizedProducts.currentItem())
        return
    def MoveDown(self):
        currentRow = self.dlg.list_MaximizedProducts.currentRow()
        currentItem = self.dlg.list_MaximizedProducts.takeItem(currentRow)
        self.dlg.list_MaximizedProducts.insertItem(currentRow + 1, currentItem)
        self.dlg.list_MaximizedProducts.setCurrentItem(self.dlg.list_MaximizedProducts.currentItem())
        return
    def ClearProducts(self):
        self.dlg.list_MaximizedProducts.clear()
        self.dlg.list_MaximizedProducts.hide()
        self.dlg.cmd_MoveUpMaximizedProduct.hide()
        self.dlg.cmd_MoveDownMaximizedProduct.hide()
        self.dlg.cmd_ClearMaximizedProducts.hide()
        self.dlg.lbl_MaximizationHierarchy.hide()
        return
    def ShowMaximizationHierarchy(self):
        self.dlg.list_MaximizedProducts.show()
        self.dlg.cmd_MoveUpMaximizedProduct.show()
        self.dlg.cmd_MoveDownMaximizedProduct.show()
        self.dlg.cmd_ClearMaximizedProducts.show()
        self.dlg.lbl_MaximizationHierarchy.show()
        return
    def Close(self):
        self.dlg.close()
        exit()
        return

    #Blinded Side Draw Command
    def ClearBlindedSideDraws(self):
        self.dlg.list_BlindedSideDraws.clear()
        self.dlg.list_BlindedSideDraws.hide()
        self.dlg.cmd_ClearBlindedSideDraws.hide()
        self.dlg.lbl_BlindedSideDraws.hide()
        if len(self.Refinery_List) == 0: return
        #self.Refinery_List[self.SelectedRefinery].ProductsInstance.InitializeBlindedSideDraws()
        return
    def ShowBlindedSideDraws(self):
        self.dlg.list_BlindedSideDraws.show()
        self.dlg.cmd_ClearBlindedSideDraws.show()
        self.dlg.lbl_BlindedSideDraws.show()
        return
    # Desired Process Units on Process Unit Breakdown Tab
    def CreateCustomUnitList(self):
        try:
        

            self.CustomRefinery.__init__(self.MyProducts)

            print(self.CustomRefinery.Specified_Unit_List)
            if self.dlg.txt_sideDrawsAmtCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(SideDraw())
                self.CustomRefinery.Specified_Unit_List[-1].Name = "Atmospheric Side Draw"
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_CDUCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_sideDrawsAmtCustom.text()) - 3
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_sideDrawsPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice * self.CustomRefinery.Specified_Unit_List[-1].NumUnits))
                self.dlg.txt_sideDrawsCapCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].InletFlow))
            if self.dlg.txt_BlindedSideDrawsAmtCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(BlindedSideDraw())
                self.CustomRefinery.Specified_Unit_List[-1].Name = "Blinded Side Draw"
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_CDUCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_BlindedSideDrawsAmtCustom.text()) - 3
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_sideDrawsPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice * self.CustomRefinery.Specified_Unit_List[-1].NumUnits))
                self.dlg.txt_sideDrawsCapCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].InletFlow))
            if self.dlg.txt_desalterAmtCustom.text() != '' and self.dlg.txt_desalterCapCustom.text() != '':           
                self.CustomRefinery.Specified_Unit_List.append(Desalter())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_desalterCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_desalterAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_desalterPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_CDUAmtCustom.text() != '' and self.dlg.txt_CDUCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(CDU())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_CDUCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_CDUAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                if self.dlg.txt_sideDrawsAmtCustom.text() != '':
                    for Unit in self.CustomRefinery.Specified_Unit_List:
                        if Unit.Name == "Atmospheric Side Draw":
                            print("Price before: " + to_str(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
                            self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice = self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice + Unit.DomesticTotalPrice * Unit.NumUnits
                        if UnitName == "Blinded Side Draw":
                            self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice = self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice + Unit.DomesticTotalPrice * Unit.NumUnits
                self.dlg.txt_CDUPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_VDUAmtCustom.text() != '' and self.dlg.txt_VDUCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(VDU())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_VDUCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_VDUAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_VDUPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_NHTAmtCustom.text() != '' and self.dlg.txt_NHTCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(NaphthaHydrotreater())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_NHTCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_NHTAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_NHTPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_NaphSpAmtCustom.text() != '' and self.dlg.txt_NaphSpCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(RefluxedSplitter())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_NaphSpCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_NaphSpAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_NaphSpPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_DHTAmtCustom.text() != '' and self.dlg.txt_DHTCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(DieselHydrotreater())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_DHTCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_DHTAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_DHTPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_distSpAmtCustom.text() != '' and self.dlg.txt_distSpCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(RefluxedSplitter())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_distSpCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_distSpAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_distSpPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_CRUAmtCustom.text() != '' and self.dlg.txt_CRUCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(CRU())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_CRUCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_CRUAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_CRUPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_isomAmtCustom.text() != '' and self.dlg.txt_isomCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(IsomerizationUnit())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_isomCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_isomAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_isomPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_stabilizerAmtCustom.text() != '' and self.dlg.txt_stabilizerCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(RefluxedStabilizer())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_stabilizerCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_stabilizerAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_stabilizerPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_KHTAmtCustom.text() != '' and self.dlg.txt_KHTCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(KeroseneHydrotreater())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_KHTCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_KHTAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_KHTPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_alloyAmtCustom.text() != '' and self.dlg.txt_alloyCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(AdditionalAlloy())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_alloyCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_alloyAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_alloyPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_CausticTreaterAmtCustom.text() != '' and self.dlg.txt_CausticTreaterCapCustom.text() != '':
                self.CustomRefinery.Specified_Unit_List.append(CausticTreater())
                self.CustomRefinery.Specified_Unit_List[-1].InletFlow = int(self.dlg.txt_CausticTreaterCapCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].NumUnits = int(self.dlg.txt_CausticTreaterAmtCustom.text())
                self.CustomRefinery.Specified_Unit_List[-1].TotalFlow = self.CustomRefinery.Specified_Unit_List[-1].InletFlow * self.CustomRefinery.Specified_Unit_List[-1].NumUnits
                self.CustomRefinery.Specified_Unit_List[-1].CalculatePrice()
                self.dlg.txt_CausticTreaterPriceCustom.setText(special_format(self.CustomRefinery.Specified_Unit_List[-1].DomesticTotalPrice))
            if self.dlg.txt_ThiolexCapCustom.text() != '':
                self.CustomRefinery.Utilities.SulfurOutTotal = float(self.dlg.txt_ThiolexCapCustom.text())
                self.CustomRefinery.AddSRU()
                if self.CustomRefinery.MyMerichemThiolex.NumUnits >= 1:
                    self.CustomRefinery.MyMerichemThiolex.CalculatePrice()
                    self.CustomRefinery.Specified_Unit_List.append(CustomRefinery.MyMerichemThiolex)
                    self.dlg.txt_alloyPriceCustom.setText(special_format(CustomRefinery.MyMerichemThiolex))
                if self.CustomRefinery.MyMerichemLoCat.NumUnits >= 1:
                    CustomRefinery.MyMerichemLoCat.CalculatePrice()
                    CustomRefinery.Specified_Unit_List.append(CustomRefinery.MyMerichemLoCat)
                self.CustomRefinery.UnitListBreakdown()
                self.dlg.txt_ThiolexPriceCustom.setText(special_format(CustomRefinery.ThiolexPrice))
                self.dlg.txt_ThiolexAmtCustom.setText(special_format(CustomRefinery.ThiolexAmount))
                self.dlg.txt_LoCatPriceCustom.setText(special_format(CustomRefinery.LoCatPrice))
                self.dlg.txt_LoCatAmtCustom.setText(special_format(CustomRefinery.LoCatAmount))
            if self.dlg.txt_LoCatCapCustom.text() != '':
                self.CustomRefinery.Utilities.SulfurOutTotal = float(self.dlg.txt_LoCatCapCustom.text())
                self.CustomRefinery.AddSRU()
                if self.CustomRefinery.MyMerichemThiolex.NumUnits >= 1:
                    self.CustomRefinery.MyMerichemThiolex.CalculatePrice()
                    self.CustomRefinery.Specified_Unit_List.append(CustomRefinery.MyMerichemThiolex)
                    self.dlg.txt_alloyPriceCustom.setText(special_format(CustomRefinery.MyMerichemThiolex))
                if self.CustomRefinery.MyMerichemLoCat.NumUnits >= 1:
                    CustomRefinery.MyMerichemLoCat.CalculatePrice()
                    CustomRefinery.Specified_Unit_List.append(CustomRefinery.MyMerichemLoCat)
                self.CustomRefinery.UnitListBreakdown()
                self.dlg.txt_ThiolexPriceCustom.setText(special_format(CustomRefinery.ThiolexPrice))
                self.dlg.txt_ThiolexAmtCustom.setText(special_format(CustomRefinery.ThiolexAmount))
                self.dlg.txt_LoCatPriceCustom.setText(special_format(CustomRefinery.LoCatPrice))
                self.dlg.txt_LoCatAmtCustom.setText(special_format(CustomRefinery.LoCatAmount))


            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Creating Custom Unit. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def ClearCustomUnit(self):
        try:
            self.CustomRefinery.ReInitialze()

            self.dlg.txt_desalterAmtCustom.setText('')
            self.dlg.txt_desalterCapCustom.setText('')
            self.dlg.txt_desalterPriceCustom.setText('')
            self.dlg.txt_CDUAmtCustom.setText('')
            self.dlg.txt_CDUCapCustom.setText('')
            self.dlg.txt_CDUPriceCustom.setText('')
            self.dlg.txt_BlindedSideDrawsAmtCustom.clear()
            self.dlg.txt_BlindedSideDrawsCapCustom.clear()
            self.dlg.txt_BlindedSideDrawsPriceCustom.clear()
            self.dlg.txt_sideDrawsAmtCustom.setText('')
            self.dlg.txt_sideDrawsCapCustom.setText('')
            self.dlg.txt_sideDrawsPriceCustom.setText('')
            self.dlg.txt_VDUAmtCustom.setText('')
            self.dlg.txt_VDUCapCustom.setText('')
            self.dlg.txt_VDUPriceCustom.setText('')
            self.dlg.txt_NHTAmtCustom.setText('')
            self.dlg.txt_NHTCapCustom.setText('')
            self.dlg.txt_NHTPriceCustom.setText('')
            self.dlg.txt_NaphSpAmtCustom.setText('')
            self.dlg.txt_NaphSpCapCustom.setText('')
            self.dlg.txt_NaphSpPriceCustom.setText('')
            self.dlg.txt_DHTAmtCustom.setText('')
            self.dlg.txt_DHTCapCustom.setText('')
            self.dlg.txt_DHTPriceCustom.setText('')
            self.dlg.txt_distSpAmtCustom.setText('')
            self.dlg.txt_distSpCapCustom.setText('')
            self.dlg.txt_distSpPriceCustom.setText('')
            self.dlg.txt_CRUAmtCustom.setText('')
            self.dlg.txt_CRUCapCustom.setText('')
            self.dlg.txt_CRUPriceCustom.setText('')
            self.dlg.txt_gasPlantAmtCustom.setText('')
            self.dlg.txt_gasPlantCapCustom.setText('')
            self.dlg.txt_gasPlantPriceCustom.setText('')
            self.dlg.txt_isomAmtCustom.setText('')
            self.dlg.txt_isomCapCustom.setText('')
            self.dlg.txt_isomPriceCustom.setText('')
            self.dlg.txt_stabilizerAmtCustom.setText('')
            self.dlg.txt_stabilizerCapCustom.setText('')
            self.dlg.txt_stabilizerPriceCustom.setText('')
            self.dlg.txt_KHTAmtCustom.setText('')
            self.dlg.txt_KHTCapCustom.setText('')
            self.dlg.txt_KHTPriceCustom.setText('')
            self.dlg.txt_ARAHCUAmtCustom.setText('')
            self.dlg.txt_ARAHCUCapCustom.setText('')
            self.dlg.txt_ARAHCUPriceCustom.setText('')
            self.dlg.txt_ARACHAmtCustom.setText('')
            self.dlg.txt_ARACHCapCustom.setText('')
            self.dlg.txt_ARACHPriceCustom.setText('')
            self.dlg.txt_alloyCapCustom.setText('')
            self.dlg.txt_alloyPriceCustom.setText('')
            self.dlg.txt_SMRAmtCustom.setText('')
            self.dlg.txt_SMRCapCustom.setText('')
            self.dlg.txt_SMRPriceCustom.setText('')
            self.dlg.txt_ThiolexAmtCustom.setText('')
            self.dlg.txt_ThiolexCapCustom.setText('')
            self.dlg.txt_ThiolexPriceCustom.setText('')
            self.dlg.txt_LoCatAmtCustom.setText('')
            self.dlg.txt_LoCatCapCustom.setText('')
            self.dlg.txt_LoCatPriceCustom.setText('')
            self.dlg.txt_SNRAmtCustom.setText('')
            self.dlg.txt_SNRCapCustom.setText('')
            self.dlg.txt_SNRPriceCustom.setText('')
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Adding Assay. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
        return
    def ClearChemexUnit(self):
        self.dlg.txt_desalterAmtCHEMEX.setText('')
        self.dlg.txt_desalterCapCHEMEX.setText('')
        self.dlg.txt_desalterPriceCHEMEX.setText('')
        self.dlg.txt_CDUAmtCHEMEX.setText('')
        self.dlg.txt_CDUCapCHEMEX.setText('')
        self.dlg.txt_CDUPriceCHEMEX.setText('')
        self.dlg.txt_BlindedSideDrawsAmtCHEMEX.clear()
        self.dlg.txt_BlindedSideDrawsCapCHEMEX.clear()
        self.dlg.txt_BlindedSideDrawsPriceCHEMEX.clear()
        self.dlg.txt_sideDrawsAmtCHEMEX.setText('')
        self.dlg.txt_sideDrawsCapCHEMEX.setText('')
        self.dlg.txt_sideDrawsPriceCHEMEX.setText('')
        self.dlg.txt_VDUAmtCHEMEX.setText('')
        self.dlg.txt_VDUCapCHEMEX.setText('')
        self.dlg.txt_VDUPriceCHEMEX.setText('')
        self.dlg.txt_NHTAmtCHEMEX.setText('')
        self.dlg.txt_NHTCapCHEMEX.setText('')
        self.dlg.txt_NHTPriceCHEMEX.setText('')
        self.dlg.txt_NaphSpAmtCHEMEX.setText('')
        self.dlg.txt_NaphSpCapCHEMEX.setText('')
        self.dlg.txt_NaphSpPriceCHEMEX.setText('')
        self.dlg.txt_DHTAmtCHEMEX.setText('')
        self.dlg.txt_DHTCapCHEMEX.setText('')
        self.dlg.txt_DHTPriceCHEMEX.setText('')
        self.dlg.txt_distSpAmtCHEMEX.setText('')
        self.dlg.txt_distSpCapCHEMEX.setText('')
        self.dlg.txt_distSpPriceCHEMEX.setText('')
        self.dlg.txt_CRUAmtCHEMEX.setText('')
        self.dlg.txt_CRUCapCHEMEX.setText('')
        self.dlg.txt_CRUPriceCHEMEX.setText('')
        self.dlg.txt_gasPlantAmtCHEMEX.setText('')
        self.dlg.txt_gasPlantCapCHEMEX.setText('')
        self.dlg.txt_gasPlantPriceCHEMEX.setText('')
        self.dlg.txt_isomAmtCHEMEX.setText('')
        self.dlg.txt_isomCapCHEMEX.setText('')
        self.dlg.txt_isomPriceCHEMEX.setText('')
        self.dlg.txt_stabilizerAmtCHEMEX.setText('')
        self.dlg.txt_stabilizerCapCHEMEX.setText('')
        self.dlg.txt_stabilizerPriceCHEMEX.setText('')
        self.dlg.txt_KHTAmtCHEMEX.setText('')
        self.dlg.txt_KHTCapCHEMEX.setText('')
        self.dlg.txt_KHTPriceCHEMEX.setText('')
        self.dlg.txt_ARAHCUAmtCHEMEX.setText('')
        self.dlg.txt_ARAHCUCapCHEMEX.setText('')
        self.dlg.txt_ARAHCUPriceCHEMEX.setText('')
        self.dlg.txt_ARACHAmtCHEMEX.setText('')
        self.dlg.txt_ARACHCapCHEMEX.setText('')
        self.dlg.txt_ARACHPriceCHEMEX.setText('')
        self.dlg.txt_alloyCapCHEMEX.setText('')
        self.dlg.txt_alloyPriceCHEMEX.setText('')
        self.dlg.txt_SMRAmtCHEMEX.setText('')
        self.dlg.txt_SMRCapCHEMEX.setText('')
        self.dlg.txt_SMRPriceCHEMEX.setText('')
        self.dlg.txt_ThiolexAmtCHEMEX.setText('')
        self.dlg.txt_ThiolexCapCHEMEX.setText('')
        self.dlg.txt_ThiolexPriceCHEMEX.setText('')
        self.dlg.txt_LoCatAmtCHEMEX.setText('')
        self.dlg.txt_LoCatCapCHEMEX.setText('')
        self.dlg.txt_LoCatPriceCHEMEX.setText('')
        self.dlg.txt_SNRAmtCHEMEX.setText('')
        self.dlg.txt_SNRCapCHEMEX.setText('')
        self.dlg.txt_SNRPriceCHEMEX.setText('')
        self.dlg.txt_processUnitsCostCHEMEX.setText('')
        return
    #Desired/CHEMEX Calculated Process Units Utilities
    def CustomCurrentlySelectedRefineryCombobox(self):
        if self.dlg.cmb_CustomCurrentlySelectedRefineryCombobox.currentText() == "Select Refinery":
            self.dlg.cmb_CustomUnitOperation.clear()
            self.dlg.cmb_CustomUnitOperation.addItems(['Select Unit Operation'])
            return

        self.SelectedRefinery = self.dlg.cmb_CustomCurrentlySelectedRefineryCombobox.currentText()

        return
    def CustomCurrentlySelectedUnitOperationCombobox(self):
        a=5
        return
    def CurrentlySelectedRefineryCombobox(self):
        #if self.dlg.cmb_CurrentlySelectedRefinery.currentText() == "Select Refinery": #If the user goes back to current refinery we want to clear everything so they can make another decision
         #   self.dlg.cmb_UnitOperation.clear()
          #  self.ClearChemexCalculatedProcessUnitsUtilities()
            
        if self.dlg.cmb_CurrentlySelectedRefinery.currentText() != self.SelectedRefinery:
            self.dlg.cmb_UnitOperation.clear()
            self.ClearChemexCalculatedProcessUnitsUtilities()
            if self.dlg.cmb_CurrentlySelectedRefinery.currentText() == "Select Refinery": #If the user goes back to current refinery we want to clear everything so they can make another decision
                self.ClearChemexUnit()
                return

        self.SelectedRefinery = self.dlg.cmb_CurrentlySelectedRefinery.currentText()
        self.dlg.txt_SelectedRefinery.setText(self.SelectedRefinery) #We must do this so that the SelectedRefinery Method will work
        self.dlg.cmb_UnitOperation.addItems(['Select Unit Operation', self.SelectedRefinery])
        if self.FillOutCHEMEXUnitVariables() == None:
            self.ClearChemexUnit()
        self.FillOutDashboardVariables()
        self.UploadImage()
        
        self.dlg.cmb_UnitOperation.clear()
        self.dlg.cmb_UnitOperation.addItems(['Select Unit Operation', self.SelectedRefinery])
        for UnitOp in self.Refinery_List[self.SelectedRefinery].Unit_List:
            self.dlg.cmb_UnitOperation.addItems([UnitOp.Name])

        return
    def CurrentlySelectedUnitOperationCombobox(self):
        if self.dlg.cmb_UnitOperation.currentText() == "Select Unit Operation":
            self.ClearChemexCalculatedProcessUnitsUtilities()
            return
        self.FillOutChemexCalculatedProcessUnitsUtilities()
        return
    def ClearChemexCalculatedProcessUnitsUtilities(self):
        self.dlg.txt_SulfurOut.clear()
        self.dlg.txt_SCFH2pD.clear()
        self.dlg.txt_Steam.clear()
        self.dlg.txt_Power.clear()
        self.dlg.txt_CoolingTowerPower.clear()
        self.dlg.txt_CoolingWater.clear()
        self.dlg.txt_Fuel.clear()
        self.dlg.txt_DesalterWater.clear()
        return
    def FillOutChemexCalculatedProcessUnitsUtilities(self):
        if self.dlg.cmb_UnitOperation.currentText() == self.dlg.cmb_CurrentlySelectedRefinery.currentText():
            self.dlg.txt_SulfurOut.setText(to_str(self.Refinery_List[self.SelectedRefinery].TotalUtilities.SulfurOutTotal))
            self.dlg.txt_SCFH2pD.setText(to_str(self.Refinery_List[self.SelectedRefinery].TotalUtilities.SCFH2pD))
            self.dlg.txt_Steam.setText(to_str(self.Refinery_List[self.SelectedRefinery].TotalUtilities.Steam))
            self.dlg.txt_Power.setText(to_str(self.Refinery_List[self.SelectedRefinery].TotalUtilities.Power))
            self.dlg.txt_CoolingTowerPower.setText(to_str(self.Refinery_List[self.SelectedRefinery].TotalUtilities.CoolingTowerPower))
            self.dlg.txt_CoolingWater.setText(to_str(self.Refinery_List[self.SelectedRefinery].TotalUtilities.CoolingWater))
            self.dlg.txt_Fuel.setText(to_str(self.Refinery_List[self.SelectedRefinery].TotalUtilities.Fuel))
            self.dlg.txt_DesalterWater.setText(to_str(self.Refinery_List[self.SelectedRefinery].TotalUtilities.DesalterWater))
            return
        for UnitOp in self.Refinery_List[self.SelectedRefinery].Unit_List:
            if UnitOp.Name == self.dlg.cmb_UnitOperation.currentText():
                self.dlg.txt_SulfurOut.setText(to_str(UnitOp.Utilities.SulfurOutTotal))
                self.dlg.txt_SCFH2pD.setText(to_str(UnitOp.Utilities.SCFH2pD))
                self.dlg.txt_Steam.setText(to_str(UnitOp.Utilities.Steam))
                self.dlg.txt_Power.setText(to_str(UnitOp.Utilities.Power))
                self.dlg.txt_CoolingTowerPower.setText(to_str(UnitOp.Utilities.CoolingTowerPower))
                self.dlg.txt_CoolingWater.setText(to_str(UnitOp.Utilities.CoolingWater))
                self.dlg.txt_Fuel.setText(to_str(UnitOp.Utilities.Fuel))
                self.dlg.txt_DesalterWater.setText(to_str(UnitOp.Utilities.DesalterWater))
                break
        return
    def ClearDesiredProcessUnitsUtilities(self):
        self.dlg.txt_CustomSulfurOut.clear()
        self.dlg.txt_CustomSCFH2pD.clear()
        self.dlg.txt_CustomSteam.clear()
        self.dlg.txt_CustomPower.clear()
        self.dlg.txt_CustomCoolingTowerPower.clear()
        self.dlg.txt_CustomCoolingWater.clear()
        self.dlg.txt_CustomFuel.clear()
        self.dlg.txt_CustomDesalterWater.clear()
        return
    def FillOutDesiredProcessUnitsUtilities(self):
        a=5
        return

    #Blending Manager/Sandbox
    def SettingsSelectedRefinery(self):
        if self.dlg.cmb_SettingsSelectedRefinery.currentText() == "Select Refinery":
            self.ClearTemperatureCutPoints()
            return
        self.dlg.list_assays.setCurrentRow(self.dlg.cmb_SettingsSelectedRefinery.currentIndex()-1)
        self.SelectRefinery()
        return
    def ClearTemperatureCutPoints(self):
        self.dlg.txt_SettingsIBP.clear()
        for i in range(0,len(self.Temperature_Cut_Point_List)):
            self.Temperature_Cut_Point_List[i].clear()
        self.dlg.txt_SettingsFBP.clear()
        return
    def SaveSettingsForCurrentRefinery(self):
        self.SaveProductInfo()
        return
    def SaveSettingsForAllRefineries(self):
        self.SaveAllProductInfo()
        return

    # transfer over needed information from the assay database
    def TransferClicked(self):
        wb = load_workbook(r'C:\Users\ekaufmann\source\repos\Assay OCR Project\database\Assays.xlsm')
        ws = wb['All']
        for j in range(0, len(self.assaysToTest)):
            while ws.cell(1, row).value != self.assaysToTest[j]:
                row = row + 1      
    # Generates Proposal with User Input
    def addImages(self, doc, firstPage): 
        if dlg.cmb_companys.currentText() == 'Chemex Modular, LLC': 
            self.Dictionary.update({'OurAbbrev': 'Modular'})
            if firstPage:
                r = doc.tables[0].rows[0].cells[0].paragraphs[0].add_run()
                r.add_picture(logo)
            r = doc.sections[0].header.tables[0].rows[0].cells[1].paragraphs[0].add_run()
            r.add_picture(logo)
        elif dlg.cmb_companys.currentText() == 'Pan Africa Chemex Limited': 
            self.Dictionary.update({'OurAbbrev': 'Africa'})
            if firstPage:
                r = doc.tables[0].rows[0].cells[0].paragraphs[0].add_run()
                r.add_picture(logo)
            r = doc.sections[0].header.tables[0].rows[0].cells[1].paragraphs[0].add_run()
            r.add_picture(logo)
        elif dlg.cmb_companys.currentText() == 'Chemex Modular India Private Limited': 
            self.Dictionary.update({'OurAbbrev': 'India'})
            if firstPage:
                r = doc.tables[0].rows[0].cells[0].paragraphs[0].add_run()
                r.add_picture(logo)
            r = doc.sections[0].header.tables[0].rows[0].cells[1].paragraphs[0].add_run()
            r.add_picture(logo)
        return
    def remove_row(self,table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)
        return
    def removeEmptyRows(self,table, column, row, offset, endValue, doc):
        while table.cell(row, column - offset).text != endValue:
            if table.cell(row, column).text == '' or table.cell(row, column).text == 'None': 
                rowToDelete = table.rows[row]
                self.remove_row(table, rowToDelete)
            else: 
                row = row + 1
        if table.cell(row, column).text == '' or table.cell(row, column).text == 'None': 
            rowToDelete = table.rows[row]
            self.remove_row(table, rowToDelete)
        else: 
            row = row + 1
        table.style = doc.styles['List Table 1 Light'] # removes empty rows
        return
    # get input values from user interface
    def getCompany(self): 
        global logo
        if self.dlg.cmb_companys.currentText() == 'Chemex Modular, LLC': 
            logo = 'chemex modular.png'
            self.Dictionary.update({'Company': self.dlg.cmb_companys.currentText(), 'OurAbbrev': 'CML'})
        elif self.dlg.cmb_companys.currentText() == 'Pan Africa Chemex Limited':  
            logo = 'pan africa.jpg'
            self.Dictionary.update({'Company': self.dlg.cmb_companys.currentText(), 'OurAbbrev': 'Africa'})
        elif self.dlg.cmb_companys.currentText() == 'Chemex Modular India Private Limited':  
            logo = 'india.jpg'
            self.Dictionary.update({'Company': self.dlg.cmb_companys.currentText(), 'OurAbbrev': 'India'})
        return
    def getContactInfo(self): 
        if self.dlg.cmb_contact.currentText() == 'Johnny Hallford':  self.Dictionary.update({'Name': 'Johnny Hallford', 'Title': 'Director of Sales', 'Phone': '(979) 777-1951', 'Email': 'jhallford@chemexmodular.com'})
        elif self.dlg.cmb_contact.currentText() == 'Matt Rodgers':  self.Dictionary.update({'Name': 'Matt Rodgers', 'Title': 'Chief Commercial Officer', 'Phone': '(936) 520-7068', 'Email': 'mrodgers@chemexmodular.com'})
        elif self.dlg.cmb_contact.currentText() == 'Joseph DeSpain': self.Dictionary.update({'Name': 'Joseph DeSpain', 'Title': 'Sales/Applications Engineer', 'Phone': '(630) 697-2343', 'Email': 'jdespain@chemexmodular.com'})
        elif self.dlg.cmb_contact.currentText() == 'Other': self.Dictionary.update({'Name': self.dlg.txt_name.text(), 'Title': self.dlg.txt_title.text(), 'Phone': self.dlg.txt_phone.text(), 'Email': self.dlg.txt_email.text()})
        return
    def combine_all_docx(self,filename_master):
        number_of_sections=len(self.files_list)
        master = Document(filename_master)
        composer = Composer(master)
        for i in range(0, number_of_sections):
            doc_temp = Document(self.files_list[i])
            composer.append(doc_temp)
        composer.save(os.getcwd() + '\\Proposal Templates\\generation\\' + 'Generated_Proposal.docx')
        return

    def getTemplate(self):
        path = os.getcwd() + '\\Proposal Templates\\'
        output = self.dlg.txt_folder.text() + '/'
        #try:
        if self.dlg.chk_proposal.isChecked():
            output_file = output + 'CML - ' + self.Dictionary["ReferenceNum"] + ' - ' + self.Dictionary["Year"] + self.Dictionary["Month"] + self.Dictionary["Day"] + ' (' + self.Dictionary["Rev#"] + ') Equipment Supply Proposal' + '.docx'
            template = path + 'APE-D2 Proposal Template.docx'
            self.proposal(template,output_file)
            QtWidgets.QMessageBox.information(self.dlg, 'Success', 'Proposal Successfully Created', QtWidgets.QMessageBox.Ok)
        if self.dlg.chk_feasibilityStudy.isChecked(): 
            output_file = output + 'CML - ' + self.Dictionary["ReferenceNum"] + ' - ' + self.Dictionary["Year"] + self.Dictionary["Month"] + self.Dictionary["Day"] + ' (' + self.Dictionary["Rev#"] + ') Feasibility Study' + '.docx'
            template = path + 'APE-D2 Feasibility Study Template.docx'
            self.feasibilityStudy(template,output_file)
        if self.dlg.chk_definitionStudy.isChecked() : 
            output_file = output + 'CML - ' + self.Dictionary["ReferenceNum"] + ' - ' + self.Dictionary["Year"] + self.Dictionary["Month"] + self.Dictionary["Day"] + ' (' + self.Dictionary["Rev#"] + ') Definition Study' + '.docx'
            template = path + 'APE-D2 Definition Study Template.docx'
            self.definitionStudy(template,output_file)
        if self.dlg.chk_emissionStudy.isChecked(): 
            output_file = output + 'CML - ' + self.Dictionary["ReferenceNum"] + ' - ' + self.Dictionary["Year"] + self.Dictionary["Month"] + self.Dictionary["Day"] + ' (' + self.Dictionary["Rev#"] + ') Emissions Study' + '.docx'
            template = path + 'APE-D2 Emission Study Template.docx'
            self.emissionStudy(template,output_file)
        if self.dlg.chk_OSBLISBL.isChecked(): 
            output_file = output + 'CML - ' + self.Dictionary["ReferenceNum"] + ' - ' + self.Dictionary["Year"] + self.Dictionary["Month"] + self.Dictionary["Day"] + ' (' + self.Dictionary["Rev#"] + ') FEED Study (ISBL + OSBL)' + '.docx'
            template = path + 'APE-D2 OSBL ISBL Study Template.docx'
            self.osblisblStudy(template,output_file)
        if self.dlg.chk_costData.isChecked(): 
            output_file = output + 'CML - ' + self.Dictionary["ReferenceNum"] + ' - ' + self.Dictionary["Year"] + self.Dictionary["Month"] + self.Dictionary["Day"] + ' (' + self.Dictionary["Rev#"] + ') Cost Data' + '.docx'
            template = path + 'APE-D2 Cost Data Template.docx'
            self.costData(template,output_file)
        if self.dlg.chk_letterOfInterest.isChecked(): 
            output_file = output + 'CML - ' + self.Dictionary["ReferenceNum"] + ' - ' + self.Dictionary["Year"] + self.Dictionary["Month"] + self.Dictionary["Day"] + ' (' + self.Dictionary["Rev#"] + ') Letter Of Interest' + '.docx'
            template = path + 'APE-D2 Letter of Interest Template.docx'
            self.letterOfInterest(template,output_file)
        if self.dlg.chk_servicesOrder.isChecked(): 
            output_file = output + 'CML - ' + self.Dictionary["ReferenceNum"] + ' - ' + self.Dictionary["Year"] + self.Dictionary["Month"] + self.Dictionary["Day"] + ' (' + self.Dictionary["Rev#"] + ') Services Order' + '.docx'
            template = path + 'APE-D2 Services Order Template.docx'
            self.servicesOrder(template,output_file)
        #except:
        #    QtWidgets.QMessageBox.information(self.dlg, 'Error', 'Please use a valid path', QtWidgets.QMessageBox.Ok)
        return
    #________________________________________________________________________________Proposal Generation____
    def proposal(self,template,output_file):

        
        TemporaryFileName = "Temporary File.docx"
        #mail_merge_document = MailMerge(template)
        #convert_dictionary_values_to_str(self.Dictionary)
        MyDocument = Document(template)
        MyDocument.save(TemporaryFileName)
        YieldsTable = MyDocument.tables[2]
        print("*****************************************************************************************************")
        print("This is the Proposal Dictionary:")
        print(self.Dictionary)
        TableIndex = 2 #Index number accounting for word doc headers, which are also technically tables
        for key in self.Proposal_Dictionary_List.keys():
            print("This is the Assay Dictionary: ")
            print(self.Proposal_Dictionary_List[key])
            self.Dictionary.update(self.Proposal_Dictionary_List[key].copy())
            mail_merge_document = MailMerge(os.getcwd() + '\\' + TemporaryFileName)
            mail_merge_document.merge(**self.Dictionary)
            mail_merge_document.write(output_file)
            MyDocument = Document(output_file)
            new_p = OxmlElement("w:p")
            MyDocument.tables[TableIndex]._tbl.addnext(new_p)
            paragraph = Paragraph(new_p,MyDocument.tables[TableIndex])
            copy_table_after(YieldsTable,paragraph)
            print("*****************************************************************************************************")
            print("This is the Proposal Dictionary:")
            print(self.Dictionary)
            self.removeEmptyRows(MyDocument.tables[TableIndex], 2, 1, 2, 'Total', MyDocument) # yield calculations table
            TableIndex += 1
            MyDocument.save(TemporaryFileName)
        
        mail_merge_document.close()
        os.remove(TemporaryFileName)

        #mail_merge_document.merge(**self.Dictionary)
        #mail_merge_document.write(output_file) # merges fields

        doc = Document(output_file)

        # self.addImages(doc, True) # adds correspoding company logo
        self.removeEmptyRows(doc.tables[TableIndex-1], 2, 1, 2, 'Total', doc) # yield calculations table
        self.removeEmptyRows(doc.tables[TableIndex], 2, 2, 1, 'Gas Plant, MMSCFD', doc) # major equipment summary
        self.removeEmptyRows(doc.tables[TableIndex+1], 3, 2, 1, 'Hydrogen, scf/bbl', doc ) # utilities preliminary estimate


        #self.removeEmptyRows(doc.tables[2], 2, 1, 2, 'Total', doc) # yield calculations table
        #self.removeEmptyRows(doc.tables[3], 2, 2, 1, 'Gas Plant, MMSCFD', doc) # major equipment summary
        #self.removeEmptyRows(doc.tables[4], 3, 2, 1, 'Annual Spares Cost, USD', doc ) # utilities preliminary estimate
        doc.save(output_file) # saves file
        os.startfile(output_file) # opens file


        return
    #________________________________________________________________________Feasibility Study Generation____    
    def feasibilityStudy(self,template,output_file):
        document = MailMerge(template)
        document.merge(**self.Dictionary)
        document.write(output_file) # merges fields
        doc = Document(output_file)

        self.addImages(doc, True) # adds correspoding company logo

        doc.save(output_file) # saves file
        os.startfile(output_file) # opens file
    #___________________________________________________________________________Emission Study Generation____    ---bullet points
    def emissionStudy(self,template,output_file):
        document = MailMerge(template)
        document.merge(**self.Dictionary)
        document.write(output_file) # merges fields
        doc = Document(output_file)

        self.addImages(doc, True) # adds correspoding company logo

        doc.save(output_file) # saves file
        os.startfile(output_file) # opens file
    #_________________________________________________________________________Definition Study Generation____
    def definitionStudy(self,template,output_file):
        document = MailMerge(template)
        document.merge(**self.Dictionary)
        document.write(output_file) # merges fields
        doc = Document(output_file)

        self.addImages(doc, True) # adds correspoding company logo

        doc.save(output_file) # saves file
        os.startfile(output_file) # opens file
    #_______________________________________________________________________Letter of Interest Generation____
    def letterOfInterest(self,template,output_file):
        document = MailMerge(template)
        document.merge(**self.Dictionary)
        document.write(output_file) # merges fields
        doc = Document(output_file)

        self.addImages(doc, False) # adds correspoding company logo
        if self.dlg.cmb_contact.currentText() == 'Matt Rodgers': signature = 'signature3.png'
        elif self.dlg.cmb_contact.currentText() == 'Roscoe Vasquez': signature = 'signature1.png'
        elif self.dlg.cmb_contact.currentText() == 'Johnny Hallford': signature = 'signature2.png'
        r = doc.tables[1].rows[0].cells[0].paragraphs[0].add_run()
        r.add_picture(signature) # adds signature image
        doc.save(output_file) # saves file
        os.startfile(output_file) # opens file
    #__________________________________________________________________________OSBL ISBL Study Generation____
    def osblisblStudy(self,template,output_file):
        document = MailMerge(template)
        document.merge(**self.Dictionary)
        document.write(output_file) # merges fields
        doc = Document(output_file)

        self.addImages(doc, True) # adds correspoding company logo

        doc.save(output_file) # saves file
        os.startfile(output_file) # opens file
    #___________________________________________________________________________Services Order Generation____    ---description
    def servicesOrder(self,template,output_file):
        document = MailMerge(template)
        document.merge(**self.Dictionary)
        document.write(output_file) # merges fields
        doc = Document(output_file)

        self.addImages(doc, True) # adds correspoding company logo
       
        doc.save(output_file) # saves file
        os.startfile(output_file) # opens file
    #________________________________________________________________________________Cost Data Generation____
    def costData(self,template,output_file):
        document = MailMerge(template)
        document.merge(**self.Dictionary)
        document.write(output_file) # merges fields
        doc = Document(output_file)
        print("*****************************************************************************************************")
        print("This is the Proposal Dictionary:")
        print(self.Dictionary)
        self.addImages(doc, False) # adds correspoding company logo
        self.removeEmptyRows(doc.tables[1], 5, 1, 4, 'Isom Unit', doc) # cost data table
        #self.removeEmptyRows(doc.tables[1], 2, 1, 1, 'Isom Unit', doc) # cost data table
        doc.save(output_file) # saves file
        os.startfile(output_file) # opens file    

    def GenerateClicked(self): 

        try:
            SentBy = self.dlg.cmb_contact.currentText() #Assigned Correctly       
            #InstallLocation = "Puntas de Mata, Monagas States, Venezuela" #Easily Assigned
            ReferenceNum = self.dlg.txt_ReferenceNumber.text() # Correctly assigned (for now)
            RevNumber = self.dlg.txt_RevisionNumber.text() # Correctly Assigned (for now)
            Company = self.dlg.cmb_companys.currentText() #Assigned Correctly
            InternationalRate = 0.15
            DefinitionStudyCost = self.dlg.txt_DefinitionStudyCost.text()
            FeasibilityStudyCost = self.dlg.txt_FeasibilityStudyCost.text()
            ProcessDesignStudyCost = self.dlg.txt_ProcessDesignStudyCost.text()
            if self.dlg.btn_budgetary.isChecked():
                self.Dictionary.update({'Budget_Fixed': 'budgetary', '+or-': '± %' + dlg.txt_budgetary.text(), "Pricing": 'Budgetary'})
            elif self.dlg.btn_fixed.isChecked():
                self.Dictionary.update({'Budget_Fixed': 'fixed', '+or-': '', "Pricing": 'Fixed'})

            self.Dictionary.update({"ClientName": self.clientName, "SENT_BY": SentBy })
            self.Dictionary.update({"ReferenceNum": to_str(ReferenceNum)})
            self.Dictionary.update({"Rev#": to_str(RevNumber), "Company": Company, "Def_Cost": to_str(self.DefinitionStudyCost), "Feas_Cost": to_str(self.FeasibilityStudyCost), "PDP_Cost": to_str(self.ProcessDesignStudyCost), "Commission_Rate": to_str(self.commissionValue)})

            for RefineryInstance in self.Refinery_List.values():
                if RefineryInstance.ProductsInstance.IsCalculated == False:
                    QtWidgets.QMessageBox.information(self.dlg, 'Error', 'You must Calculate Products before writing proposal', QtWidgets.QMessageBox.Ok)
                    return
            self.AssignDashboardVariables()
            self.CreateDictionary()
            self.getCompany() # assigns company abbreviation
            self.getContactInfo() # assigns contact info
            self.getTemplate() # gets type of proposal and corresponding template
            return
        except:
            QtWidgets.QMessageBox.information(self.dlg, 'UNEXPECTED CRASH', 'Error Occured while Generating Proposal. Please Check Inputs/Consult Code', QtWidgets.QMessageBox.Ok)
            raise
        return
    # Creates unit and updates graph
    def UploadImage(self):
        #pixmap = QPixmap('TemporaryPieChart.svgz')
        #pixmap = pixmap.scaled(250,300)
        #self.dlg.lbl_dashboardimagelabel.setPixmap(pixmap)
        if len(self.Refinery_List.values()) == 0:
            self.dlg.wdg_DashboardPieChart.canvas.axes.clear()
            self.dlg.wdg_DashboardPieChart.canvas.axes.pie([1], colors = ['blue'])
            self.dlg.wdg_DashboardPieChart.canvas.draw()
            return
        if self.Refinery_List[self.SelectedRefinery].ProductsInstance.IsCalculated:
            self.dlg.wdg_DashboardPieChart.canvas.axes.clear()
            colorArray = ['goldenrod','turquoise','springgreen','green','red','maroon','sienna','dimgray']
            labels = self.Refinery_List[self.SelectedRefinery].ProductsInstance.Product_Names
            cuts = self.Refinery_List[self.SelectedRefinery].ProductsInstance.Yield_List
            self.dlg.wdg_DashboardPieChart.canvas.axes.pie(cuts, labels=labels, shadow = True, autopct ='%1.1f%%', colors = colorArray,)
            self.dlg.wdg_DashboardPieChart.canvas.axes.axis('equal')
            self.dlg.wdg_DashboardPieChart.canvas.draw()
        else:
            self.dlg.wdg_DashboardPieChart.canvas.axes.clear()
            self.dlg.wdg_DashboardPieChart.canvas.axes.pie([1], colors = ['blue'])
            self.dlg.wdg_DashboardPieChart.canvas.draw()
        #self.dlg.show()  
        return    
    def ConnectObjects(self):

        #Maximization Buttons
        self.dlg.cmd_MoveUpMaximizedProduct.clicked.connect(self.MoveUp) 
        self.dlg.cmd_MoveDownMaximizedProduct.clicked.connect(self.MoveDown)
        self.dlg.cmd_ClearMaximizedProducts.clicked.connect(self.ClearProducts) 
        self.dlg.cmd_FuelGasMax.clicked.connect(self.FuelGasMax)
        self.dlg.cmd_LPGMax.clicked.connect(self.LPGMax)
        self.dlg.cmd_LightNaphthaMax.clicked.connect(self.LightNaphthaMax)
        self.dlg.cmd_HeavyNaphthaMax.clicked.connect(self.HeavyNaphthaMax)
        self.dlg.cmd_SRNaphthaMax.clicked.connect(self.SRNaphthaMax)
        self.dlg.cmd_GasolineMax.clicked.connect(self.GasolineMax)
        self.dlg.cmd_KeroseneMax.clicked.connect(self.KeroseneMax)
        self.dlg.cmd_JetAMax.clicked.connect(self.JetAMax)
        self.dlg.cmd_DieselMax.clicked.connect(self.DieselMax)
        self.dlg.cmd_HSDieselMax.clicked.connect(self.HSDieselMax)
        self.dlg.cmd_ULSDieselMax.clicked.connect(self.ULSDieselMax)
        self.dlg.cmd_MDOMax.clicked.connect(self.MDOMax)
        self.dlg.cmd_LSMDOMax.clicked.connect(self.LSMDOMax)
        self.dlg.cmd_AGOMax.clicked.connect(self.AGOMax)
        self.dlg.cmd_ATBMax.clicked.connect(self.ATBMax)
        self.dlg.cmd_LVGOMax.clicked.connect(self.LVGOMax)
        self.dlg.cmd_HVGOMax.clicked.connect(self.HVGOMax)
        self.dlg.cmd_VTBMax.clicked.connect(self.VTBMax)

        #Blinded Side Draw Buttons
        self.dlg.cmd_KeroseneBlindedSideDraw.clicked.connect(self.KeroseneBlindedSideDraw)
        self.dlg.cmd_JetABlindedSideDraw.clicked.connect(self.JetABlindedSideDraw)
        self.dlg.cmd_DieselBlindedSideDraw.clicked.connect(self.DieselBlindedSideDraw)
        self.dlg.cmd_ULSDieselBlindedSideDraw.clicked.connect(self.ULSDieselBlindedSideDraw)
        self.dlg.cmd_HSDieselBlindedSideDraw.clicked.connect(self.HSDieselBlindedSideDraw)
        self.dlg.cmd_LSMDOBlindedSideDraw.clicked.connect(self.LSMDOBlindedSideDraw)
        self.dlg.cmd_MDOBlindedSideDraw.clicked.connect(self.MDOBlindedSideDraw)
        self.dlg.cmd_AGOBlindedSideDraw.clicked.connect(self.AGOBlindedSideDraw)
        self.dlg.cmd_ClearBlindedSideDraws.clicked.connect(self.ClearBlindedSideDraws)

        # assay selector commands
        self.ShowRegionOptions()
        self.dlg.cmb_AssayRegion.activated.connect(self.ShowAssayOptions)
        self.dlg.cmb_AssayName.activated.connect(self.ShowSourceOptions)
        self.dlg.cmd_Add.clicked.connect(self.AddAssay)
        self.dlg.cmd_ClearAll.clicked.connect(self.ClearAssays)
        self.dlg.cmd_SelectRefinery.clicked.connect(self.SelectRefinery)
        self.dlg.list_assays.itemActivated.connect(self.SelectRefinery)
        self.dlg.cmd_Clear.clicked.connect(self.ClearRefinery)

        ###Joe: price index date selector commands
        self.ShowPriceIndexOptions()
        self.dlg.cmb_IndexDate.activated.connect(self.ShowPriceIndexOptions)

        # general information commands
        self.dlg.cmb_contact.currentIndexChanged.connect(self.OtherContact)
        self.dlg.cmd_Browse.clicked.connect(self.File_Select)
        self.dlg.btn_domestic.toggled.connect(self.DomesticSelected)
        self.dlg.btn_international.toggled.connect(self.InternationalSelected)
        self.dlg.btn_budgetary.toggled.connect(self.BudgetarySelected)
        self.dlg.btn_fixed.toggled.connect(self.FixedSelected)

        # function commands
        self.dlg.cmd_createUnit.clicked.connect(self.ProductsBreakdownCreateUnit)
        self.dlg.cmd_createAllUnits.clicked.connect(self.ProductsBreakdownCreateAllUnits)
        self.dlg.cmd_Save.clicked.connect(self.SaveProductInfo)
        self.dlg.cmd_SaveAllUnits.clicked.connect(self.SaveAllProductInfo)
        self.dlg.cmd_Generate.clicked.connect(self.GenerateClicked)
        self.dlg.cmd_processUnitsCreateUnit.clicked.connect(self.CreateCustomUnitList)
        self.dlg.cmd_clearUnit.clicked.connect(self.ClearCustomUnit)

        # assay table commands
        self.dlg.btn_InputDataToDatabase.clicked.connect(self.InputDataToDatabase)
        self.dlg.btn_DeleteDataFromDatabase.clicked.connect(self.DeleteDataFromDatabase)
        self.dlg.cmd_insertRow.clicked.connect(self.AddRow)
        self.dlg.cmd_removeRow.clicked.connect(self.RemoveRow)

        #Process Unit Breakdown Page 
        self.dlg.cmb_CurrentlySelectedRefinery.activated.connect(self.CurrentlySelectedRefineryCombobox) #This activates the Unit Operation combobox below it and fills out Calculated Process Units
        self.dlg.cmb_UnitOperation.activated.connect(self.CurrentlySelectedUnitOperationCombobox) #This fills out the text boxes with the correct Utilities
        self.dlg.cmb_CustomCurrentlySelectedRefinery.activated.connect(self.CustomCurrentlySelectedRefineryCombobox)
        self.dlg.cmb_CustomUnitOperation.activated.connect(self.CustomCurrentlySelectedUnitOperationCombobox)

        #Upload Method on Assay Viewer Page
        self.dlg.cmb_SourceName.activated.connect(self.UpdateAssayViewer)
        self.dlg.btn_ocr.toggled.connect(self.OcrSelected)
        self.dlg.btn_UploadFromDatabase.toggled.connect(self.UploadFromDatabaseSelected)
        self.dlg.cmd_UploadUsingOCR.clicked.connect(self.UploadUsingOCR)
        self.dlg.cmd_GenerateCHEMEXTemplate.clicked.connect(self.GenerateCHEMEXTemplate)

        #Blending Manager/Sandbox Page
        self.dlg.cmb_SettingsSelectedRefinery.activated.connect(self.SettingsSelectedRefinery)
        self.dlg.cmd_SaveSettingsForCurrentRefinery.clicked.connect(self.SaveSettingsForCurrentRefinery)
        self.dlg.cmd_SaveSettingsForAllRefineries.clicked.connect(self.SaveSettingsForAllRefineries)

        return
    def BuildPopup(self):
        self.exPopup = Popup(self)
        self.exPopup.show()
        return




class Popup(QMainWindow):
    
    def __init__(self,DashboardInstance):

        parent=None
        super(Popup,self).__init__(parent)
        self.setWindowFlags(QtCore.Qt.WindowStaysOnTopHint)

        #super().__init__()
        self.MyDashboard = DashboardInstance
        resolution = QDesktopWidget().screenGeometry()
        self.List_Of_Assays = self.MyDashboard.List_Of_Assay_Sources
        self.setGeometry((resolution.width() / 2) - (self.frameSize().width() / 2),
                  (resolution.height() / 2) - (self.frameSize().height() / 2),700,700)
        lblName = QLabel('Assay List', self)
        self.RepeatAssaysList = self.CreateAssayList()
        self.ReplaceButton = self.CreateReplaceButton()
        self.AddToAssayDatabase = self.CreateAddToAssayDatabaseButton()
        self.ReplacePressed = False
        self.AddAssayPressed = False
        self.SourceNumber = 1
        
    def CreateAssayList(self):
        RepeatAssaysList = QListWidget(self)
        RepeatAssaysList.setGeometry(100,60,500,500)
        for i in range(0,len(self.MyDashboard.List_Of_Assay_Sources)):
            RepeatAssaysList.addItem(self.MyDashboard.List_Of_Assay_Sources[i][0])
        return RepeatAssaysList
    def CreateReplaceButton(self):
        ReplaceButton = QPushButton("Replace Assay",self)
        ReplaceButton.clicked.connect(self.ReplaceButton)
        ReplaceButton.resize(200,60)
        ReplaceButton.move(100,580)
        ReplaceButton.show()
        return ReplaceButton
    def CreateAddToAssayDatabaseButton(self):
        AddToAssayDatabaseButton = QPushButton("Add To Assay Database",self)
        AddToAssayDatabaseButton.clicked.connect(self.AddToAssayDatabaseButton)
        AddToAssayDatabaseButton.resize(200,60)
        AddToAssayDatabaseButton.move(400,580)
        AddToAssayDatabaseButton.show()
        return AddToAssayDatabaseButton

    def ReplaceButton(self):
        if len(self.RepeatAssaysList.selectedItems()) == 0:
            QtWidgets.QMessageBox.information(self, 'Error', 'Please select an assay to replace', QtWidgets.QMessageBox.Ok)
            return
        self.ReplacePressed = True

        self.DeleteDataFromDatabase()
        self.InputDataToDatabase()

        QtWidgets.QMessageBox.information(self, 'Success!', 'Assay Successfully Replaced', QtWidgets.QMessageBox.Ok)
        self.ReplacePressed = False
        self.close()
        return
    def AddToAssayDatabaseButton(self):
        self.AddAssayPressed = True
        self.InputDataToDatabase()
        QtWidgets.QMessageBox.information(self, 'Success!', 'Assay Successfully Added to Database', QtWidgets.QMessageBox.Ok)
        self.AddAssayPressed = False
        self.close()
        return
    def closeEvent(self, *args, **kwargs):
        #super(QMainWindow, self).closeEvent(*args, **kwargs)
        #print ("you just closed the pyqt window!!! you are awesome!!!")
        return

    def InputDataToDatabase(self):

        #https://stackoverflow.com/questions/3501382/checking-whether-a-variable-is-an-integer-or-not
        if self.AddAssayPressed == True:
            for i in range(self.RepeatAssaysList.count()):
                SourceComponents = re.findall('\<(.*?)\>',self.RepeatAssaysList.item(i).text())
                try:
                    SourceComponents[-1] = int(SourceComponents[-1])
                except:
                    print("LAST ARGUMENT NOT AN INTEGER")
                    continue
                if isinstance(SourceComponents[-1], int) and SourceComponents[-1] // 1000 < 1: #The second condition is to ensure we dont have a year as a final input ex: 2019
                    if SourceComponents[-1] >= self.MyDashboard.SourceNumber:
                        self.MyDashboard.SourceNumber = SourceComponents[-1] + 1

        elif self.ReplacePressed == True:
            SourceComponents = re.findall('\<(.*?)\>',self.RepeatAssaysList.selectedItems()[0].text())
            try:
                SourceComponents[-1] = int(SourceComponents[-1])
            except:
                self.MyDashboard.SourceNumber = 1
                print("Didnt Work")
            if isinstance(SourceComponents[-1], int) and SourceComponents[-1] // 1000 < 1: #The second condition is to ensure we dont have a year as a final input ex: 2019
                self.MyDashboard.SourceNumber = SourceComponents[-1]
            else:
                self.MyDashboard.SourceNumber = 1
        self.MyDashboard.PopupCheck = True
        self.MyDashboard.InputDataToDatabase()
        return      
    def DeleteDataFromDatabase(self):
        #This function inserts assay information into the database
        self.MyDashboard.MyDatabase.ConnectToDatabase()

        #This next section checks the Source Name combo box to see if the data that is currently on the table will create a new entry in the database, or should prompt the user to select Replace/Add New
        
        #At this point we need to check if there are mutliple instances of the same assay
        #https://dba.stackexchange.com/questions/117609/looking-for-simple-contains-method-when-searching-text-in-postgresql
        print(add_quotes(self.RepeatAssaysList.selectedItems()[0].text()))
        self.MyDashboard.MyDatabase.cursor.execute("DELETE FROM " + self.MyDashboard.MyDatabase.tableName + " WHERE assaysource = " + add_quotes(self.RepeatAssaysList.selectedItems()[0].text()) + " AND assayname LIKE " + "'%" + to_str(self.MyDashboard.dlg.txt_assay.text()) + "%'" + " AND assayregion LIKE " + "'%" + to_str(self.MyDashboard.dlg.txt_region.text()) + "%'" + ";")
        self.MyDashboard.MyDatabase.connection.commit()
        #self.List_Of_Assay_Sources = self.MyDatabase.cursor.fetchall()
        self.MyDashboard.MyDatabase.DisconnectFromDatabase()
        return


#Dont mind this. This is an attempt (not yet successful) to get the UI to load properly on lower resolution screens
if hasattr(QtCore.Qt, 'AA_EnableHighDpiScaling'):
    QtWidgets.QApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling, True)
if hasattr(QtCore.Qt, 'AA_UseHighDpiPiximaps'):
    QtWidgets.QApplication.setAttribute(QtCore.Qt.AA_UseHighDpiPixmaps, True)

#5 Lines of code that are actually called when you run the script. The rest is object-oriented
app = QtWidgets.QApplication([])
dlg = uic.loadUi("ape_dashboard.ui")
MyDashboard = Dashboard(dlg)
dlg.showMaximized()
app.exec()



